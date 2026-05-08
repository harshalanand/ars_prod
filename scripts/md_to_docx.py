"""Convert LISTING_PROCESS_SOP.md to a styled .docx file."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


SRC = Path(r"e:/ARS/docs/LISTING_PROCESS_SOP.md")
DST = Path(r"e:/ARS/docs/LISTING_PROCESS_SOP.docx")


# ─── Style helpers ─────────────────────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "888888")
        borders.append(b)
    tbl_pr.append(borders)


def add_runs_with_inline(paragraph, text, base_bold=False):
    """Parse `code`, **bold**, *italic*, [text](url) and add runs to paragraph."""
    # Mask links first: replace [t](u) with placeholder to simplify
    pattern = re.compile(
        r"(\[[^\]]+\]\([^)]+\))"          # markdown links
        r"|(`[^`]+`)"                      # code
        r"|(\*\*[^*]+\*\*)"                # bold
        r"|(\*[^*]+\*)"                    # italic
    )
    pos = 0
    for m in pattern.finditer(text):
        if m.start() > pos:
            run = paragraph.add_run(text[pos:m.start()])
            if base_bold:
                run.bold = True
        token = m.group(0)
        if token.startswith("["):
            link_m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_m:
                run = paragraph.add_run(link_m.group(1))
                run.font.color.rgb = RGBColor(0x0B, 0x57, 0xD0)
                run.underline = True
                if base_bold:
                    run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token.strip("`"))
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif token.startswith("**"):
            run = paragraph.add_run(token.strip("*"))
            run.bold = True
        elif token.startswith("*"):
            run = paragraph.add_run(token.strip("*"))
            run.italic = True
        pos = m.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        if base_bold:
            run.bold = True


def split_table_row(line):
    """Split a markdown table row, stripping leading/trailing pipes."""
    parts = line.strip().strip("|").split("|")
    return [p.strip() for p in parts]


# ─── Main parser ───────────────────────────────────────────────────────────
def md_to_docx(md_text: str, out_path: Path):
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(1.8)
        section.right_margin = Cm(1.8)

    # Default body style
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    lines = md_text.splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]

        # ── Code fence ────────────────────────────────────────────────────
        if line.lstrip().startswith("```"):
            i += 1
            buf = []
            while i < n and not lines[i].lstrip().startswith("```"):
                buf.append(lines[i])
                i += 1
            i += 1  # skip closing fence
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.4)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("\n".join(buf))
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            # light grey shading via paragraph border (simple alternative)
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "F4F4F4")
            pPr.append(shd)
            continue

        # ── Headings ──────────────────────────────────────────────────────
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            level = len(m.group(1))
            text = m.group(2).strip()
            heading = doc.add_heading(level=min(level, 4))
            run = heading.add_run("")
            add_runs_with_inline(heading, text, base_bold=False)
            for run in heading.runs:
                run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)
            i += 1
            continue

        # ── Horizontal rule ───────────────────────────────────────────────
        if re.match(r"^\s*---+\s*$", line):
            p = doc.add_paragraph()
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "BBBBBB")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── Tables ────────────────────────────────────────────────────────
        if "|" in line and i + 1 < n and re.match(r"^\s*\|?\s*[:\-]+", lines[i + 1]):
            header_cells = split_table_row(line)
            i += 2  # skip header + separator
            rows = []
            while i < n and "|" in lines[i] and lines[i].strip():
                rows.append(split_table_row(lines[i]))
                i += 1
            tbl = doc.add_table(rows=1 + len(rows), cols=len(header_cells))
            tbl.autofit = True
            set_table_borders(tbl)
            # Header
            for j, h in enumerate(header_cells):
                cell = tbl.rows[0].cells[j]
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                shade_cell(cell, "1F3A68")
                cell.text = ""
                p = cell.paragraphs[0]
                add_runs_with_inline(p, h, base_bold=True)
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                    run.font.size = Pt(10)
            # Body
            for r_idx, row in enumerate(rows, start=1):
                for j in range(len(header_cells)):
                    cell = tbl.rows[r_idx].cells[j]
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    if r_idx % 2 == 0:
                        shade_cell(cell, "F7F9FC")
                    cell.text = ""
                    p = cell.paragraphs[0]
                    val = row[j] if j < len(row) else ""
                    add_runs_with_inline(p, val)
                    for run in p.runs:
                        run.font.size = Pt(9.5)
            doc.add_paragraph()  # spacer
            continue

        # ── Bullet list ───────────────────────────────────────────────────
        m = re.match(r"^(\s*)([-*])\s+(.+)$", line)
        if m:
            indent_spaces = len(m.group(1))
            text = m.group(3).strip()
            # Checkbox style for [ ] / [x]
            cb = re.match(r"^\[([ xX])\]\s+(.+)$", text)
            if cb:
                box = "☒" if cb.group(1).lower() == "x" else "☐"
                text = f"{box}  {cb.group(2)}"
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Cm(0.6 + 0.6 * (indent_spaces // 2))
            add_runs_with_inline(p, text)
            i += 1
            continue

        # ── Numbered list ─────────────────────────────────────────────────
        m = re.match(r"^(\s*)(\d+)\.\s+(.+)$", line)
        if m:
            text = m.group(3).strip()
            p = doc.add_paragraph(style="List Number")
            add_runs_with_inline(p, text)
            i += 1
            continue

        # ── Block quote ───────────────────────────────────────────────────
        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(line[2:].strip())
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # ── Blank line ────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── Paragraph (may be multi-line until blank/heading/list/table) ──
        para_lines = [line]
        i += 1
        while i < n:
            nxt = lines[i]
            if (not nxt.strip()
                or re.match(r"^#{1,6}\s", nxt)
                or re.match(r"^(\s*)([-*])\s+", nxt)
                or re.match(r"^(\s*)\d+\.\s+", nxt)
                or nxt.lstrip().startswith("```")
                or nxt.startswith("> ")
                or ("|" in nxt and i + 1 < n
                    and re.match(r"^\s*\|?\s*[:\-]+", lines[i + 1] if i + 1 < n else ""))
                or re.match(r"^\s*---+\s*$", nxt)):
                break
            para_lines.append(nxt)
            i += 1
        p = doc.add_paragraph()
        add_runs_with_inline(p, " ".join(line.strip() for line in para_lines))

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"Saved {out_path}")


if __name__ == "__main__":
    md = SRC.read_text(encoding="utf-8")
    md_to_docx(md, DST)
