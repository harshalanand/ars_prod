"""
Generates `ARS_Listing_Allocation_Process.docx` — the full developer-facing
process document for the V2 Retail ARS listing + allocation pipeline.

Audience: developers maintaining listing.py, rule_engine_new.py,
rule_engine_pandas.py, parked_history.py, pend_alc_service.py.
Goal:     no part of the pipeline is left undocumented. Every Part, every
          stage helper, every audit token, every config flag is captured
          with file:line references so you can jump straight to the code.

Run:
    cd backend && ./venv/Scripts/python.exe scripts/build_listing_alloc_doc.py
Output:
    d:/ARS_PROD/ars_prod/ARS_Listing_Allocation_Process.docx
"""
from __future__ import annotations
import os, sys
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = r"d:/ARS_PROD/ars_prod/ARS_Listing_Allocation_Process.docx"

# ---- color palette ----
C_TITLE       = RGBColor(0x10, 0x35, 0x6B)   # deep navy
C_SECTION     = RGBColor(0x15, 0x4E, 0x9E)   # blue
C_SUB         = RGBColor(0x2E, 0x70, 0xC0)   # lighter blue
C_BODY        = RGBColor(0x33, 0x33, 0x33)
C_MUTED       = RGBColor(0x66, 0x66, 0x66)
C_ACCENT      = RGBColor(0xB8, 0x3A, 0x3A)   # red-ish, used for "watch out"
C_HEAD_FILL   = "E7EEF7"
C_RULE_FILL   = "F8F8F8"


# ---------------------------------------------------------------------------
# Low-level docx helpers
# ---------------------------------------------------------------------------
def _shade(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_heading(doc, text: str, level: int = 1) -> None:
    sizes  = {0: 28, 1: 18, 2: 14, 3: 12}
    colors = {0: C_TITLE, 1: C_SECTION, 2: C_SUB, 3: C_SUB}
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8 if level else 0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes.get(level, 12))
    r.font.color.rgb = colors.get(level, C_BODY)


def add_para(doc, text: str, *, bold: bool = False, italic: bool = False,
             size: int = 10, color: RGBColor = C_BODY,
             space_after: int = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color


def add_bullet(doc, text: str, *, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    if not p.runs:
        r = p.add_run(text)
    else:
        r = p.runs[0]
        r.text = text
    r.font.size = Pt(10)
    r.font.color.rgb = C_BODY


def add_mono(doc, text: str, *, size: int = 9) -> None:
    """Render a code/monospace block."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "D0D7DE")
        pBdr.append(b)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F6F8FA")
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_table(doc, headers: list[str], rows: list[list[str]],
              *, col_widths: list[float] | None = None) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_borders(t)
    # header row
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        _shade(cell, C_HEAD_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = C_TITLE
    # data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.color.rgb = C_BODY
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)


def add_callout(doc, text: str, *, color: RGBColor = C_SUB) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.15)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    pBdr.append(left)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F6FB")
    pPr.append(shd)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = C_BODY


def add_codepath(doc, label: str, location: str) -> None:
    """Small inline reference: 'Code: file.py : function_name (line N)'."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.15)
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = C_SUB
    r2 = p.add_run(location)
    r2.font.name = "Consolas"
    r2.font.size = Pt(9)
    r2.font.color.rgb = C_BODY


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------
def build() -> None:
    doc = Document()
    for s in doc.sections:
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)

    # =====================================================================
    # 0. TITLE + HOW TO READ
    # =====================================================================
    add_heading(doc, "ARS V2 Retail", level=0)
    add_heading(doc, "Listing + Allocation — Developer Process Reference", level=1)
    add_para(
        doc,
        f"Generated {date.today().isoformat()}  •  DB: HOPC560 / Rep_Data  •  "
        f"Audience: backend developers",
        italic=True, color=C_MUTED, size=9,
    )

    add_heading(doc, "How to read this document", level=2)
    add_para(doc,
        "This is the implementation-level walkthrough of the entire ARS pipeline. "
        "Every section names the Python function and approximate line in the "
        "source file, so you can hop straight from the doc to the code. Where a "
        "formula appears, it is the formula actually emitted by the engine "
        "(SQL UPDATE expressions or pandas vector ops). Where an audit token "
        "appears (e.g., SKIP_PRI_BROKEN), you can search the codebase for that "
        "literal to find the writer."
    )
    add_bullet(doc, "Sections 3–4 cover LISTING BUILD — Parts 1 through 8.6 inside listing.py.")
    add_bullet(doc, "Sections 5–9 cover ALLOCATION — Stage A → B → C → D inside rule_engine_new.py / rule_engine_pandas.py.")
    add_bullet(doc, "Section 10 covers the Fallback phase (F0–F5).")
    add_bullet(doc, "Section 11 covers Park → Approve / Reject inside parked_history.py.")
    add_bullet(doc, "Sections 12–16 are reference tables (status values, audit tokens, columns, config, code map).")
    add_callout(doc,
        "All numbers and live-row counts come from the most recent successful run "
        "(batch_id=20260513_094250_955, MAJ_CAT=M_W_PYJAMA, DONE in 3.88s). "
        "These are illustrative — when you debug a new run, swap them out for your own."
    )

    # =====================================================================
    # 1. PIPELINE OVERVIEW
    # =====================================================================
    add_heading(doc, "1. Pipeline overview", level=1)
    add_para(doc, "End-to-end the request travels through three subsystems on the API host:")
    add_mono(doc,
        "POST /listing/generate                                  (listing.py : line 363)\n"
        "      │\n"
        "      ▼ spawns background thread → _generate_listing_impl()  (listing.py : line 508)\n"
        "      │\n"
        "  ┌───┴──────────────────────────────────────────────────────────────┐\n"
        "  │ LISTING BUILD (listing.py — synchronous SQL, single connection)  │\n"
        "  │   Part 1     INSERT grid universe into ARS_LISTING               │\n"
        "  │   Part 2     INSERT MSA-missing rows                             │\n"
        "  │   Part 2.5   conditional clustered index                         │\n"
        "  │   Part 3.5x  enrich ACS_D, ALC_D, I_ROD, AGE, AUTO_GEN_ART_SALE  │\n"
        "  │   Part 3.54  populate RL_HOLD_QTY from hold-tracking             │\n"
        "  │   Part 3.55  populate MSA_FNL_Q + VAR_COUNT/VAR_FNL_COUNT        │\n"
        "  │   Part 3.6   OPT_TYPE classifier (RL / TBC / TBL / MIX)          │\n"
        "  │   Part 3.7   MIX aggregation                                     │\n"
        "  │   Part 4–4b  grid joins (each active grid contributes columns)   │\n"
        "  │   Part 4c    OPT_MBQ, OPT_REQ, OPT_MBQ_WH, OPT_REQ_WH            │\n"
        "  │   Part 4d    ART_EXCESS                                          │\n"
        "  │   Part 4e    per-grid REQ with excess deduction                  │\n"
        "  │   Part 5     final indexes                                       │\n"
        "  │   Part 6     store ranking → ARS_STORE_RANKING + ST_RANK         │\n"
        "  │   Part 7     ARS_LISTING_WORKING + GH/H/PRI_CT/SEC_CT/ALLOC_FLAG │\n"
        "  └───┬──────────────────────────────────────────────────────────────┘\n"
        "      │\n"
        "  ┌───┴──────────────────────────────────────────────────────────────┐\n"
        "  │ ALLOCATION  (rule_engine_new.py OR rule_engine_pandas.py)        │\n"
        "  │   Stage A   list (R01–R09) + rank (tier, OPT_PRIORITY_RANK)      │\n"
        "  │   Stage B   explode OPT → (VAR_ART × SZ); fill CONT + targets    │\n"
        "  │   Stage C   waterfall RL → TBC → TBL from shared #nre_pool       │\n"
        "  │             (with sec-grid cap, MJ_REQ cap, MBQ cap, pack-round) │\n"
        "  │   Stage D   reflect to listing; assign ALLOC_STATUS + ALLOC_SEQ  │\n"
        "  │   F0–F5     optional fallback phase (boost MBQs + re-run)        │\n"
        "  └───┬──────────────────────────────────────────────────────────────┘\n"
        "      │\n"
        "  ┌───┴──────────────────────────────────────────────────────────────┐\n"
        "  │ PARK + REVIEW  (parked_history.py)                               │\n"
        "  │   Part 8.4   snapshot_session_to_parked(session_id)              │\n"
        "  │   Part 8.5   OPT_STATUS post-alloc classifier (RL/NL/MIX/TBL)    │\n"
        "  │   Part 8.6   ARS_NL_TBL_HOLD_TRACKING schema                     │\n"
        "  │   APPROVE    promote *_PARKED → *_HISTORY, PEND_ALC, hold sync   │\n"
        "  │   REJECT     delete parked rows, revert hold snapshot, audit log │\n"
        "  └──────────────────────────────────────────────────────────────────┘\n"
    )
    add_callout(doc,
        "Listing build is synchronous on one connection. Allocation can run via two backends — "
        "sequential SQL (rule_engine_new) or per-MAJ_CAT pandas workers (rule_engine_pandas) — "
        "controlled by the allocation_mode parameter. The pandas path is the production default; "
        "see Section 16 for the trade-offs."
    )

    # =====================================================================
    # 2. THE WORKED EXAMPLE (carried through the whole doc)
    # =====================================================================
    add_heading(doc, "2. The example we'll follow throughout", level=1)
    add_para(doc,
        "One real OPT is traced from listing build through Stage C waterfall, "
        "Stage D status, and finally park. Numbers come from batch "
        "20260513_094250_955."
    )
    add_table(doc,
        ["Attribute", "Value", "Where it came from"],
        [
            ["Store (WERKS)",       "HB15",                    "grid universe (Part 1)"],
            ["MAJ_CAT",             "M_W_PYJAMA",              "grid universe"],
            ["GEN_ART_NUMBER",      "1115099069",              "grid universe"],
            ["GEN_ART_DESC",        "C-POP-PLN-BTM_ELS-R_FIT-C_PKT-5*3", "MASTER_GEN_ART"],
            ["CLR",                 "D_GRY",                   "grid universe"],
            ["RNG_SEG",             "V",                       "grid join (Part 4a)"],
            ["MACRO_MVGR",          "JGR",                     "grid join (Part 4a)"],
            ["STK_TTL",             "0",                       "store_stock SLOC roll-up (Part 3.5)"],
            ["ACS_D / ALC_D",       "14 / 9",                  "ARS_CALC_ST_MAJ_CAT (Part 3.5)"],
            ["MSA_FNL_Q",           "703",                     "ARS_MSA_GEN_ART (Part 3.55)"],
            ["RL_HOLD_QTY",         "0",                       "ARS_NL_TBL_HOLD_TRACKING (Part 3.54)"],
            ["Active sizes",        "S / M / L / XL / 2XL",    "ARS_MSA_VAR_ART (Stage B)"],
            ["OPT_TYPE",            "TBL",                     "Part 3.6 classifier"],
            ["OPT_MBQ / OPT_REQ",   "≈25 / 25",                "Part 4c"],
            ["OPT_MBQ_WH",          "≈43  (hold_days=15)",     "Part 4c"],
            ["PRI_CT% / ALLOC_FLAG","100% / 1",                "Part 7"],
            ["Stage A outcome",     "LISTED_FLAG=1 (all rules pass)", "rule_engine.Stage A"],
            ["Final outcome",       "ALLOCATED — 18 SHIP + 6 HOLD across 5 sizes", "Stage D"],
        ],
        col_widths=[1.8, 2.6, 2.4],
    )

    # =====================================================================
    # 3. LISTING BUILD — full Part-by-Part walkthrough
    # =====================================================================
    add_heading(doc, "3. Listing build — every Part inside listing.py", level=1)
    add_codepath(doc, "Entry point",
                 "backend/app/api/v1/endpoints/listing.py : _generate_listing_impl  (~line 508)")
    add_codepath(doc, "API endpoint",
                 "POST /listing/generate  (~line 363) — non-blocking; spawns background thread")
    add_para(doc,
        "Progress is published via GET /listing/sessions/{session_id} and "
        "GET /listing/alloc-progress?batch_id=… (see SessionRegistry in listing_sessions.py). "
        "The Part labels below are the comment markers used inside _generate_listing_impl — "
        "search the file for '# ──── Part X' to find each."
    )

    # ----- 3.1 PARTS 1 + 2 (universe) -----
    add_heading(doc, "3.1  Parts 1, 2, 2.5 — assemble the row universe", level=2)
    add_bullet(doc, "Part 1 (~line 826): INSERT every (WERKS × MAJ_CAT × GEN_ART_NUMBER × CLR) row from the active grids of ARS_GRID_MJ_GEN_ART into ARS_LISTING.")
    add_bullet(doc, "Part 2 (~line 859): INSERT rows that exist in the MSA missing-option pool but not yet in any grid. Source table chosen via req.msa_table.")
    add_bullet(doc, "Part 2.5 (~line 876): if INSERTed rows ≥ 5,000, build a clustered index on (MAJ_CAT, WERKS) to make subsequent UPDATEs sargable. Skipped for small datasets.")
    add_para(doc, "Latest run: 108,920 (Part 1) + 20,627 (Part 2) = 129,547 rows in ARS_LISTING.",
             italic=True, color=C_MUTED)

    # ----- 3.2 PART 3.5 enrichment -----
    add_heading(doc, "3.2  Parts 3.5, 3.5a, 3.5b, 3.5c — per-row enrichment", level=2)
    add_table(doc, ["Part", "What it populates", "Source"], [
        ["3.5  (~line 914)",  "ACS_D, ALC_D",                            "ARS_CALC_ST_MAJ_CAT + ARS_CALC_ST_ART"],
        ["3.5a (~line 916)",  "I_ROD, CLR_MIN, CLR_MAX, FOCUS_W_CAP, FOCUS_WO_CAP", "ARS_CALC_ST_MAJ_CAT"],
        ["3.5b (~line 973)",  "AUTO_GEN_ART_SALE",                        "MASTER_GEN_ART_SALE.SAL_PD"],
        ["3.5c (~line 996)",  "AGE",                                      "MASTER_GEN_ART_AGE (key = ST_CD, MAJ_CAT, GEN_ART_NUMBER, CLR)"],
    ], col_widths=[1.6, 2.6, 2.6])
    add_callout(doc,
        "ACS_D = accessories density = one OPT display quantity, NOT a daily sale figure. "
        "Velocity comes from MAX_DAILY_SALE = MAX(L-7-daily, AUTO_GEN_ART_SALE), or PER_OPT_SALE "
        "when AGE < age_threshold (default 15)."
    )

    # ----- 3.3 Parts 3.54 + 3.55 -----
    add_heading(doc, "3.3  Parts 3.54, 3.55 — hold + MSA pool", level=2)
    add_table(doc, ["Part", "What", "Detail"], [
        ["3.54 (~line 1068)", "RL_HOLD_QTY",
         "SUM(HOLD_REM) of open (IS_CLOSED=0) rows in ARS_NL_TBL_HOLD_TRACKING joined on "
         "(WERKS, MAJ_CAT, GEN_ART_NUMBER, CLR). Closed holds assumed already reflected in STK_TTL."],
        ["3.55 (~line 1105)", "MSA_FNL_Q + VAR_COUNT + VAR_FNL_COUNT",
         "MSA_FNL_Q = sum across sizes from ARS_MSA_GEN_ART. "
         "VAR_COUNT = total color-size variants; VAR_FNL_COUNT = variants with positive FNL_Q."],
    ], col_widths=[1.4, 1.8, 3.6])

    # ----- 3.4 PART 3.6 OPT_TYPE classifier -----
    add_heading(doc, "3.4  Part 3.6 — OPT_TYPE classifier", level=2)
    add_codepath(doc, "Function",
                 "_classify_opt_type()  inside listing.py (~line 1162) — emits one SQL UPDATE with a 6-branch CASE")
    add_para(doc,
        "Top-down, first match wins. threshold = req.stock_threshold_pct (default 0.6). "
        "ACS_D falls back to req.default_acs_d (default 18) when missing/zero — done inline "
        "via ISNULL(NULLIF(ACS_D, 0), …)."
    )
    add_table(doc, ["#", "Condition", "OPT_TYPE", "Plain meaning"], [
        ["1", "STK_TTL < threshold×ACS_D  AND  MSA_FNL_Q=0  AND  RL_HOLD_QTY=0", "MIX (safety)", "Nothing anywhere — discontinue."],
        ["2", "VAR_FNL_COUNT / VAR_COUNT < size_threshold  OR  VAR_FNL_COUNT < min_size_count", "MIX (sparse)", "Not enough sizes to ship — bad customer experience."],
        ["3", "(STK_TTL ≥ threshold×ACS_D  OR  RL_HOLD_QTY > 0)  AND  MSA_FNL_Q > 0", "RL", "Adequate stock, supply available — top up."],
        ["4", "0 < STK_TTL < threshold×ACS_D  AND  (MSA_FNL_Q > 0  OR  RL_HOLD_QTY > 0)", "TBC", "Some stock, below target — partial top-up."],
        ["5", "STK_TTL ≤ 0  AND  (MSA_FNL_Q > 0  OR  RL_HOLD_QTY > 0)", "TBL", "Empty store, warehouse has stock — fresh listing."],
        ["6", "default (anything else)", "MIX", "Catch-all."],
    ], col_widths=[0.3, 3.0, 1.0, 2.4])
    add_para(doc,
        "Important (2026-05-11 change): RL now requires MSA_FNL_Q > 0. A store with adequate stock "
        "but no MSA supply falls through to MIX (branch 6) — we don't want the allocator considering "
        "rows it cannot ship.", italic=True, color=C_ACCENT
    )
    add_para(doc, "Latest run distribution: MIX 100,333 → aggregated to 365 lines in Part 3.7 | TBL 16,494 | TBC 2,358 | RL 10,362.",
             italic=True, color=C_MUTED)
    add_para(doc, "Our example: STK=0 AND MSA_FNL_Q=703 → branch 5 → OPT_TYPE = TBL.", bold=True, color=C_SUB)

    # ----- 3.5 PART 3.7 MIX aggregation -----
    add_heading(doc, "3.5  Part 3.7 — MIX aggregation", level=2)
    add_codepath(doc, "Function",
                 "listing.py (~line 1269) — collapse MIX rows per mix_mode")
    add_para(doc, "MIX rows are rolled up so they don't bloat the working table. mix_mode controls the grain:")
    add_table(doc, ["mix_mode", "Aggregation grain"], [
        ["each",         "no aggregation — every MIX row preserved"],
        ["st_maj_rng",   "one MIX row per (WERKS × MAJ_CAT × RNG_SEG)"],
        ["maj_cat_rng",  "one MIX row per (WERKS × MAJ_CAT)  (default — most aggressive)"],
    ], col_widths=[1.4, 4.6])
    add_para(doc, "Numeric columns are summed; categorical columns kept as 'MIX'. Aggregated rows always have ALLOC_FLAG=0 downstream.",
             italic=True, color=C_MUTED)

    # ----- 3.6 PART 4 grid joins -----
    add_heading(doc, "3.6  Parts 4, 4a, 4b — grid joins", level=2)
    add_codepath(doc, "Function",
                 "listing.py (~line 1453 onwards) — driven by ARS_GRID_BUILDER + ARS_GRID_HIERARCHY")
    add_para(doc,
        "Every active grid registered in ARS_GRID_BUILDER joins its STK_TTL / MBQ / OPT_CNT / "
        "DISP_Q columns onto ARS_LISTING as <prefix>_<column>. Part 4a executes the JOIN per grid "
        "(direct join on the last hierarchy column avoids expensive multi-table joins). Part 4b "
        "populates PER_OPT_SALE from the grid flagged use_for_opt_sale=1."
    )
    add_table(doc, ["Grid", "Hierarchy", "grid_group"], [
        ["MJ",            "WERKS × MAJ_CAT",                       "Primary"],
        ["MJ_RNG_SEG",    "WERKS × MAJ_CAT × RNG_SEG",             "Primary"],
        ["MJ_MACRO_MVGR", "WERKS × MAJ_CAT × MACRO_MVGR",          "Primary"],
        ["MJ_MICRO_MVGR", "WERKS × MAJ_CAT × MICRO_MVGR",          "Secondary"],
        ["MJ_FAB",        "WERKS × MAJ_CAT × FAB",                 "Secondary"],
        ["MJ_CLR",        "WERKS × MAJ_CAT × CLR",                 "Secondary"],
        ["MJ_M_VND_CD",   "WERKS × MAJ_CAT × M_VND_CD",            "Secondary"],
    ], col_widths=[1.6, 3.0, 1.4])
    add_callout(doc,
        "Primary grids drive ALLOC_FLAG (PRI_CT% must reach 100). Secondary grids are informational at "
        "Stage A but cap dispatch at Stage C via the sec-grid cap (Section 9). The grid_group "
        "column in ARS_GRID_BUILDER is the single source of truth for that bucketing — don't hard-code it."
    )

    # ----- 3.7 PART 4c formulas -----
    add_heading(doc, "3.7  Part 4c — OPT_MBQ, OPT_REQ, OPT_MBQ_WH, OPT_REQ_WH", level=2)
    add_codepath(doc, "Function",
                 "listing.py (~line 1703) — emits one UPDATE per article-level field")
    add_mono(doc,
        "rate_expr   = CASE WHEN AGE < age_threshold THEN MAX(PER_OPT_SALE, L7_daily)\n"
        "                   ELSE                          MAX(L7_daily, AUTO_GEN_ART_SALE)\n"
        "              END\n"
        "\n"
        "OPT_MBQ     = ROUND(ISNULL(ACS_D, 0) + rate_expr × ISNULL(ALC_D, 0), 0)\n"
        "\n"
        "OPT_REQ     = CASE WHEN OPT_MBQ − STK_TTL > 0\n"
        "                   THEN ROUND(OPT_MBQ − STK_TTL, 0)\n"
        "                   ELSE 0 END\n"
        "\n"
        "OPT_MBQ_WH  = ROUND(\n"
        "                ISNULL(ACS_D, 0) +\n"
        "                rate_expr × (ISNULL(ALC_D, 0) + (CASE WHEN OPT_TYPE='TBL' THEN hold_days ELSE 0 END)),\n"
        "                0)\n"
        "\n"
        "OPT_REQ_WH  = MAX(OPT_MBQ_WH − STK_TTL, 0)\n"
        "\n"
        "MAX_DAILY_SALE = MAX(L7_daily, AUTO_GEN_ART_SALE)        ← used by Stage A ranking\n"
    )
    add_para(doc, "For our example: rate ≈ 1.22, ACS_D=14, ALC_D=9, OPT_TYPE=TBL, hold_days=15.")
    add_mono(doc,
        "OPT_MBQ    = ROUND(14 + 1.22 × 9,        0) ≈ 25\n"
        "OPT_REQ    = ROUND(25 − 0,               0)  = 25\n"
        "OPT_MBQ_WH = ROUND(14 + 1.22 × (9 + 15), 0) ≈ 43\n"
        "OPT_REQ_WH = MAX(43 − 0, 0)                  = 43\n"
    )
    add_callout(doc,
        "OPT_MBQ_WH only adds hold_days when OPT_TYPE='TBL' (zero-stock + MSA "
        "available — needs a warehouse reserve alongside the display set). "
        "RL / TBC / MIX rows have OPT_MBQ_WH = OPT_MBQ → no hold reserve at Stage C.",
        color=C_ACCENT,
    )

    # ----- 3.8 PART 4d ART_EXCESS -----
    add_heading(doc, "3.8  Part 4d — ART_EXCESS (article-level over-stock)", level=2)
    add_codepath(doc, "Function",
                 "listing.py (~line 1727) — fills ART_EXCESS and EXCESS_STK")
    add_mono(doc,
        "ART_EXCESS = CASE\n"
        "  WHEN OPT_TYPE = 'MIX'                                       THEN 0\n"
        "  WHEN STK_TTL − excess_multiplier × OPT_MBQ > 0\n"
        "       THEN ROUND(STK_TTL − excess_multiplier × OPT_MBQ, 0)\n"
        "  ELSE 0 END                                  -- excess_multiplier default 2.0\n"
        "EXCESS_STK = ART_EXCESS\n"
    )
    add_para(doc, "ART_EXCESS feeds the per-grid REQ deduction in Part 4e — we don't want to "
                  "request stock for grids whose underlying article is already over-stocked.")

    # ----- 3.9 PART 4e -----
    add_heading(doc, "3.9  Part 4e — per-grid REQ with excess deduction", level=2)
    add_para(doc,
        "For each active grid, <prefix>_REQ is the deficit at that grid's grain, "
        "minus the article-level excess. Stored on ARS_LISTING for use by Part 7."
    )
    add_table(doc, ["<prefix>_REQ", "Grain", "Our example"], [
        ["MJ_REQ",            "WERKS × MAJ_CAT",                   "848"],
        ["RNG_SEG_REQ",       "WERKS × MAJ_CAT × RNG_SEG ('V')",   "161"],
        ["MACRO_MVGR_REQ",    "WERKS × MAJ_CAT × MACRO_MVGR ('JGR')", "50"],
        ["MICRO_MVGR_REQ",    "WERKS × MAJ_CAT × MICRO_MVGR",      "Secondary — informational"],
        ["FAB_REQ / CLR_REQ / M_VND_CD_REQ", "Secondary grids",     "Secondary — informational"],
    ], col_widths=[2.0, 3.0, 1.5])

    # ----- 3.10 PART 5 indexes -----
    add_heading(doc, "3.10  Part 5 — final indexes", level=2)
    add_para(doc,
        "Nonclustered index on [RDC] for downstream RDC-keyed lookups (pool join). "
        "Created at ~line 1805. No-op if already exists."
    )

    # ----- 3.11 PART 6 store ranking -----
    add_heading(doc, "3.11  Part 6 — store ranking (ST_RANK)", level=2)
    add_codepath(doc, "Function",
                 "listing.py (~line 1870) — builds ARS_STORE_RANKING then UPDATEs ST_RANK onto ARS_LISTING")
    add_mono(doc,
        "Per (MAJ_CAT, WERKS):\n"
        "  FILL_RATE  = MJ_STK_TTL / MJ_MBQ                          -- fill efficiency\n"
        "  REQ_RANK   = ROW_NUMBER over MAJ_CAT ORDER BY MJ_REQ  DESC\n"
        "  FILL_RANK  = ROW_NUMBER over MAJ_CAT ORDER BY FILL_RATE ASC\n"
        "  SCORE      = REQ_RANK × req_weight + FILL_RANK × fill_weight\n"
        "  ST_RANK    = ROW_NUMBER over MAJ_CAT ORDER BY SCORE ASC      (1 = best)\n"
        "\n"
        "Defaults: req_weight = 0.4   fill_weight = 0.6\n"
    )
    add_para(doc, "Our example: HB15 ranked 53 across stores for M_W_PYJAMA.",
             italic=True, color=C_MUTED)

    # ----- 3.12 PART 7 working table + ALLOC_FLAG -----
    add_heading(doc, "3.12  Part 7 — ARS_LISTING_WORKING + ALLOC_FLAG", level=2)
    add_codepath(doc, "Function",
                 "listing.py (~line 2085) — creates the filtered working table the allocator reads")
    add_para(doc, "Working table filter — only rows with meaningful work survive:")
    add_mono(doc,
        "INSERT INTO ARS_LISTING_WORKING\n"
        "SELECT * FROM ARS_LISTING\n"
        "WHERE LISTING = 1\n"
        "  AND (MSA_FNL_Q > 0 OR RL_HOLD_QTY > 0)\n"
        "  AND OPT_REQ_WH >= 1\n"
        "  AND (VAR_FNL_COUNT × 1.0 / NULLIF(VAR_COUNT,0)) >= var_ratio_threshold\n"
    )
    add_para(doc, "Latest run: 22,123 working rows survived (out of 129,547).",
             italic=True, color=C_MUTED)
    add_para(doc, "Coverage-flag math (set per row, per grid):")
    add_mono(doc,
        "For each active grid g (Primary or Secondary):\n"
        "  GH_<g> = 1 if row belongs to g's hierarchy (else raw hierarchy flag 0/1)\n"
        "  H_<g>  = (1 if <g>_REQ > 0 else 0)  ×  GH_<g>\n"
        "\n"
        "Then per row:\n"
        "  PRI_CT%   = ROUND(Σ H_<g_in_Primary>   / NULLIF(Σ GH_<g_in_Primary>, 0)   × 100, 1)\n"
        "  SEC_CT%   = ROUND(Σ H_<g_in_Secondary> / NULLIF(Σ GH_<g_in_Secondary>, 0) × 100, 1)\n"
        "  ALLOC_FLAG = CASE WHEN ISNULL(PRI_CT%, 0) >= 100 THEN 1 ELSE 0 END\n"
        "\n"
        "Our example:\n"
        "  GH_MJ=1, GH_RNG_SEG=1, GH_MACRO_MVGR=1     (row covers all 3 Primary grids)\n"
        "  H_MJ=1,  H_RNG_SEG=1,  H_MACRO_MVGR=1     (each has REQ > 0)\n"
        "  PRI_CT% = (1+1+1) / (1+1+1) × 100 = 100%\n"
        "  ALLOC_FLAG = 1                            → eligible for allocation\n"
    )
    add_callout(doc,
        "ALLOC_FLAG is a SUGGESTION, not a gate. Stage A applies R06 which only ENFORCES PRI_CT≥100 "
        "for OPT_TYPE='TBL' by default. RL/TBC rows can list even when PRI_CT<100 — they're capped "
        "later by MBQ-cap (rl_mbq_cap_pct / tbc_mbq_cap_pct) at Stage C."
    )

    # ----- 3.13 PART 8.4 / 8.5 / 8.6 -----
    add_heading(doc, "3.13  Parts 8.4, 8.5, 8.6 — post-alloc steps", level=2)
    add_table(doc, ["Part", "What", "Code"], [
        ["8.4 (~line 2198)", "Snapshot listing + alloc to *_PARKED",
         "parked_history.snapshot_session_to_parked(session_id) — non-fatal if it errors"],
        ["8.5 (~line 2247)", "OPT_STATUS post-alloc reclassifier",
         "CASE: RL→RL ; TBC alloc>0 & (STK+ALC)≥t×ACS_D→RL else MIX ; TBL alloc>0 & ≥t→NL else TBL. "
         "Sets TBL_LISTED_DATE=GETDATE() when TBL→NL transition happens."],
        ["8.6 (~line 2349)", "Hold-tracking schema ensure",
         "Creates/alters ARS_NL_TBL_HOLD_TRACKING (key WERKS+VAR_ART+SZ). Data writes happen at approve, not here."],
    ], col_widths=[1.3, 2.0, 3.6])
    add_callout(doc,
        "OPT_STATUS is the POST-allocation label and lives on the WORKING table. It is what the "
        "review UI displays. The pre-allocation OPT_TYPE never changes — it is the input to Stage A."
    )

    # =====================================================================
    # 4. ALLOCATION OVERVIEW + STAGE A
    # =====================================================================
    add_heading(doc, "4. Allocation — Stage A (list + rank)", level=1)
    add_codepath(doc, "Sequential backend",
                 "backend/app/services/rule_engine_new.py : _stage_a_apply_rules (~line 276), _stage_a_assign_rank (~line 394)")
    add_codepath(doc, "Pandas backend",
                 "backend/app/services/rule_engine_pandas.py : run_listing_and_allocation_pandas (~line 374) — calls Stage A in SQL once before fanning out")

    add_heading(doc, "4.1  Stage A rules", level=2)
    add_para(doc,
        "Each rule is gated by a module-level boolean flag (RULE_R01_LISTING, RULE_R02_NOT_MIX, …) "
        "at the top of rule_engine_new.py (~lines 31–39) — flip a flag to False to disable that "
        "rule entirely. LISTED_REASON is a concatenation of every rule that fired; "
        "LISTED_FLAG = 1 iff LISTED_REASON is empty."
    )
    add_table(doc, ["Rule", "Skip when…", "Flag", "Our example"], [
        ["R01_LISTING",        "LISTING ≠ 1",                                                         "RULE_R01_LISTING",        "pass"],
        ["R02_NOT_MIX",        "OPT_TYPE = 'MIX'",                                                    "RULE_R02_NOT_MIX",        "pass (TBL)"],
        ["R04_MSA_POS",        "MSA_FNL_Q ≤ 0 AND RL_HOLD_QTY ≤ 0",                                  "RULE_R04_MSA_POS",        "pass (703 > 0)"],
        ["R05_REQ_POS",        "OPT_REQ_WH < 1",                                                      "RULE_R05_REQ_POS",        "pass (≈43)"],
        ["R06_PRI_100",        "PRI_CT% < 100 AND ALLOC_FLAG ≠ 1 AND OPT_TYPE ∈ enforced_set",        "RULE_R06_PRI_100",        "pass (100%)"],
        ["R07_VAR_RATIO_TBL",  "OPT_TYPE='TBL' AND (VAR_FNL_COUNT/VAR_COUNT < size_threshold OR VAR_FNL_COUNT < min_size_count)", "RULE_R07_VAR_RATIO_TBL", "pass"],
        ["R08_MJ_REQ_BOOSTED", "OPT_TYPE ∈ {RL,TBC} AND (OPT_MBQ×cap_pct − MJ_STK_TTL) < ACS_SKIP_FACTOR×ACS_D  (only when PRI gate off)", "RULE_R08_MJ_REQ_BOOSTED", "n/a (TBL)"],
        ["R09_TBL_TRIVIAL",    "OPT_TYPE='TBL' AND MJ_REQ < tbl_trivial_factor × ACS_D",              "RULE_R09_TBL_TRIVIAL",    "pass (848 > 7)"],
    ], col_widths=[1.5, 3.0, 1.4, 0.8])
    add_callout(doc,
        "R06 enforced_set is computed in _stage_a_apply_rules (~line 313). It is always {TBL}; "
        "{RL} is added when pri_ct_check_rl is True; {TBC} is added when pri_ct_check_tbc is True. "
        "Default (cap mode) keeps R06 strict for TBL only, and RL/TBC fall into MBQ-cap mode at Stage C."
    )
    add_para(doc, "Latest run: 'listed=6,767  dropped=15,356  total=22,123' (~30% pass through Stage A).",
             italic=True, color=C_MUTED)

    add_heading(doc, "4.2  Stage A ranking", level=2)
    add_codepath(doc, "Tier",  "_stage_a_assign_tier  (~line 381)")
    add_codepath(doc, "Rank",  "_stage_a_assign_rank  (~line 394)")
    add_mono(doc,
        "OPT_PRIORITY_TIER = CASE\n"
        "  WHEN FOCUS_WO_CAP = 1 THEN 1     -- focus, no cap\n"
        "  WHEN FOCUS_W_CAP  = 1 THEN 2     -- focus, capped\n"
        "  ELSE                       3     -- regular\n"
        "END    -- when ENABLE_FOCUS_TIERING = False, ALL rows get tier 3\n"
        "\n"
        "OPT_PRIORITY_RANK = ROW_NUMBER() OVER (\n"
        "  PARTITION BY WERKS, OPT_TYPE, MAJ_CAT\n"
        "  ORDER BY\n"
        "    OPT_PRIORITY_TIER ASC,\n"
        "    SEC_CT%           DESC,\n"
        "    MAX_DAILY_SALE    DESC,\n"
        "    OPT_REQ_WH        DESC\n"
        ")\n"
        "\n"
        "ST_RANK is read in from ARS_LISTING (computed in Part 6).\n"
        "\n"
        "Our example:\n"
        "  OPT_PRIORITY_TIER = 3   (no focus)\n"
        "  OPT_PRIORITY_RANK = 1   (first in HB15 × M_W_PYJAMA × TBL)\n"
        "  ST_RANK           = 53\n"
    )

    # =====================================================================
    # 5. STAGE B
    # =====================================================================
    add_heading(doc, "5. Allocation — Stage B (explode OPT → size rows)", level=1)
    add_codepath(doc, "Explode",
                 "rule_engine_new.py : _stage_b_explode  (~line 565)")
    add_codepath(doc, "Contributions",
                 "rule_engine_new.py : _stage_b_fill_cont  (~line 643)")
    add_codepath(doc, "Targets",
                 "rule_engine_new.py : _stage_b_fill_targets  (~line 675)")
    add_para(doc,
        "Each LISTED OPT cross-joins ARS_MSA_VAR_ART to produce one row per "
        "(VAR_ART × SZ) into ARS_ALLOC_WORKING. CONT (size contribution) is "
        "looked up in Master_CONT_SZ; if missing, fallback = 1/distinct_sz_count."
    )
    add_mono(doc,
        "CONT          = lookup(ST_CD, MAJ_CAT, SZ)   |   fallback 1 / distinct_sz_count\n"
        "SZ_MBQ        = ROUND(OPT_MBQ    × CONT, 0)\n"
        "SZ_MBQ_WH     = ROUND(OPT_MBQ_WH × CONT, 0)\n"
        "SZ_STK        = per-size stock at store (from ARS_GRID_MJ_VAR_ART) — optional, defaults to 0\n"
        "SZ_REQ        = MAX(SZ_MBQ    − SZ_STK, 0)\n"
        "SZ_REQ_WH     = MAX(SZ_MBQ_WH − SZ_STK, 0)\n"
    )
    add_para(doc, "Latest run: 9,920 size rows in ARS_ALLOC_WORKING at Stage B start.",
             italic=True, color=C_MUTED)

    add_heading(doc, "5.1  Our example after Stage B", level=2)
    add_table(doc,
        ["VAR_ART", "SZ", "CONT", "SZ_MBQ", "SZ_STK", "SZ_REQ", "FNL_Q (pool)"],
        [
            ["1115099069030", "2XL", "0.0957", "2", "0", "2", "339"],
            ["1115099069028", "L",   "0.2239", "4", "0", "4", "128"],
            ["1115099069027", "M",   "0.2901", "6", "0", "6", "268"],
            ["1115099069026", "S",   "0.1871", "4", "0", "4", "153"],
            ["1115099069029", "XL",  "0.1172", "2", "0", "2", "154"],
            ["",              "TOTAL", "",     "18", "0", "18", ""],
        ],
        col_widths=[1.3, 0.5, 0.7, 0.8, 0.8, 0.8, 1.0],
    )

    # =====================================================================
    # 6. STAGE C
    # =====================================================================
    add_heading(doc, "6. Allocation — Stage C (the waterfall)", level=1)
    add_codepath(doc, "Sequential orchestrator",
                 "rule_engine_new.py : _stage_c_waterfall  (~line 1662)")
    add_codepath(doc, "Pandas worker",
                 "rule_engine_pandas.py : _run_majcat_waterfall  (~line 1128) — mirrors sequential statement-for-statement, in numpy")

    add_heading(doc, "6.1  Pool initialisation (#nre_pool)", level=2)
    add_mono(doc,
        "CREATE TABLE #nre_pool (\n"
        "  RDC, MAJ_CAT, GEN_ART_NUMBER, CLR, VAR_ART, SZ,\n"
        "  FNL_Q       int,        -- seeded from MAX(FNL_Q) per key\n"
        "  FNL_Q_REM   int         -- decremented as bands take units\n"
        ")\n"
        "Seeded from ARS_MSA_VAR_ART (or sometimes ARS_MSA_GEN_ART for the rolled-up grain).\n"
    )

    add_heading(doc, "6.2  Iteration order", level=2)
    add_mono(doc,
        "for opt_type in [RL, TBC, TBL]:                  # cross-opt-type ordering is fixed\n"
        "    for r in 1..I_ROD:                            # rounds per OPT (I_ROD from ARS_CALC_*)\n"
        "        for rank_band in 1..max_rank:             # BAND_SIZE = 1 (option-by-option)\n"
        "            _stage_c_run_band(...)                # cumulative pool take\n"
        "            _revalidate_after_band(...)           # update _REM; fire skip rules\n"
        "    _rerank_for_next_opt_type(...)                # skip live-low-ratio; re-rank survivors\n"
        "    _revalidate_cross_type(...)                   # write CROSS_SKIP_<ot>_* on next types\n"
        "    if fallback_mode:                             # only inside F4/F5\n"
        "        _apply_pack_round(...)\n"
        "        _apply_sec_grid_cap(...)\n"
        "_stage_c_apply_mbq_cap(...)                       # post-waterfall: RL/TBC over cap_pct of MJ_MBQ\n"
        "_stage_c_apply_mj_req_cap(...)                    # global: store-level cap on combined ship\n"
        "ALLOC_QTY := SHIP_QTY                             # finalize on alloc_table\n"
    )

    add_heading(doc, "6.3  _stage_c_run_band — the actual pool take", level=2)
    add_codepath(doc, "Function",
                 "rule_engine_new.py : _stage_c_run_band  (~line 1929)")
    add_mono(doc,
        "For this band's rows (one rank, one round, one opt_type):\n"
        "\n"
        "  need_ship  = MAX(r × SZ_MBQ − SZ_STK − SHIP_QTY, 0)\n"
        "\n"
        "  need_pool  = CASE\n"
        "                 -- TBL hold buffer (round 2+)\n"
        "                 WHEN OPT_TYPE='TBL' AND ship_already_satisfied AND hold_demand_remaining\n"
        "                   THEN SZ_MBQ_WH + (r-1)×SZ_MBQ − SZ_MBQ − SZ_STK − POOL_CONSUMED\n"
        "                 WHEN r × SZ_MBQ − SZ_STK > POOL_CONSUMED\n"
        "                   THEN r × SZ_MBQ − SZ_STK − POOL_CONSUMED\n"
        "                 ELSE 0\n"
        "               END\n"
        "\n"
        "  cumulative window across pool key (RDC, MAJ_CAT, GEN_ART, CLR, VAR_ART, SZ)\n"
        "  ORDER BY ST_RANK, OPT_PRIORITY_RANK   →  take = MIN(need_pool, FNL_Q_REM_at_window_position)\n"
        "\n"
        "  Apply take:\n"
        "    SHIP_QTY      += MIN(take, need_ship)         (ship portion of take)\n"
        "    HOLD_QTY      += MAX(take − need_ship, 0)     (TBL-only hold portion)\n"
        "    POOL_CONSUMED += take\n"
        "    FNL_Q_REM     -= take\n"
    )

    add_heading(doc, "6.4  Our example walked through Stage C", level=2)
    add_mono(doc,
        "TBL  rank 1   round 1 / 2     (HB15 / 1115099069 / D_GRY)\n"
        "  SZ    pool   take   pool_after\n"
        "  2XL   339    2      337\n"
        "  L     128    4      124\n"
        "  M     268    6      262\n"
        "  S     153    4      149\n"
        "  XL    154    2      152\n"
        "  SHIP  = 18  (matches Σ SZ_MBQ)\n"
        "\n"
        "TBL  rank 1   round 2 / 2     (hold buffer: SZ_MBQ_WH − already-shipped)\n"
        "  HOLD additions: 2XL=0 L=2 M=2 S=1 XL=1  →  HOLD = 6\n"
    )

    add_heading(doc, "6.5  _revalidate_after_band — propagating to next band", level=2)
    add_codepath(doc, "Function",
                 "rule_engine_new.py : _revalidate_after_band  (~line 840)")
    add_mono(doc,
        "After every band:\n"
        "  Step 1  MSA_FNL_Q_REM   -=  ROUND_SHIP + ROUND_HOLD             (per OPT)\n"
        "  Step 2  <grid>_REQ_REM  -=  ROUND_SHIP at grid grain            (MJ, RNG_SEG, MACRO_MVGR)\n"
        "  Step 3  H_<grid>_REM    =   (REQ_REM > ACS_SKIP_FACTOR × ACS_D) AND (GH = 1)\n"
        "  Step 4  PRI_CT_REM      =   Σ H_REM (Primary) / Σ GH (Primary) × 100\n"
        "\n"
        "Skip rules — apply to next band only — write SKIP_REASON + ALLOC_REMARKS:\n"
        "  MSA_FNL_Q_REM ≤ 0                                  → SKIP_MSA_EXHAUSTED(rem=X)\n"
        "  PRI_CT_REM    < 100  AND opt_type in enforced_set  → SKIP_PRI_BROKEN(pri_ct=X%)\n"
        "  MJ_REQ_REM    < tbl_trivial_factor × ACS_D          → SKIP_STORE_BROKEN(req_rem=X)\n"
    )

    add_heading(doc, "6.6  Cross-opt-type revalidation", level=2)
    add_codepath(doc, "Function",
                 "rule_engine_new.py : _revalidate_cross_type  (~line 1567)")
    add_para(doc,
        "After a full OPT_TYPE finishes all its rounds, any next-type rows that now fail MSA / PRI / "
        "store-broken get pre-tagged so they short-circuit before their first band runs. The remarks "
        "use CROSS_SKIP_<completed_ot>_<reason> tokens (see Section 13).")

    add_heading(doc, "6.7  Post-waterfall caps", level=2)
    add_table(doc, ["Cap", "Function", "What it trims"], [
        ["MBQ-cap (RL/TBC)",
         "_stage_c_apply_mbq_cap  (~line 1489)",
         "When pri_ct_check_<ot> is off, total SHIP of that OPT_TYPE per store ≤ <ot>_mbq_cap_pct × MJ_MBQ. "
         "Excess returned to #nre_pool. Lowest-priority rows trimmed first."],
        ["MJ_REQ cap (global)",
         "_stage_c_apply_mj_req_cap  (~line 1808)",
         "Total SHIP across all OPT_TYPEs per (WERKS, MAJ_CAT) ≤ MJ_REQ. Final clip; rows above budget "
         "zeroed (writes MJ_REQ_CAP_HIT) or partially trimmed (writes MJ_REQ_CAP_PARTIAL)."],
        ["Secondary-grid cap",
         "_apply_sec_grid_cap  (~line ~2570; sec-grid-cap module)",
         "For each Secondary grid g (MJ_MICRO_MVGR, MJ_FAB, MJ_CLR, MJ_M_VND_CD): "
         "SUM(SHIP) at g's grain ≤ g_MBQ × cap_pct. Default cap_pct=130. "
         "Trims lowest-priority first (DESC by OPT_PRIORITY_RANK, ST_RANK, GEN_ART_NUMBER)."],
        ["Pack-round (fallback only)",
         "_apply_pack_round  (~line ~2780)",
         "Tops up any partial ship up to SZ_MBQ, capped by remaining pool. "
         "Universal — no ACS_D threshold. Fires only inside F4/F5 of fallback."],
    ], col_widths=[1.5, 2.4, 3.1])

    # =====================================================================
    # 7. STAGE D
    # =====================================================================
    add_heading(doc, "7. Allocation — Stage D (reflect to listing)", level=1)
    add_codepath(doc, "Function",
                 "rule_engine_new.py : _stage_d_reflect  (~line 2101)")
    add_para(doc, "Aggregates ARS_ALLOC_WORKING up to OPT grain and writes back onto ARS_LISTING_WORKING.")
    add_mono(doc,
        "ALLOC_STATUS = CASE\n"
        "  WHEN SHIP_QTY + HOLD_QTY = 0                                  THEN 'NOT_ALLOCATED'\n"
        "  WHEN filled_size_rows < total_size_rows                       THEN 'PARTIAL'\n"
        "  WHEN SHIP_QTY + HOLD_QTY > 0 AND filled_size_rows = total_size_rows\n"
        "                                                                THEN 'ALLOCATED'\n"
        "  WHEN LISTED_FLAG = 0                                          THEN 'INELIGIBLE'\n"
        "  ELSE existing_status\n"
        "END\n"
        "ALLOC_QTY    = Σ SHIP_QTY across sizes\n"
        "HOLD_QTY     = Σ HOLD_QTY across sizes\n"
        "ALLOC_REMARKS = 'ship=X; hold=Y; sizes=N/M; seq=K' [+ audit detail from Stage C]\n"
        "\n"
        "ALLOC_SEQ = ROW_NUMBER() OVER (\n"
        "  PARTITION BY MAJ_CAT\n"
        "  ORDER BY\n"
        "    CASE OPT_TYPE WHEN 'RL' THEN 1 WHEN 'TBC' THEN 2 WHEN 'TBL' THEN 3 END,\n"
        "    MIN(ALLOC_ROUND) FROM alloc_table,\n"
        "    ST_RANK,\n"
        "    OPT_PRIORITY_RANK\n"
        ")\n"
    )

    add_heading(doc, "7.1  Our example after Stage D", level=2)
    add_table(doc,
        ["Column", "Value"],
        [
            ["ALLOC_STATUS",  "ALLOCATED"],
            ["ALLOC_QTY",     "18"],
            ["HOLD_QTY",      "6"],
            ["ALLOC_SEQ",     "1550"],
            ["ALLOC_REMARKS", "ship=18; hold=6; sizes=5/5; seq=1550"],
        ],
        col_widths=[1.8, 4.5],
    )

    # =====================================================================
    # 8. ALLOC_STATUS REFERENCE + WORKED SAMPLES
    # =====================================================================
    add_heading(doc, "8. ALLOC_STATUS reference (with live samples)", level=1)
    add_table(doc, ["Status", "Live count (latest run)", "Meaning"], [
        ["ALLOCATED",     "2,081 rows  (ship 6,982 + hold 1,019)",  "Every size hit target."],
        ["PARTIAL",       "53 rows  (ship 175 + hold 12)",          "Some sizes shipped, less than target. Usually clipped by MJ_REQ cap or MBQ cap."],
        ["SKIPPED",       "5,083 rows  (ship 0)",                   "Skipped mid-waterfall — see SKIP_REASON."],
        ["NOT_ALLOCATED", "(see by_table)",                          "Listed but no pool match at any band."],
        ["INELIGIBLE",    "(rest of listing)",                       "Failed Stage A (LISTED_FLAG=0)."],
    ], col_widths=[1.4, 2.4, 2.9])

    add_heading(doc, "8.1  Sample PARTIAL row", level=2)
    add_mono(doc,
        "Store HM37 / M_W_PYJAMA / article 1115100697 / SZ=XS / OPT_TYPE=TBC\n"
        "  SZ_MBQ=2  SZ_REQ=1  pool=233  SHIP=2  HOLD=0\n"
        "  ALLOC_STATUS = PARTIAL\n"
        "  ALLOC_REMARKS = MJ_REQ_CAP_PARTIAL(kept=2, trimmed=1, budget=25)\n"
        "→ MJ_REQ budget for HM37 × M_W_PYJAMA was only 25. This row was trimmed from\n"
        "  3 to 2 to keep the store total ≤ budget. The trimmed unit returned to #nre_pool.\n"
    )
    add_heading(doc, "8.2  Sample SKIPPED row", level=2)
    add_mono(doc,
        "Store HJ27 / M_W_PYJAMA / article 1115108826 / SZ=XL / OPT_TYPE=TBL\n"
        "  SZ_MBQ=2  SZ_REQ=2  pool=290  SHIP=0\n"
        "  ALLOC_STATUS = SKIPPED\n"
        "  SKIP_REASON  = NO_POOL_MSA       (was NO_POOL_OR_DEMAND before May-2026)\n"
        "  ALLOC_REASON = BLOCKED_NO_POOL_MSA\n"
        "  ALLOC_REMARKS includes: NO_POOL_MSA(SZ_REQ=2, SZ_MBQ=2, FNL_Q=290, FNL_Q_REM=0)\n"
        "→ By the time this OPT's rank came up, higher-priority rows had consumed the relevant\n"
        "  cumulative window, or a skip rule (SKIP_PRI_BROKEN / SKIP_MSA_EXHAUSTED) already fired.\n"
        "  If SZ_REQ had been 0, SKIP_REASON would have been 'NO_REQ' (no demand) instead.\n"
    )

    # =====================================================================
    # 9. CAPS + ROUNDING (DEEPER)
    # =====================================================================
    add_heading(doc, "9. Caps & rounding — deeper", level=1)

    add_heading(doc, "9.1  Secondary-grid dispatch cap", level=2)
    add_para(doc,
        "Toggle: req.apply_sec_cap_in_normal (default True). Runs once after each opt_type's "
        "rounds complete. Default cap_pct = 130, per-grid override via "
        "ARS_GRID_BUILDER.sec_cap_pct (blank → 130). Only grids with "
        "sec_cap_applicable = 1 participate."
    )
    add_mono(doc,
        "per Secondary grid g (in this order: MJ_MICRO_MVGR, MJ_FAB, MJ_CLR, MJ_M_VND_CD):\n"
        "  budget = g_MBQ × cap_pct / 100\n"
        "  dispatched = SUM(SHIP_QTY at g's grain)\n"
        "  if dispatched > budget:\n"
        "      trim rows lowest-priority first  (DESC by OPT_PRIORITY_RANK, ST_RANK, GEN_ART_NUMBER)\n"
        "      return trimmed units to #nre_pool.FNL_Q_REM\n"
        "      ALLOC_REMARKS += SEC_CAP_HIT(grid=g, cap_pct=X, dispatched=D, budget=B, trim_row=T)\n"
        "      SKIP_REASON   =  SEC_CAP_<g>     (only when row is fully zeroed)\n"
    )

    # Section 9.2 "Pack-round" and Section 10 "Fallback F0–F5" were
    # removed 2026-05-16 when the fallback phase itself was deleted.
    # See backend/app/docs/processes/fallback_archived.md for the design
    # frozen at the time of removal.

    add_heading(doc, "9.2  Pack-round — REMOVED 2026-05-16", level=2)
    add_para(doc,
        "Pack-round was only used inside the fallback phase. With fallback removed, "
        "this step no longer exists. See fallback_archived.md if rebuilding."
    )

    # =====================================================================
    # 10. FALLBACK — REMOVED 2026-05-16
    # =====================================================================
    add_heading(doc, "10. Fallback phase — REMOVED 2026-05-16", level=1)
    add_para(doc,
        "The fallback phase (F0–F5: MBQ boost, REQ recompute, newly-eligible OPT insert, "
        "boosted re-run waterfall, TBL Primary-grid demotion) has been removed from the engine. "
        "The orchestrator goes straight from Stage C waterfall + caps to Stage D reflect. "
        "If under-allocated stock needs to surface, raise the per-OPT-type MBQ cap sliders "
        "(rl_mbq_cap_pct / tbc_mbq_cap_pct / tbl_mbq_cap_pct) — they serve the same effect "
        "at the main pass. The full historical design lives in "
        "backend/app/docs/processes/fallback_archived.md."
    )

    # =====================================================================
    # 11. PARK → APPROVE → REJECT LIFECYCLE
    # =====================================================================
    add_heading(doc, "11. Park → Approve → Reject lifecycle", level=1)
    add_codepath(doc, "All functions in",
                 "backend/app/services/parked_history.py")

    add_heading(doc, "11.1  The 5 snapshot pairs", level=2)
    add_para(doc, "Defined in _SNAPSHOT_TARGETS (~lines 57–88). Source → Parked → History columns are tracked in the same struct.")
    add_table(doc, ["Source", "Parked", "History"], [
        ["ARS_ALLOC_WORKING",   "ARS_ALLOC_PARKED",           "ARS_ALLOC_HISTORY"],
        ["ARS_LISTING_WORKING", "ARS_LISTING_WORKING_PARKED", "ARS_LISTING_WORKING_HISTORY"],
        ["ARS_LISTING",         "ARS_LISTING_PARKED",         "ARS_LISTING_HISTORY"],
        ["ARS_MSA_GEN_ART",     "ARS_MSA_GEN_ART_PARKED",     "ARS_MSA_GEN_ART_HISTORY"],
        ["ARS_MSA_VAR_ART",     "ARS_MSA_VAR_ART_PARKED",     "ARS_MSA_VAR_ART_HISTORY"],
    ], col_widths=[2.3, 2.3, 2.3])
    add_para(doc, "Every snapshot row carries control columns: SESSION_ID, PARKED_AT (GETDATE), PARK_STATUS ('PARKED').", italic=True, color=C_MUTED)

    add_heading(doc, "11.2  SESSION_ID format", level=2)
    add_codepath(doc, "Generator",
                 "listing_sessions.py : make_session_id  (~line 129)")
    add_mono(doc,
        "session_id = datetime.now().strftime('%Y%m%d_%H%M%S_%f')[:-3]\n"
        "→ format:  YYYYMMDD_HHMMSS_mmm   (e.g., 20260513_094250_955)\n"
    )

    add_heading(doc, "11.3  Snapshot (Part 8.4)", level=2)
    add_codepath(doc, "Function",
                 "parked_history.py : snapshot_session_to_parked(session_id)  (~line 377)")
    add_bullet(doc, "Ensures *_PARKED schema for each target (idempotent ALTER TABLE column reconciliation via INFORMATION_SCHEMA).")
    add_bullet(doc, "INSERT each source's filtered rows into *_PARKED with SESSION_ID + PARKED_AT + PARK_STATUS='PARKED'.")
    add_bullet(doc, "No explicit transaction — individual inserts per target. Non-fatal if one target fails.")

    add_heading(doc, "11.4  Approve (atomic transaction)", level=2)
    add_codepath(doc, "Function",
                 "parked_history.py : approve_parked(session_id, user)  (~line 472)")
    add_para(doc, "All 8 steps run on one SQLAlchemy connection inside one commit:")
    add_table(doc, ["#", "Step", "What it does"], [
        ["1", "Ensure schemas",
         "_ensure_parked_table + _ensure_history_table for each of 5 targets."],
        ["2", "Idempotency check",
         "If every history table already has rows for this SESSION_ID → return {already_approved: True}. Prevents double-apply."],
        ["3", "Promote *_PARKED → *_HISTORY",
         "_promote_one_within_conn(conn, tgt, session_id, user) for each target. INSERTs with NOT EXISTS guard. "
         "Writes APPROVED_AT=GETDATE(), APPROVED_BY=user. Then DELETEs matching rows from *_PARKED."],
        ["4", "Write ARS_PEND_ALC",
         "write_pend_alc(conn, session_id) in pend_alc_service.py (~line 1328). "
         "Grain: (SESSION_ID, RDC, ST_CD, ARTICLE_NUMBER, ALLOC_MODE). Maps WERKS → RDC via Master_ALC_INPUT_ST_MASTER."],
        ["5", "Snapshot hold-tracking (scoped)",
         "Only rows for VAR_ARTs touched by this session (May 2026 optimization). "
         "INSERT into ARS_NL_TBL_HOLD_TRACKING_SNAPSHOT + marker row in ARS_NL_TBL_HOLD_TRACKING_SNAPSHOT_SESSIONS."],
        ["6", "Step A — RL/TBC consume hold",
         "From ARS_ALLOC_HISTORY rows with OPT_TYPE ∈ {RL, TBC}: HOLD_REM -= FROM_HOLD_QTY (floor 0); "
         "IS_CLOSED=1 + CLOSED_DATE=GETDATE() when HOLD_REM ≤ 0."],
        ["6b","Step B — TBL create new hold",
         "From ARS_ALLOC_HISTORY rows with OPT_TYPE='TBL': MERGE ARS_NL_TBL_HOLD_TRACKING; "
         "MATCHED → HOLD_QTY_INITIAL/HOLD_REM += hold_qty, IS_CLOSED=0; NOT MATCHED → INSERT new hold row."],
        ["7", "MSA delta update",
         "apply_pend_alc_delta_by_session(conn, session_id, sign=+1) in pend_alc_service.py (~line 2618). "
         "PEND_QTY += alloc_qty; FNL_Q = MAX(STK_QTY − PEND_QTY − HOLD_QTY, 0) on ARS_MSA_TOTAL / VAR_ART / GEN_ART."],
        ["8", "MSA hold sync (scoped)",
         "bootstrap_msa_hold_sync(conn, session_id) in pend_alc_service.py (~line 2735). "
         "Reseeds HOLD_QTY in MSA tables from open rows in ARS_NL_TBL_HOLD_TRACKING; recomputes FNL_Q. "
         "Scoped to (RDC, ARTICLE) keys touched by this session."],
    ], col_widths=[0.3, 1.7, 4.6])
    add_callout(doc,
        "All 8 steps share a single connection and single commit. If any step throws, the entire approve "
        "rolls back — *_PARKED remains untouched and *_HISTORY stays empty.",
        color=C_ACCENT,
    )

    add_heading(doc, "11.5  Reject (atomic transaction)", level=2)
    add_codepath(doc, "Function",
                 "parked_history.py : reject_parked(session_id, user, note=None)  (~line 670)")
    add_table(doc, ["#", "Step", "What it does"], [
        ["1", "Delete parked rows",
         "DELETE FROM *_PARKED WHERE SESSION_ID=:sid AND PARK_STATUS='PARKED' — for each of 5 targets."],
        ["2", "Revert hold-tracking",
         "_revert_hold_tracking(conn, session_id) (~line 1310). Only runs if a snapshot row exists in "
         "_HOLD_SNAPSHOT_SESSIONS for this session. "
         "DELETE rows present in live table but NOT in snapshot (rows we created via approve); "
         "UPDATE rows present in snapshot back to pre-run state. Then DELETE the session's snapshot rows."],
        ["3", "Audit log (best-effort)",
         "INSERT into System DB audit_log: action='REJECT_PARKED_ALLOC', resource_id=session_id, notes=note. "
         "Exception is swallowed — does not fail the reject."],
    ], col_widths=[0.3, 1.7, 4.6])
    add_para(doc, "Note: rejecting a session that was already approved is a no-op for hold-tracking (the snapshot doesn't exist for promoted sessions).",
             italic=True, color=C_MUTED)

    add_heading(doc, "11.6  Return shapes", level=2)
    add_mono(doc,
        "approve_parked returns:\n"
        "  {\n"
        "    approved_rows: int,                     # total across all 5 targets\n"
        "    by_table: {label: count, ...},\n"
        "    already_approved: bool,\n"
        "    error: None | str\n"
        "  }\n"
        "\n"
        "reject_parked returns:\n"
        "  {\n"
        "    rejected_rows: int,\n"
        "    by_table: {label: count, ..., 'hold_revert': {...}},\n"
        "    hold_revert: {reverted, deleted_new, restored, error},\n"
        "    error: None | str\n"
        "  }\n"
    )

    # =====================================================================
    # 12. ALLOC_STATUS / SKIP / REMARKS REFERENCE
    # =====================================================================
    add_heading(doc, "12. Audit remark catalogue (every token, where it's written)", level=1)
    add_para(doc, "Search the codebase for any of these literal strings to find the exact writer location.")
    add_table(doc, ["Token", "Stage", "Example", "Code"], [
        ["R01_LISTING",            "Stage A",        "R01_LISTING",                                              "_stage_a_apply_rules:298"],
        ["R02_NOT_MIX",            "Stage A",        "R02_NOT_MIX",                                              "_stage_a_apply_rules:300"],
        ["R04_MSA_POS",            "Stage A",        "R04_MSA_POS",                                              "_stage_a_apply_rules:305-309"],
        ["R05_REQ_POS",            "Stage A",        "R05_REQ_POS",                                              "_stage_a_apply_rules:311-312"],
        ["R06_PRI_100",            "Stage A",        "R06_PRI_100",                                              "_stage_a_apply_rules:313-324"],
        ["R07_VAR_RATIO_TBL",      "Stage A",        "R07_VAR_RATIO_TBL",                                        "_stage_a_apply_rules:325-331"],
        ["R07_SIZE_RATIO_LIVE",    "Stage C re-rank","R07_SIZE_RATIO_LIVE",                                      "_rerank_for_next_opt_type"],
        ["R08_MJ_REQ_BOOSTED",     "Stage A",        "R08_MJ_REQ_BOOSTED",                                       "_stage_a_apply_rules:343-358"],
        ["R09_TBL_TRIVIAL",        "Stage A",        "R09_TBL_TRIVIAL",                                          "_stage_a_apply_rules:333-337"],
        ["SKIP_MSA_EXHAUSTED",     "Stage C band",   "SKIP_MSA_EXHAUSTED(rem=0)",                                "_revalidate_after_band ~line 1033"],
        ["SKIP_PRI_BROKEN",        "Stage C band",   "SKIP_PRI_BROKEN(pri_ct=67%)",                              "_revalidate_after_band ~line 1038"],
        ["SKIP_STORE_BROKEN",      "Stage C band",   "SKIP_STORE_BROKEN(req_rem=3, acs_d=14)",                   "_revalidate_after_band ~line 1049"],
        ["CROSS_SKIP_<ot>_MSA",    "Cross-opt-type", "CROSS_SKIP_RL_MSA(rem=0)",                                 "_revalidate_cross_type ~line 1606"],
        ["CROSS_SKIP_<ot>_PRI",    "Cross-opt-type", "CROSS_SKIP_TBC_PRI(pri_ct=80%)",                           "_revalidate_cross_type ~line 1611"],
        ["CROSS_SKIP_<ot>_STORE_BROKEN", "Cross-opt-type", "CROSS_SKIP_RL_STORE_BROKEN(req_rem=2)",              "_revalidate_cross_type ~line 1625"],
        ["NO_REQ",                 "Stage D",        "NO_REQ — store didn't need stock (SZ_REQ<=0)",             "_stage_d_reflect"],
        ["NO_POOL_MSA",            "Stage D",        "NO_POOL_MSA — had demand, RDC pool empty at this size",    "_stage_d_reflect"],
        ["ALREADY_STOCKED",        "Stage D",        "ALREADY_STOCKED",                                          "_stage_d_reflect"],
        ["MBQ_CAP_HIT",            "Post-waterfall", "MBQ_CAP_HIT(opt_type=RL, cap_pct=110, cum_ship=120, budget=110, trim_row=10)", "_stage_c_apply_mbq_cap ~line 1538"],
        ["MBQ_CAP_PARTIAL",        "Post-waterfall", "MBQ_CAP_PARTIAL(kept=8, trimmed=2, budget=10)",            "_stage_c_apply_mbq_cap ~line 1546"],
        ["MJ_REQ_CAP_HIT",         "Post-waterfall", "MJ_REQ_CAP_HIT(cum_ship=40, budget=25, trim_row=5)",       "_stage_c_apply_mj_req_cap ~line 1871"],
        ["MJ_REQ_CAP_PARTIAL",     "Post-waterfall", "MJ_REQ_CAP_PARTIAL(kept=2, trimmed=1, budget=25)",         "_stage_c_apply_mj_req_cap ~line 1877"],
        ["SEC_CAP_HIT",            "Sec-grid cap",   "SEC_CAP_HIT(grid=MJ_FAB, cap_pct=130, dispatched=180, budget=130, trim_row=50)", "_apply_sec_grid_cap ~line 2627"],
        ["SEC_CAP_<grid>",         "Sec-grid cap (SKIP_REASON)", "SEC_CAP_MJ_FAB",                               "_apply_sec_grid_cap ~line 2645"],
        ["PACK_ROUND_SHIP",        "Fallback only",  "PACK_ROUND_SHIP(gap=2, added=2, sz_mbq=10)",               "_apply_pack_round ~line 2805"],
        ["PACK_ROUND_NO_POOL",     "Fallback only",  "PACK_ROUND_NO_POOL(gap=2, available=0, sz_mbq=10)",        "_apply_pack_round ~line 2811"],
    ], col_widths=[1.8, 1.4, 2.4, 1.4])

    # =====================================================================
    # 13. COLUMN REFERENCE — TWO WORKING TABLES
    # =====================================================================
    add_heading(doc, "13. Column reference — what to look at where", level=1)

    add_heading(doc, "13.1  ARS_LISTING_WORKING (one row per OPT)", level=2)
    add_para(doc, "Grain: (WERKS, MAJ_CAT, GEN_ART_NUMBER, CLR).", italic=True, color=C_MUTED)
    add_table(doc, ["Column", "Source / writer", "Meaning"], [
        ["OPT_TYPE",                "Part 3.6",           "RL / TBC / TBL / MIX (classifier output)"],
        ["IS_NEW",                  "Part 3.6 / upstream","1 if first-time listing this batch (informational; hold_days now gated by OPT_TYPE='TBL')"],
        ["LISTED_FLAG",             "Stage A",            "1 = passed Stage A; 0 = INELIGIBLE"],
        ["LISTED_REASON",           "Stage A",            "Concatenated rule codes (R01_LISTING; R04_MSA_POS; …) when LISTED_FLAG=0"],
        ["OPT_MBQ / OPT_REQ",       "Part 4c",            "Article-level target / gap (no hold)"],
        ["OPT_MBQ_WH / OPT_REQ_WH", "Part 4c",            "Same with hold_days for TBL"],
        ["MJ_REQ, <prefix>_REQ",    "Part 4 + 4e",        "Per-grid REQ (excess-deducted)"],
        ["GH_<grid>, H_<grid>",     "Part 7",             "Per-row grid coverage flags"],
        ["PRI_CT% / SEC_CT%",       "Part 7",             "Primary / Secondary coverage %"],
        ["ALLOC_FLAG",              "Part 7",             "1 when PRI_CT% ≥ 100"],
        ["ST_RANK",                 "Part 6",             "Store rank within MAJ_CAT (1 = best)"],
        ["OPT_PRIORITY_TIER",       "Stage A",            "1 / 2 / 3 (focus tiering)"],
        ["OPT_PRIORITY_RANK",       "Stage A",            "Position within (WERKS, OPT_TYPE, MAJ_CAT)"],
        ["OPT_STATUS",              "Part 8.5",           "Post-alloc relabel: RL / NL / MIX / TBL"],
        ["TBL_LISTED_DATE",         "Part 8.5",           "GETDATE() at TBL→NL transition"],
        ["ALLOC_STATUS",            "Stage D",            "ALLOCATED / PARTIAL / SKIPPED / NOT_ALLOCATED / INELIGIBLE"],
        ["ALLOC_QTY / HOLD_QTY",    "Stage D",            "Final SHIP / HOLD totals (sum across sizes)"],
        ["ALLOC_SEQ",               "Stage D",            "Per-OPT stable sequence number (review UI)"],
        ["ALLOC_REMARKS",           "Stage D + Stage C",  "Human-readable summary + audit detail"],
        ["FALLBACK_LVL",            "Fallback (F0)",      "0 = main pass; 1 = F1 boost; 2+ = F5 demotion levels"],
        ["*_PRE_FB",                "Fallback (F0)",      "Snapshots of MBQ / ST_RANK before boost"],
    ], col_widths=[2.4, 1.8, 3.0])

    add_heading(doc, "13.2  ARS_ALLOC_WORKING (one row per size)", level=2)
    add_para(doc, "Grain: (WERKS, MAJ_CAT, GEN_ART_NUMBER, CLR, VAR_ART, SZ).", italic=True, color=C_MUTED)
    add_table(doc, ["Column", "Writer", "Meaning"], [
        ["VAR_ART, SZ",           "Stage B explode",  "Variant article + size (pool key)"],
        ["CONT",                  "Stage B fill_cont","Size-mix proportion (Σ ≈ 1 per store)"],
        ["SZ_MBQ / SZ_REQ",       "Stage B fill_targets", "Per-size target / gap"],
        ["SZ_MBQ_WH / SZ_REQ_WH", "Stage B fill_targets", "Same with hold buffer (TBL)"],
        ["SZ_STK",                "Stage B",          "Per-size stock at store"],
        ["FNL_Q",                 "Stage C pool init","RDC pool snapshot at start of opt_type"],
        ["FNL_Q_REM",             "Stage C run_band", "Remaining pool (decremented as bands take)"],
        ["POOL_CONSUMED",         "Stage C run_band", "Cumulative units taken by this row across rounds"],
        ["SHIP_QTY",              "Stage C run_band", "Units that will dispatch"],
        ["HOLD_QTY",              "Stage C run_band", "Warehouse hold reserve (TBL only; 0 in fallback)"],
        ["FROM_HOLD_QTY",         "Stage C run_band", "Portion of SHIP that came from existing warehouse hold (RL/TBC)"],
        ["ALLOC_QTY",             "Stage D",          "= SHIP_QTY at finalize"],
        ["ALLOC_STATUS",          "Stage D",          "ALLOCATED / PARTIAL / SKIPPED / NOT_ALLOCATED"],
        ["SKIP_REASON",           "Stage C / D",      "MBQ_CAP_<otype> / MJ_REQ_CAP / SEC_CAP_<grid> / NO_REQ / NO_POOL_MSA / ALREADY_STOCKED"],
        ["ALLOC_REMARKS",         "Stage C / D",      "Detailed audit lines"],
        ["FALLBACK_LVL",          "Fallback",         "0 main / 1 F1 / 2+ F5"],
        ["ALLOC_ROUND",           "Stage C",          "Round number (1..I_ROD) in which this row's SHIP was assigned"],
        ["ALLOC_SEQ",             "Stage D",          "OPT-level sequence (mirrored from working table)"],
    ], col_widths=[2.2, 1.8, 3.2])

    # =====================================================================
    # 14. CONFIG / REQUEST FIELDS CHEAT-SHEET
    # =====================================================================
    add_heading(doc, "14. UI / request fields cheat-sheet", level=1)
    add_table(doc, ["Field", "Default", "Effect"], [
        ["apply_sec_cap_in_normal",  "True",             "Apply Sec-grid cap during main pass. Default 130%; per-grid override via ARS_GRID_BUILDER.sec_cap_pct. Only grids with sec_cap_applicable=1 participate."],
        ["pri_ct_check_rl",          "False",            "Strict PRI_CT≥100 gate for RL (default off → MBQ-cap mode)"],
        ["pri_ct_check_tbc",         "False",            "Strict PRI_CT≥100 gate for TBC"],
        ["rl_mbq_cap_pct",           "110",              "RL ship total ≤ X% of MJ_MBQ (when pri_ct_check_rl off)"],
        ["tbc_mbq_cap_pct",          "110",              "TBC ship total ≤ X% of MJ_MBQ (when pri_ct_check_tbc off)"],
        ["stock_threshold_pct",      "0.6",              "OPT_TYPE threshold: STK ≥ X × ACS_D → RL"],
        ["size_threshold",           "0.6",              "VAR_FNL_COUNT/VAR_COUNT below this → MIX (sparse rule)"],
        ["min_size_count",           "3",                "VAR_FNL_COUNT < X → MIX (hard floor)"],
        ["default_acs_d",            "18",               "Fallback when ACS_D is missing / zero"],
        ["hold_days",                "15",               "Extra days added to ALC_D for TBL hold buffer"],
        ["age_threshold",            "15",               "AGE < X → use PER_OPT_SALE in OPT_MBQ rate"],
        ["excess_multiplier",        "2.0",              "ART_EXCESS triggers when STK > X × OPT_MBQ"],
        ["tbl_trivial_factor",       "0.5",              "R09 threshold and SKIP_STORE_BROKEN factor (× ACS_D)"],
        ["acs_skip_factor",          "0.5",              "Skip rule: REQ_REM < X × ACS_D → store broken"],
        ["req_weight / fill_weight", "0.4 / 0.6",        "Part 6 store ranking weights"],
        ["allocation_mode",          "pandas",           "pandas (per-MAJ_CAT workers) vs new (sequential SQL)"],
        ["parallel_workers",         "4 (env: ARS_PARALLEL_WORKERS)", "Pandas worker count (clamped 2..8)"],
        ["use_writer_queue",         "ON",               "Serialise writes through single thread when use_pool=True"],
        ["mix_mode",                 "maj_cat_rng",      "MIX aggregation grain (each / st_maj_rng / maj_cat_rng)"],
    ], col_widths=[2.4, 1.4, 3.0])

    # =====================================================================
    # 15. PANDAS vs SEQUENTIAL
    # =====================================================================
    add_heading(doc, "15. rule_engine_new vs rule_engine_pandas", level=1)
    add_table(doc, ["Aspect", "rule_engine_new (sequential)", "rule_engine_pandas (parallel)"], [
        ["Entry",            "_stage_a … _stage_d called inline by orchestrator",
                              "run_listing_and_allocation_pandas (~line 374)"],
        ["Where Stages run", "All four stages: SQL UPDATE / INSERT on the connection",
                              "Stage A+B in SQL once; Stage C runs PER-MAJ_CAT in worker pool"],
        ["Stage C model",    "_stage_c_waterfall iterates RL→TBC→TBL globally; bands update #nre_pool in place",
                              "_run_majcat_waterfall  (~line 1128) — numpy vector ops; mirrors run_band statement-for-statement"],
        ["Parallelism",      "None (one connection, serial UPDATEs)",
                              "ProcessPoolExecutor (or threads when MAJ_CATs < 3). 2..8 workers. Stable mergesort reproduces SQL ROW_NUMBER tie-break exactly."],
        ["Writes",           "Direct per-statement UPDATE",
                              "Bulk write at end, or fed through _writer_thread_fn  (~line 278) when defer_writes=True — avoids page-lock deadlock"],
        ["Use when",         "≤5 MAJ_CATs; debugging; reproducibility checks",
                              "Production default. 10+ MAJ_CATs go from minutes to seconds."],
    ], col_widths=[1.4, 2.6, 2.8])
    add_callout(doc,
        "Both backends produce IDENTICAL allocations for the same inputs — pandas uses stable mergesort "
        "and the same ORDER BY chain. If results diverge, that's a bug in pandas (most often a "
        "tie-breaker mismatch) — diff the working tables produced by each backend on the same input."
    )

    # =====================================================================
    # 16. LATEST RUN SUMMARY
    # =====================================================================
    add_heading(doc, "16. Latest run summary (from this database)", level=1)
    add_table(doc, ["Metric", "Value"], [
        ["batch_id",                "20260513_094250_955"],
        ["status",                  "DONE"],
        ["MAJ_CATs processed",      "1  (M_W_PYJAMA — single-MAJ_CAT inline path)"],
        ["Listing rows (Part 1+2)", "129,547  (108,920 grid + 20,627 MSA-missing)"],
        ["Working rows (Part 7)",   "22,123"],
        ["Stage A listed OPTs",     "6,767  (30.6% pass through)"],
        ["Stage B size rows",       "9,920"],
        ["Total SHIP",              "10,278"],
        ["Total HOLD",              "1,643"],
        ["Total ROWS_AFFECTED",     "7,086"],
        ["Stage C duration",        "3.88 s"],
        ["ALLOC_STATUS ALLOCATED",  "2,081 rows  (ship 6,982 + hold 1,019)"],
        ["ALLOC_STATUS PARTIAL",    "53 rows  (MJ_REQ_CAP_PARTIAL audit detail visible)"],
        ["ALLOC_STATUS SKIPPED",    "5,083 rows"],
    ], col_widths=[2.2, 4.2])

    # =====================================================================
    # 17. WHERE TO LOOK IN CODE
    # =====================================================================
    add_heading(doc, "17. Where to look in code", level=1)
    add_table(doc, ["Concern", "File : function"], [
        ["API entry — POST /listing/generate",   "backend/app/api/v1/endpoints/listing.py : generate_listing (~line 363)"],
        ["Listing build orchestrator",           "listing.py : _generate_listing_impl (~line 508)"],
        ["OPT_TYPE classifier",                  "listing.py : _classify_opt_type — Part 3.6 (~line 1162)"],
        ["MIX aggregation",                      "listing.py : Part 3.7 (~line 1269)"],
        ["Grid joins",                           "listing.py : Part 4 / 4a / 4b (~line 1453 onwards)"],
        ["OPT_MBQ / OPT_REQ / OPT_MBQ_WH",       "listing.py : Part 4c (~line 1703)"],
        ["ART_EXCESS",                           "listing.py : Part 4d (~line 1727)"],
        ["Per-grid REQ",                         "listing.py : Part 4e (~line 1785)"],
        ["Store ranking (ST_RANK)",              "listing.py : Part 6 (~line 1870) + ARS_STORE_RANKING (~line 1808)"],
        ["Working table + ALLOC_FLAG",           "listing.py : Part 7 (~line 2085)"],
        ["Post-alloc OPT_STATUS",                "listing.py : Part 8.5 (~line 2247)"],
        ["Stage A rules",                        "rule_engine_new.py : _stage_a_apply_rules (~line 276)"],
        ["Stage A ranking",                      "rule_engine_new.py : _stage_a_assign_tier (~line 381), _stage_a_assign_rank (~line 394)"],
        ["Stage B explode",                      "rule_engine_new.py : _stage_b_explode (~line 565)"],
        ["Stage B contributions",                "rule_engine_new.py : _stage_b_fill_cont (~line 643)"],
        ["Stage B targets",                      "rule_engine_new.py : _stage_b_fill_targets (~line 675)"],
        ["Stage C waterfall (sequential)",       "rule_engine_new.py : _stage_c_waterfall (~line 1662)"],
        ["Stage C band",                         "rule_engine_new.py : _stage_c_run_band (~line 1929)"],
        ["Revalidate after band",                "rule_engine_new.py : _revalidate_after_band (~line 840)"],
        ["Re-rank between opt_types",            "rule_engine_new.py : _rerank_for_next_opt_type (~line 435)"],
        ["Cross-opt-type revalidate",            "rule_engine_new.py : _revalidate_cross_type (~line 1567)"],
        ["MJ_REQ cap",                           "rule_engine_new.py : _stage_c_apply_mj_req_cap (~line 1808)"],
        ["MBQ cap (RL/TBC)",                     "rule_engine_new.py : _stage_c_apply_mbq_cap (~line 1489)"],
        ["Secondary-grid cap",                   "rule_engine_new.py : _apply_sec_grid_cap (~line 2570)"],
        ["Pack-round (fallback only)",           "rule_engine_new.py : _apply_pack_round (~line 2780)"],
        ["REQ recompute chain (fallback)",       "rule_engine_new.py : _recompute_req_chain"],
        ["Fallback orchestrator (F0–F5)",        "rule_engine_new.py : _run_fallback_new (~line 3013)"],
        ["Stage D reflect",                      "rule_engine_new.py : _stage_d_reflect (~line 2101)"],
        ["Pandas entry",                         "rule_engine_pandas.py : run_listing_and_allocation_pandas (~line 374)"],
        ["Pandas per-MAJ_CAT waterfall",         "rule_engine_pandas.py : _run_majcat_waterfall (~line 1128)"],
        ["Pandas writer thread",                 "rule_engine_pandas.py : _writer_thread_fn (~line 278)"],
        ["Park snapshots",                       "parked_history.py : snapshot_session_to_parked (~line 377)"],
        ["Approve",                              "parked_history.py : approve_parked (~line 472)"],
        ["Reject",                               "parked_history.py : reject_parked (~line 670)"],
        ["Hold snapshot",                        "parked_history.py : snapshot_hold_tracking (~line 1148)"],
        ["Step A (RL/TBC consume hold)",         "parked_history.py : _apply_hold_tracking_from_history (~line 1193)"],
        ["Step B (TBL create hold)",             "parked_history.py : (~line 1245)"],
        ["Revert hold (reject)",                 "parked_history.py : _revert_hold_tracking (~line 1310)"],
        ["Write PEND_ALC",                       "pend_alc_service.py : write_pend_alc (~line 1328)"],
        ["MSA delta on approve",                 "pend_alc_service.py : apply_pend_alc_delta_by_session (~line 2618)"],
        ["MSA hold sync",                        "pend_alc_service.py : bootstrap_msa_hold_sync (~line 2735)"],
        ["Session ID generator",                 "listing_sessions.py : make_session_id (~line 129)"],
    ], col_widths=[2.6, 4.4])

    # =====================================================================
    # 18. COMMON DEV PITFALLS
    # =====================================================================
    add_heading(doc, "18. Common developer pitfalls / FAQ", level=1)
    add_table(doc, ["Symptom", "Probable cause", "Where to look"], [
        ["All rows OPT_TYPE=MIX",
         "ACS_D=0 or null everywhere AND no fallback default_acs_d, OR MSA_FNL_Q=0 globally (no warehouse stock)",
         "ARS_CALC_ST_MAJ_CAT freshness; req.default_acs_d; ARS_MSA_GEN_ART"],
        ["TBL rows listed but never ship",
         "PRI_CT% < 100 — at least one Primary grid has zero REQ for that row's hierarchy",
         "<grid>_REQ in Part 4e; ALLOC_FLAG in Part 7"],
        ["RL/TBC clipped to ~110% of MJ_MBQ",
         "Default cap mode — rl_mbq_cap_pct=110, tbc_mbq_cap_pct=110",
         "_stage_c_apply_mbq_cap — search ALLOC_REMARKS for MBQ_CAP_HIT"],
        ["ALLOC_STATUS=PARTIAL with MJ_REQ_CAP_PARTIAL remark",
         "Store hit MJ_REQ budget mid-row — partial trim. Check store's MJ_REQ vs. cumulative ship",
         "_stage_c_apply_mj_req_cap (~line 1808)"],
        ["Pandas and sequential produce different counts",
         "Tie-break mismatch in sort key, OR worker count race writing to alloc_table",
         "Stable mergesort in rule_engine_pandas.py; defer_writes / writer queue"],
        ["Approve says already_approved but UI shows parked rows",
         "Partial promotion — 1 of 5 *_HISTORY tables already populated. approve_parked is idempotent per-target",
         "approve_parked step 2; check each *_HISTORY for the session_id"],
        ["Reject leaves rows in ARS_NL_TBL_HOLD_TRACKING",
         "No snapshot existed (session was never approved, or revert ran but new rows were created outside the snapshot scope)",
         "_revert_hold_tracking; ARS_NL_TBL_HOLD_TRACKING_SNAPSHOT_SESSIONS"],
        ["MSA out of sync after manual data edits",
         "bootstrap_msa_hold_sync without session_id is unscoped (full reseed) — use to recover",
         "pend_alc_service.py : bootstrap_msa_hold_sync"],
    ], col_widths=[2.0, 3.0, 2.0])

    add_para(doc, "", space_after=12)
    add_para(doc,
        "End of document. To regenerate after code changes, edit "
        "scripts/build_listing_alloc_doc.py and re-run with backend/venv python. "
        "Line numbers will drift as code changes — they are deliberately written as "
        "approximate (~line N) so a small drift doesn't break the doc.",
        italic=True, color=C_MUTED, size=9,
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
