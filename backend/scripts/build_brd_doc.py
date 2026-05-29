"""
Convert backend/app/docs/BRD_ARS_V2.md into a styled Word document.

Run:
    cd backend && ./venv/Scripts/python.exe scripts/build_brd_doc.py
Output:
    d:/ARS_PROD/ars_prod/ARS_BRD_V2.docx
"""
from __future__ import annotations
import os
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


HERE = Path(__file__).resolve().parents[1]
SRC = HERE / "app" / "docs" / "BRD_ARS_V2.md"
OUT = Path(r"d:/ARS_PROD/ars_prod/ARS_BRD_V2.docx")

C_TITLE   = RGBColor(0x10, 0x35, 0x6B)
C_SECTION = RGBColor(0x15, 0x4E, 0x9E)
C_SUB     = RGBColor(0x2E, 0x70, 0xC0)
C_BODY    = RGBColor(0x33, 0x33, 0x33)
C_MUTED   = RGBColor(0x66, 0x66, 0x66)
C_ACCENT  = RGBColor(0xB8, 0x3A, 0x3A)
C_HEAD_FILL = "E7EEF7"


def _shade(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_heading(doc, text: str, level: int) -> None:
    sizes  = {0: 26, 1: 18, 2: 14, 3: 12, 4: 11}
    colors = {0: C_TITLE, 1: C_SECTION, 2: C_SUB, 3: C_SUB, 4: C_BODY}
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10 if level else 0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes.get(level, 11))
    r.font.color.rgb = colors.get(level, C_BODY)


_INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def _add_inline_runs(paragraph, text: str, *, base_size: int = 10,
                     base_color: RGBColor = C_BODY) -> None:
    """Render a markdown line with **bold** and `code` segments."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            r.font.size = Pt(base_size)
            r.font.color.rgb = base_color
        token = m.group(0)
        if token.startswith("**"):
            r = paragraph.add_run(token[2:-2])
            r.bold = True
            r.font.size = Pt(base_size)
            r.font.color.rgb = base_color
        else:  # `code`
            r = paragraph.add_run(token[1:-1])
            r.font.name = "Consolas"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
            r.font.size = Pt(base_size - 1)
            r.font.color.rgb = base_color
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        r.font.size = Pt(base_size)
        r.font.color.rgb = base_color


def add_para(doc, text: str, *, space_after: int = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    _add_inline_runs(p, text)


def add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(2)
    # remove default run inserted by style
    for r in list(p.runs):
        r.text = ""
    _add_inline_runs(p, text)


def add_code_block(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "D0D7DE")
        pBdr.append(b)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F6F8FA")
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_table(doc, headers: list[str], rows: list[list[str]]) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_borders(t)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        _shade(cell, C_HEAD_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        _add_inline_runs(p, h, base_size=9, base_color=C_TITLE)
        for r in p.runs:
            r.bold = True
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci] if ci < len(headers) else None
            if cell is None:
                continue
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            _add_inline_runs(p, str(val), base_size=9)


def _split_table_row(line: str) -> list[str]:
    # strip leading/trailing pipes, then split
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _is_separator_row(line: str) -> bool:
    cells = _split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", c) for c in cells if c)


def parse_and_render(md_text: str, doc) -> None:
    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]

        # fenced code block
        if line.strip().startswith("```"):
            if not in_code:
                in_code = True
                code_buf = []
            else:
                add_code_block(doc, "\n".join(code_buf))
                in_code = False
                code_buf = []
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        # blank line
        if not line.strip():
            i += 1
            continue

        # horizontal rule -> visual break
        if re.fullmatch(r"-{3,}", line.strip()):
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run("─" * 60)
            r.font.color.rgb = C_MUTED
            r.font.size = Pt(8)
            i += 1
            continue

        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", line)
        if m:
            level = len(m.group(1)) - 1  # # -> 0, ## -> 1, ...
            add_heading(doc, m.group(2).strip(), level=min(level, 4))
            i += 1
            continue

        # table: header line followed by separator line
        if "|" in line and i + 1 < len(lines) and _is_separator_row(lines[i + 1]):
            headers = _split_table_row(line)
            i += 2  # skip header + separator
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(_split_table_row(lines[i]))
                i += 1
            add_table(doc, headers, rows)
            continue

        # bullet
        if re.match(r"^\s*[-*]\s+", line):
            add_bullet(doc, re.sub(r"^\s*[-*]\s+", "", line))
            i += 1
            continue

        # numbered list -> render as plain paragraph keeping the number
        if re.match(r"^\s*\d+\.\s+", line):
            add_para(doc, line.strip())
            i += 1
            continue

        # plain paragraph
        add_para(doc, line.strip())
        i += 1


def add_footer(doc, text: str) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.size = Pt(8)
    r.font.color.rgb = C_MUTED


def main() -> None:
    md_text = SRC.read_text(encoding="utf-8")
    doc = Document()
    # tighten default margins
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # set default body font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    parse_and_render(md_text, doc)
    add_footer(doc, "ARS V2 Retail — Business Requirements Document  |  Confidential")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
