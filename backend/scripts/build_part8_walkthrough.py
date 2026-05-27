"""
Generates `ARS_Part8_Allocation_Walkthrough.docx` — a focused, step-by-step
walkthrough of Part 8 (allocation engine) and Part 8.4–8.6 (post-alloc).
Built around ONE simple worked example so a reviewer can trace every value.

Run:
    python backend/scripts/build_part8_walkthrough.py
Output:
    d:/ARS_PROD/ars_prod/ARS_Part8_Allocation_Walkthrough.docx
"""
from __future__ import annotations
import os
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = r"d:/ARS_PROD/ars_prod/ARS_Part8_Allocation_Walkthrough.docx"

C_TITLE = RGBColor(0x10, 0x35, 0x6B)
C_H1    = RGBColor(0x15, 0x4E, 0x9E)
C_H2    = RGBColor(0x2E, 0x70, 0xC0)
C_H3    = RGBColor(0x4A, 0x7B, 0xBA)
C_BODY  = RGBColor(0x22, 0x22, 0x22)
C_OK    = RGBColor(0x1B, 0x6E, 0x2D)
C_WARN  = RGBColor(0xB8, 0x3A, 0x3A)
C_MUTED = RGBColor(0x55, 0x55, 0x55)
F_HEAD  = "DCE6F2"
F_NOTE  = "FFF7DC"
F_OK    = "E8F5E9"
F_WARN  = "FDECEA"


def _shade(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    bdr = OxmlElement("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{e}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        bdr.append(b)
    tblPr.append(bdr)


def H(doc, t, level=1):
    sizes = {0: 26, 1: 18, 2: 14, 3: 12}
    cols  = {0: C_TITLE, 1: C_H1, 2: C_H2, 3: C_H3}
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10 if level else 0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(sizes.get(level, 12))
    r.font.color.rgb = cols.get(level, C_BODY)


def P(doc, text, bold=False, italic=False, size=10, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color or C_BODY


def Bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = C_BODY


def Code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def Box(doc, label, body, fill=F_NOTE, label_color=None):
    t = doc.add_table(rows=1, cols=1)
    _borders(t)
    cell = t.rows[0].cells[0]
    _shade(cell, fill)
    p = cell.paragraphs[0]
    r1 = p.add_run(label + ": ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = label_color or C_H1
    r2 = p.add_run(body)
    r2.font.size = Pt(10)
    doc.add_paragraph()


def Table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    _borders(t)
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        _shade(c, F_HEAD)
        r = c.paragraphs[0].add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = C_H1
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            r = t.rows[i].cells[j].paragraphs[0].add_run(str(v))
            r.font.size = Pt(9)
            r.font.color.rgb = C_BODY
    if widths:
        for row in t.rows:
            for j, w in enumerate(widths):
                row.cells[j].width = Inches(w)
    doc.add_paragraph()


def build():
    doc = Document()
    for s in doc.sections:
        s.left_margin  = Inches(0.7)
        s.right_margin = Inches(0.7)
        s.top_margin   = Inches(0.6)
        s.bottom_margin = Inches(0.6)

    # -------------------------------------------------------------- TITLE
    H(doc, "Part 8 & Post-Alloc — Step-by-Step Walkthrough", 0)
    P(doc, "ARS Listing & Allocation — Layman Review Edition", italic=True,
      color=C_MUTED)
    P(doc, f"Document date: {date.today().isoformat()}", italic=True,
      color=C_MUTED, size=9)
    doc.add_paragraph()

    P(doc, "This document walks through Part 8 (the allocation engine) and "
      "the post-allocation steps 8.4, 8.5, 8.6 — using ONE simple example "
      "that you can trace through every stage. The numbers are kept "
      "deliberately small so each rule's effect is visible.")

    # ============================================================ SETUP
    H(doc, "The Worked Example", 1)
    P(doc, "Imagine the warehouse has just one product to dispatch.", bold=True)

    Bullet(doc, "Product:  T-shirt 'TSH001' in 1 colour (BLUE), 5 sizes "
                "(S, M, L, XL, XXL).")
    Bullet(doc, "Warehouse pool today: 100 pieces total — 20 per size.")
    Bullet(doc, "We have 5 stores: S1, S2, S3, S4, S5.")
    Bullet(doc, "MAJ_CAT for the product = 'M_TSHIRT'.")

    P(doc, "Each store's situation:", bold=True)
    Table(doc,
          ["Store", "MJ_MBQ\n(what shop should hold)", "MJ_STK_TTL\n(what shop has now)",
           "MJ_REQ\n(=MBQ−STK)", "OPT_TYPE\n(see Part 3.6)", "Comment"],
          [
              ["S1", "60", "10", "50", "TBC",  "Has some stock, far below capacity"],
              ["S2", "40", "0",  "40", "TBL",  "Brand new option — never sold yet"],
              ["S3", "30", "25", "5",  "RL",   "Almost full, only small top-up needed"],
              ["S4", "50", "5",  "45", "TBC",  "Mostly empty, MSA supply exists"],
              ["S5", "20", "20", "0",  "—",    "Already at capacity, nothing to send"],
          ], widths=[0.5, 1.6, 1.7, 1.0, 1.2, 2.3])

    Box(doc, "ACS_D for all stores",
        "Assume ACS_D = 5 pieces / day (rough daily sale velocity). "
        "Threshold for 'low stock' = 60% × 5 = 3 pieces. "
        "Stores S1 and S4 have STK ≥ 3 but well below MBQ → TBC.",
        fill=F_NOTE)

    Box(doc, "Settings used for this run",
        "PRI≥100 strict ON for RL and TBC. Sec-grid cap ON (130%). "
        "TBL MBQ-cap = 100%. Fallback OFF.", fill=F_NOTE)

    doc.add_page_break()

    # ===================================================== PART 8 OVERVIEW
    H(doc, "Part 8 — The Allocation Engine", 1)
    P(doc, "Part 8 takes the listing table (built by Parts 1-7) and decides "
      "exactly how many pieces of each size go to each store. It runs in "
      "six sub-stages, named Stage A through Stage F.")

    Table(doc,
          ["Stage", "What it does", "Output"],
          [
              ["A", "Apply rules R01-R09 to drop ineligible OPTs.",
               "ARS_LISTED_OPT — surviving (store, article, colour) rows"],
              ["B", "Explode each surviving OPT into one row per (size).",
               "ARS_ALLOC_WORKING — size-level rows with SHIP=0"],
              ["C", "Waterfall: RL band → TBC band → TBL band; each band "
                    "ships from the warehouse pool until exhausted.",
               "SHIP_QTY filled per size row"],
              ["D", "Sec-grid cap pre-gate: block OPTs that push any "
                    "secondary grid past 130% of its MBQ.",
               "Some OPTs reset to SHIP=0 with reason"],
              ["F", "Finalise: classify reasons, write back to listing, snapshot.",
               "ALLOC_STATUS, ALLOC_REASON, parked snapshots"],
          ], widths=[0.5, 4.0, 2.5])

    # ============================================================ STAGE A
    H(doc, "Stage A — Filter eligible OPTs (rules R01–R09)", 2)
    P(doc, "Each OPT is tested top-to-bottom; the first rule it fails marks "
      "it 'SKIPPED' and the next OPT is checked.")

    Table(doc,
          ["Rule", "What it checks", "Outcome in our example"],
          [
              ["R01_LISTING",       "LISTING flag = 1",
                                    "All 5 stores published — all pass"],
              ["R02_NOT_MIX",       "OPT_TYPE ≠ 'MIX'",
                                    "S5 has no OPT_TYPE — dropped early"],
              ["R04_MSA_POS",       "MSA_FNL_Q>0 OR RL_HOLD_QTY>0",
                                    "Warehouse pool = 100 → all pass"],
              ["R05_REQ_POS",       "OPT_REQ_WH ≥ 1",
                                    "S5 fails (REQ=0); S1-S4 pass"],
              ["R06_PRI_100",       "PRI_CT% = 100 (strict mode)",
                                    "Suppose all stores have PRI_CT%=100 — pass"],
              ["R07_VAR_RATIO_TBL", "TBL: VAR_FNL/VAR_COUNT ≥ 60% OR "
                                    "VAR_FNL ≥ 3 sizes",
                                    "S2 (TBL) has 5 sizes with stock — pass"],
              ["R09_TBL_TRIVIAL",   "100% × MJ_MBQ − MJ_STK_TTL ≥ 0.5×ACS_D",
                                    "All pass (everyone needs ≥ 2.5)"],
          ], widths=[1.4, 3.0, 2.6])

    P(doc, "After Stage A: 4 OPTs survive (S1, S2, S3, S4). S5 was dropped.")

    # ============================================================ STAGE B
    H(doc, "Stage B — Explode to size grain", 2)
    P(doc, "Each surviving OPT becomes 5 rows (S, M, L, XL, XXL). We assume "
      "equal contribution — 20% each size — so each size's MBQ = OPT_MBQ × 20%.")

    P(doc, "After Stage B we have 4 stores × 5 sizes = 20 alloc rows. Each row "
      "starts with SHIP_QTY=0 and a 'pool' equal to the warehouse stock of "
      "that size (20 pieces per size).")

    # ============================================================ STAGE C
    H(doc, "Stage C — Waterfall (RL → TBC → TBL)", 2)
    P(doc, "All stores compete simultaneously inside each band. The order in "
      "which the pool drains is: highest-priority OPT first, ties broken by "
      "best-ranked store (ST_RANK).")

    P(doc, "Each store also gets a 'budget' = its MJ_REQ (because PRI≥100 "
      "strict pins the cap at 100% × MJ_MBQ − MJ_STK_TTL). The budget is "
      "rebuilt fresh at the start of every band from the LIVE MJ_REQ_REM "
      "column so all prior ships are counted automatically.")

    H(doc, "Stage C.1 — RL band (only OPTs of type RL)", 3)
    P(doc, "Only one RL OPT in our example: S3 (REQ=5).")
    Bullet(doc, "S3 needs 5 pieces. Each size row asks for 1 piece (5×20%).")
    Bullet(doc, "Pool per size = 20. Plenty of stock.")
    Bullet(doc, "S3 budget = 5. Engine ships 1 piece per size → 5 total.")
    Bullet(doc, "Pool after RL: 20 → 19 for each of S, M, L, XL, XXL.")
    Bullet(doc, "S3 MJ_REQ_REM was 5 → now 0 (budget exhausted).")

    H(doc, "Stage C.2 — TBC band (TBC OPTs only)", 3)
    P(doc, "Two TBC OPTs: S1 (REQ=50) and S4 (REQ=45). Both compete.")
    Bullet(doc, "Engine rebuilds budgets from live MJ_REQ_REM at start of band: "
                "S1 budget = 50, S4 budget = 45.")
    Bullet(doc, "Each store wants 10 per size (50×20% and 45×20% ≈ 9).")
    Bullet(doc, "Pool per size = 19. Both stores can be satisfied if they "
                "agree, but they compete in priority order.")
    Bullet(doc, "Assume S1 ranks better. S1 takes 10 per size = 50 total. "
                "Pool → 9 per size.")
    Bullet(doc, "S4 takes min(9 per size, what's left in budget). 9×5 = 45 "
                "= exactly S4's budget. S4 fully shipped.")
    Bullet(doc, "Pool after TBC: 0 per size for S/M/L; if anything left, used "
                "later by TBL.")
    Bullet(doc, "S1 MJ_REQ_REM 50 → 0; S4 MJ_REQ_REM 45 → 0.")

    H(doc, "Stage C.3 — TBL band (TBL OPTs only)", 3)
    P(doc, "Only S2 (REQ=40). Pool now nearly empty.")
    Bullet(doc, "S2 budget = 40 (its MJ_REQ).")
    Bullet(doc, "Each size needs 8 (40×20%). Pool has 0 per size.")
    Bullet(doc, "Result: S2 ships 0 — pool is dry. HOLD_QTY ≤ 40 may be "
                "reserved for the next shipment.")
    Bullet(doc, "If pool had stock, TBL would ship up to MJ_REQ exactly "
                "(because TBL MBQ-cap = 100%).")

    Box(doc, "Why this order works",
        "RL tops up known sellers first; TBC fixes break-pack stores; TBL "
        "is the speculative new-listing — it goes last so warehouse pool isn't "
        "wasted on unproven options. If S2 were a 'must list' store, planners "
        "can raise the TBL MBQ cap slider (tbl_mbq_cap_pct) or boost the "
        "TBL OPT's contribution % via the grid master.",
        fill=F_NOTE)

    H(doc, "After-each-band housekeeping (the cap engine)", 3)
    Bullet(doc, "_revalidate_after_band runs. It decrements MJ_REQ_REM "
                "and every primary grid REQ_REM at (WERKS, MAJ_CAT) grain.")
    Bullet(doc, "_pre_band_check runs before the NEXT band. It skips any OPT "
                "whose PRI_CT_REM dropped below 100 (primary grid broke).")
    Bullet(doc, "Store-broken check: if MJ_REQ_REM drops below 0.5×ACS_D, "
                "that store is skipped for the rest of the waterfall.")
    Bullet(doc, "The MJ-budget is REBUILT from the live MJ_REQ_REM at the "
                "start of every band — single source of truth, no stale dict.")

    # ============================================================ STAGE D
    H(doc, "Stage D — Sec-grid cap pre-gate (130%)", 2)
    P(doc, "After the waterfall, we look at every OPT that shipped and ask: "
      "did it push any secondary grid (colour, fabric, vendor, micro-MVGR) "
      "above 130% of THAT grid's MBQ at this store?")

    Bullet(doc, "In our example there's only one colour (BLUE) so MJ_CLR cap "
                "= MJ_MBQ. Suppose S1's BLUE post-alloc total is 60 vs "
                "MJ_CLR_MBQ=50 (boosted +20). That's 60/50 = 120% — under "
                "the 130% cap → ALLOWED.")
    Bullet(doc, "If it had been 70/50 = 140%, the OPT would be BLOCKED unless "
                "OPT_REQ ≥ 100% × OPT_MBQ (the override).")

    Table(doc,
          ["Outcome",       "What changes"],
          [
              ["ALLOWED",    "Row keeps SHIP_QTY"],
              ["OVERRIDDEN", "Row keeps SHIP_QTY; remark "
                              "'SEC_CAP_OVERRIDE_<grid>'"],
              ["BLOCKED",    "Row's SHIP_QTY reset to 0; "
                              "ALLOC_REASON='BLOCKED_SEC_CAP_PRE_<grid>'"],
          ], widths=[1.6, 5.0])

    # ============================================================ STAGE E
    # REMOVED 2026-05-16 — the optional fallback phase has been deleted.
    # See backend/app/docs/processes/fallback_archived.md for the historical
    # design.

    # ============================================================ STAGE F
    H(doc, "Stage F — Finalise reasons + write back", 2)
    Bullet(doc, "ALLOC_STATUS = ALLOCATED / PARTIAL / SKIPPED.")
    Bullet(doc, "ALLOC_REASON and SKIP_REASON populated for audit.")
    Bullet(doc, "_stage_d_reflect copies summary back to ARS_LISTING_WORKING "
                "so the UI rollup is consistent.")

    Box(doc, "Final result for our example",
        "S3: SHIP=5 (RL), fully allocated. "
        "S1: SHIP=50 (TBC), fully allocated. "
        "S4: SHIP=45 (TBC), fully allocated. "
        "S2: SHIP=0 (TBL), HOLD ≤ 40 reserved for next run. "
        "S5: SKIPPED (R05_REQ_POS, no demand). "
        "Total ship = 100 (the entire pool), 0 leftover.",
        fill=F_OK, label_color=C_OK)

    doc.add_page_break()

    # ===================================================== PART 8.4
    H(doc, "Part 8.4 — Park the snapshot", 1)
    P(doc, "Every run is reviewed by planners before the alloc actually goes "
      "to dispatch. Part 8.4 takes a frozen copy of every output table and "
      "tags it with the session_id.")

    Table(doc,
          ["Source table",        "Parked to",                 "Purpose"],
          [
              ["ARS_ALLOC_WORKING",  "ARS_ALLOC_PARKED",
                "Size-level alloc decisions"],
              ["ARS_LISTING_WORKING","ARS_LISTING_WORKING_PARKED",
                "Eligible-listing rollup with PRI%/SEC%"],
              ["ARS_LISTING",        "ARS_LISTING_PARKED",
                "Full listing including ineligible rows"],
              ["ARS_MSA_GEN_ART",    "ARS_MSA_GEN_ART_PARKED",
                "MSA supply at the time of run"],
              ["ARS_MSA_VAR_ART",    "ARS_MSA_VAR_ART_PARKED",
                "Variant-level MSA supply"],
          ], widths=[1.9, 2.2, 2.5])

    Box(doc, "Layman analogy",
        "Think of Part 8.4 as 'photocopy everything into a labelled folder' "
        "before sending the dispatch out. If the user clicks REJECT later, we "
        "just throw the folder away. If they APPROVE, the folder is promoted "
        "to history (and the warehouse hold table gets the new TBL reservations).",
        fill=F_NOTE)

    # ===================================================== PART 8.5
    H(doc, "Part 8.5 — Classify OPT_STATUS + record TBL_LISTED_DATE", 1)
    P(doc, "After alloc, we re-classify each OPT based on its post-alloc stock "
      "to drive next-run behaviour and reporting.")

    Table(doc,
          ["Pre-alloc OPT_TYPE", "Post-alloc condition",
           "New OPT_STATUS", "Layman meaning"],
          [
              ["RL",  "(always)",                             "RL",
                "Was top-up, stays a 'running line'"],
              ["TBC", "ALC>0 AND (STK+ALC) ≥ 60% × ACS_D",    "RL",
                "Break-pack restored — promote to running line"],
              ["TBC", "ALC>0 BUT total still below threshold","MIX",
                "Partial recovery — keep as MIX next run"],
              ["TBC", "ALC=0 (nothing shipped)",              "MIX",
                "Couldn't fix — group as MIX"],
              ["TBL", "ALC>0 AND (STK+ALC) ≥ 60% × ACS_D",    "NL",
                "Newly listed — now a Normal Line"],
              ["TBL", "ALC>0 BUT below threshold",            "TBL",
                "Partially listed — keep trying"],
              ["TBL", "ALC=0",                                "TBL",
                "Pool was empty — try again next run"],
          ], widths=[1.3, 2.3, 1.2, 2.0])

    P(doc, "Also: for any TBL row with ALLOC_QTY>0 AND TBL_LISTED_DATE IS "
      "NULL, set TBL_LISTED_DATE = GETDATE(). This is the day the option "
      "first reached the store — used for age calculations.")

    Box(doc, "In our example",
        "S3 stays RL (was RL). S1 (TBC, ALC=50, STK+ALC = 60, 60% × 5 = 3 → "
        "60 ≥ 3 ✅) → RL. S4 same logic → RL. S2 (TBL, ALC=0) → TBL still. "
        "S5 was MIX from the start.",
        fill=F_OK, label_color=C_OK)

    # ===================================================== PART 8.6
    H(doc, "Part 8.6 — NL/TBL hold tracking", 1)
    P(doc, "Some warehouse stock has to be 'reserved' for a TBL option that "
      "isn't shipping all sizes yet (size-curve incomplete). Otherwise, by "
      "the time we try again next week, another store will have eaten that "
      "stock.")

    P(doc, "ARS_NL_TBL_HOLD_TRACKING is a permanent table:", bold=True)
    Bullet(doc, "Grain: (WERKS, VAR_ART, SZ) — store + colour-size variant.")
    Bullet(doc, "HOLD_QTY_INITIAL = how much was reserved when the option was "
                "first listed.")
    Bullet(doc, "HOLD_REM = how much of that reservation still exists today.")
    Bullet(doc, "IS_CLOSED = 1 when the reservation is no longer needed.")

    P(doc, "Note: The actual WRITES into this table have moved out of Part "
      "8.6 — they now run when the user APPROVES the parked snapshot "
      "(parked_history.approve_parked). Part 8.6 only ensures the table "
      "exists with the right columns and indexes.", italic=True, color=C_MUTED)

    Box(doc, "Layman flow",
        "1. Run on Day 1: S2 is TBL, partial ship (3 of 5 sizes). Engine "
        "RESERVES 2 sizes × the missing qty into ARS_NL_TBL_HOLD_TRACKING. "
        "(Reservation happens on APPROVE.) "
        "2. Run on Day 8: Part 3.54 reads the table → S2 sees its prior "
        "reservation as RL_HOLD_QTY. R04 lets it pass even if MSA supply has "
        "dropped, because the warehouse is still holding stock for it. "
        "3. Run on Day 15: S2 finally completes — IS_CLOSED set to 1.",
        fill=F_NOTE)

    doc.add_page_break()

    # ===================================================== REVIEW MATRIX
    H(doc, "Review Matrix — toggle by toggle", 1)
    P(doc, "Quick lookup: what each toggle ACTUALLY does inside Part 8.")

    Table(doc,
          ["Toggle",           "Stage affected",       "Effect when ON",
           "Effect when OFF"],
          [
              ["PRI≥100 RL",     "Stage A R06; Stage C cap",
                "RL OPTs with PRI<100 SKIPPED; store×MJ cap pinned at 100%",
                "RL OPTs always pass R06; store×MJ cap = rl_mbq_cap_pct (e.g. 130)"],
              ["PRI≥100 TBC",    "Stage A R06; Stage C cap",
                "TBC OPTs with PRI<100 SKIPPED; cap pinned at 100%",
                "TBC OPTs always pass R06; cap = tbc_mbq_cap_pct"],
              ["Sec-grid Cap",   "Stage D",
                "Block OPTs that would exceed 130% × <sec-grid>_MBQ at OPT grain",
                "No secondary-grid policing — risk of colour/MVGR skew"],
              ["Enable Fallback","Stage E",
                "F1 boost + F3 re-waterfall + F4 pack_round + F5 grid demotion",
                "Stage E skipped entirely"],
              ["Fallback scope = exclude_mj",
                "Stage E F1; Stage F POST_FB trim",
                "MJ_MBQ untouched; POST_FB trim hard-caps SHIP at MJ_REQ",
                "n/a"],
              ["Fallback scope = include_mj",
                "Stage E F1; Stage F POST_FB trim",
                "MJ_MBQ also boosted 130%; POST_FB trim allows SHIP up to "
                "boosted MJ_REQ (fillrate headroom)",
                "n/a"],
              ["TBL MBQ-cap %",  "Stage C TBL band",
                "TBL bands cap at (cap × MJ_MBQ − MJ_STK_TTL)",
                "n/a — always >0"],
          ], widths=[1.4, 1.5, 2.0, 2.2])

    # ===================================================== HOLD CHECK
    H(doc, "How HOLD is decided (sanity check)", 1)
    P(doc, "HOLD_QTY in ARS_ALLOC_PARKED comes from the TBL band only:")
    Bullet(doc, "Inside _run_band TBL: if take_pool > need_ship (a row pulled "
                "more pool than that size needs), the surplus becomes HOLD_QTY "
                "for the next run.")
    Bullet(doc, "Sec-grid cap (Stage D) and the MJ-cap NEVER touch HOLD_QTY "
                "— it's set during the waterfall only.")
    Bullet(doc, "Re-running the same input gives the same HOLD_QTY iff the "
                "waterfall is deterministic. In observed runs (sessions "
                "20260514_152345_428 vs _152624_574) HOLD was identical = "
                "4569 in both → confirms TBL band is deterministic for the "
                "test dataset.")

    # ===================================================== REVIEW CHECKLIST
    H(doc, "Reviewer's Checklist", 1)

    P(doc, "After every run, verify these in order:", bold=True)
    Bullet(doc, "Engine header log shows the intended cap mode "
                "('PRI≥100 strict (MJ-cap 100%)' or 'MBQ-cap NN%').")
    Bullet(doc, "Run scripts/excess_breakdown.py <session>: under PRI strict, "
                "'Stores where SHIP > MJ_REQ' should be 0.")
    Bullet(doc, "If non-zero overshoot AND fallback was ON: this is the known "
                "POST_FB trim gap in the pandas path (see Stage E note).")
    Bullet(doc, "Total ship vs total MJ_REQ aggregate should be 85-100% "
                "depending on pool availability.")
    Bullet(doc, "HOLD_QTY should match TBL_round_2 'hold' value in the round "
                "logs.")
    Bullet(doc, "DONE log line totals: 'fb_ship + pack_round' should be ≤ "
                "the gap between main-pass ship and final ship.")
    Bullet(doc, "Sec-cap audit line: blocked + overridden counts should both "
                "be visible. If apply_sec_cap_in_normal was ON but no "
                "audit line appears, the helper failed silently.")
    Bullet(doc, "All snapshot tables (ARS_*_PARKED) have rows for this session_id.")
    Bullet(doc, "OPT_STATUS column populated on ARS_LISTING after Part 8.5.")
    Bullet(doc, "TBL_LISTED_DATE is set ONLY on TBL rows with ALLOC_QTY > 0.")

    # ===================================================== ITEMS TO REVIEW
    H(doc, "Items to confirm / improve", 1)
    Table(doc,
          ["#", "Item", "Why it matters"],
          [
              ["1", "Add _stage_c_apply_mj_req_cap(phase='POST_FB') call to "
                    "run_listing_and_allocation_pandas, immediately after "
                    "_run_fallback_new.",
                "Without it, pack-round and fallback-ship can push stores "
                "past MJ_REQ even when boost_scope='exclude_mj' is selected. "
                "Today this is silently allowed in the pandas path."],
              ["2", "Verify pack_round behaviour: should it respect MJ-cap "
                    "by itself, or rely on the POST_FB trim to clean up?",
                "Today pack_round only checks SZ_MBQ + pool — no MJ-cap "
                "check. The POST_FB trim is intended to do the clean-up."],
              ["3", "Determinism — store-tie-break in pool drain.",
                "Confirmed harmless drift in HO15↔HO24 example (3 units). "
                "Could be hardened with explicit (ST_RANK, WERKS) tiebreaker."],
              ["4", "TBL pack_round currently runs at F4 (inside fallback) "
                    "AND optionally once more in F5. Verify no double-count.",
                "If both fire, TBL stores may receive 2× the pack-round adds."],
              ["5", "OPT_STATUS classification threshold = stock_threshold_pct "
                    "(0.6). Confirm with planners.",
                "60% is the same threshold used in Part 3.6 OPT_TYPE rules — "
                "consistency is important."],
              ["6", "HOLD_QTY is captured but never decremented in Part 8 "
                    "itself — decrement happens on the next run via Part 3.54.",
                "If user runs twice in the same day, second run sees the "
                "stale HOLD. Verify this is the intended behaviour."],
              ["7", "Stage E (fallback) without Stage F (POST_FB trim) is "
                    "the current broken state for the pandas path. Decide: "
                    "fix the pandas path or disable fallback in pandas mode.",
                "Two options to close the gap; either is acceptable."],
              ["8", "Sec-cap pre-gate trims by setting SHIP=0 (all-or-nothing). "
                    "Consider switching to partial-trim (reduce to just-under-cap).",
                "All-or-nothing means valuable OPTs lose their entire ship "
                "due to a 1-unit breach. Partial-trim is gentler."],
          ], widths=[0.4, 3.6, 3.2])

    # ===================================================== TIMINGS
    H(doc, "Typical timings observed", 1)
    Table(doc,
          ["Step",                 "Small batch\n(~265 OPTs)",
           "Medium\n(7K OPTs)",    "Full universe\n(6.3M listing)"],
          [
              ["Part 8 (no fallback)",      "9 s",  "20 s",  "1.5 min+"],
              ["Part 8 (with fallback)",    "37 s", "60 s",  "3 min+"],
              ["Part 8.4 (park snapshots)", "2 s",  "3 s",   "20 s"],
              ["Part 8.5 (OPT_STATUS)",     "0.5 s","0.6 s", "5 s"],
              ["Part 8.6 (DDL guards)",     "0 s",  "0 s",   "0 s"],
          ], widths=[2.2, 1.7, 1.5, 1.7])

    doc.add_paragraph()
    P(doc, "End of walkthrough. For the full pipeline (Parts 1-7), see "
      "ARS_Listing_Allocation_Manual.docx.", italic=True, color=C_MUTED, size=9)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
