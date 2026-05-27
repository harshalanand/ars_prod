"""
Generates `ARS_Fallback_Process.docx` — the design review document for the
re-introduction of the fallback allocation phase into the ARS V2 pipeline.

Audience: project owner + developers reviewing the fallback rebuild before
          implementation begins.
Scope:    purpose, phase-by-phase mechanics (FB-0..FB-5), worked numeric
          example, grid-cap rename, knobs, pros/cons, open questions.

Run:
    cd backend && ./venv/Scripts/python.exe scripts/build_fallback_doc.py
Output:
    d:/ARS_PROD/ars_prod/ARS_Fallback_Process.docx
"""
from __future__ import annotations
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


import os as _os
_DEFAULT_OUT = r"d:/ARS_PROD/ars_prod/ARS_Fallback_Process.docx"
def _resolve_out(path: str) -> str:
    # If the target is locked (open in Word), fall back to a numbered sibling.
    try:
        if _os.path.exists(path):
            with open(path, "ab"):
                pass
        return path
    except PermissionError:
        base, ext = _os.path.splitext(path)
        for i in range(2, 20):
            cand = f"{base}_v{i}{ext}"
            try:
                if _os.path.exists(cand):
                    with open(cand, "ab"):
                        pass
                return cand
            except PermissionError:
                continue
        raise

OUT = _resolve_out(_DEFAULT_OUT)

C_TITLE       = RGBColor(0x10, 0x35, 0x6B)
C_SECTION     = RGBColor(0x15, 0x4E, 0x9E)
C_SUB         = RGBColor(0x2E, 0x70, 0xC0)
C_BODY        = RGBColor(0x33, 0x33, 0x33)
C_MUTED       = RGBColor(0x66, 0x66, 0x66)
C_ACCENT      = RGBColor(0xB8, 0x3A, 0x3A)
C_PRO         = RGBColor(0x1F, 0x6F, 0x3C)
C_HEAD_FILL   = "E7EEF7"


def _shade(cell, fill_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_borders(table) -> None:
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        tblBorders.append(b)
    tblPr.append(tblBorders)


def add_heading(doc, text: str, level: int = 1) -> None:
    sizes  = {0: 28, 1: 18, 2: 14, 3: 12}
    colors = {0: C_TITLE, 1: C_SECTION, 2: C_SUB, 3: C_SUB}
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8 if level else 0)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes.get(level, 12))
    r.font.color.rgb = colors.get(level, C_BODY)


def add_para(doc, text: str, *, bold: bool = False, italic: bool = False,
             size: int = 10, color: RGBColor = C_BODY,
             space_after: int = 4) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color


def add_bullet(doc, text: str, *, level: int = 0) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.25 + 0.25 * level)
    p.paragraph_format.space_after = Pt(2)
    if not p.runs:
        r = p.add_run(text)
    else:
        r = p.runs[0]
        r.text = text
    r.font.size = Pt(10)
    r.font.color.rgb = C_BODY


def add_mono(doc, text: str, *, size: int = 9) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Inches(0.15)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "D0D7DE")
        pBdr.append(b)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F6F8FA")
    pPr.append(shd)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_table(doc, headers, rows, *, col_widths=None) -> None:
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _set_borders(t)
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        _shade(cell, C_HEAD_FILL)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        r.font.color.rgb = C_TITLE
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = t.rows[ri + 1].cells[ci]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(val))
            r.font.size = Pt(9)
            r.font.color.rgb = C_BODY
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[ci].width = Inches(w)


def add_callout(doc, text: str, *, color: RGBColor = C_SUB) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.15)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "24")
    left.set(qn("w:color"), "%02X%02X%02X" % (color[0], color[1], color[2]))
    pBdr.append(left)
    pPr.append(pBdr)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F6FB")
    pPr.append(shd)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = C_BODY


def add_pros_cons(doc, pros, cons) -> None:
    t = doc.add_table(rows=2, cols=2)
    _set_borders(t)
    for ci, (label, color) in enumerate([("Pros", C_PRO), ("Cons", C_ACCENT)]):
        cell = t.rows[0].cells[ci]
        _shade(cell, C_HEAD_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(label)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = color
    for ci, items in enumerate([pros, cons]):
        cell = t.rows[1].cells[ci]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        first = True
        for it in items:
            if first:
                p = cell.paragraphs[0]
                first = False
            else:
                p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(f"•  {it}")
            r.font.size = Pt(9)
            r.font.color.rgb = C_BODY


def build() -> None:
    doc = Document()
    for s in doc.sections:
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)

    # -------------------------------------------------------------------
    # 0. TITLE
    # -------------------------------------------------------------------
    add_heading(doc, "ARS V2 Retail", level=0)
    add_heading(doc, "Fallback Allocation — Design Review (Pre-Implementation)",
                level=1)
    add_para(
        doc,
        f"Generated {date.today().isoformat()}  •  Status: REVIEW — NOT YET IMPLEMENTED  "
        "•  Owner: santosh kumar",
        italic=True, color=C_MUTED, size=9, space_after=10,
    )
    add_callout(
        doc,
        "Purpose. After the main RL→TBC→TBL waterfall finishes, the fallback "
        "phase boosts grid-level MBQs (never OPT_MBQ), recomputes the REQ chain, "
        "re-runs the waterfall on under-allocated OPTs, and optionally demotes "
        "non-MJ Primary grids one-by-one until the leftover physical pool is "
        "consumed or no more under-allocated OPTs can be helped.",
    )

    # -------------------------------------------------------------------
    # 0a. PLAIN ENGLISH (the whole fallback in 5 minutes, no jargon)
    # -------------------------------------------------------------------
    add_heading(doc, "0a. Plain English — what fallback actually does", level=1)

    add_para(doc, "Picture a warehouse, shelves, and a shipping list.", bold=True)
    add_bullet(doc, "Warehouse = the leftover stock pool (#nre_pool) — physical units we still have after the main pass.")
    add_bullet(doc, "Big shelf (MJ) = the total budget for one MAJ_CAT (e.g. \"men's T-shirts\" overall).")
    add_bullet(doc, "Sub-shelves (grids) = budgets for sub-cuts: by fabric, by vendor, by micro-category, etc.")
    add_bullet(doc, "Per-product display cap (OPT_MBQ) = the most we will ever display of one specific product (one OPT).")
    add_bullet(doc, "Wishlist (OPT_REQ) = how many of each product the store actually wants.")

    add_para(doc, "What happens in the main pass.", bold=True)
    add_bullet(doc, "We walk each product in priority order — RL (must-keep), then TBC (maintain), then TBL (top-up).")
    add_bullet(doc, "For each product, we try to ship up to OPT_REQ — but we stop if the big shelf is full, OR a sub-shelf is full, OR the product hits its own display cap, OR the warehouse runs out.")
    add_bullet(doc, "Result: some products are fully stocked, some get partial, some get zero (because some shelf was full).")

    add_para(doc, "What fallback does (in one paragraph).", bold=True)
    add_callout(doc,
        "\"Some shelves filled up before we could stock everyone. Let me make the sub-shelves 30% bigger "
        "and try again with whatever stock is still in the warehouse. I will NOT change the per-product "
        "display caps — those stay fair. If a product is still stuck because a Primary sub-shelf rule blocks it, "
        "I will downgrade that rule from 'mandatory' to 'advisory' one at a time, lowest priority first, until "
        "either everyone is stocked or there are no more rules to downgrade. The big shelf (MJ) stays the same "
        "unless you flip the include_mj switch.\"",
    )

    add_para(doc, "The two safety rails.", bold=True)
    add_bullet(doc, "OPT_MBQ never grows. A product capped at 60 stays capped at 60, no matter what.")
    add_bullet(doc, "Grid Cap (some shelves opt in) acts as a hard ceiling at e.g. 130% of the ORIGINAL shelf size — even if the boost made the shelf bigger, if you try to over-ship past the cap we trim.")

    # -------------------------------------------------------------------
    # 0b. ONE OPT'S STORY (concrete narrative)
    # -------------------------------------------------------------------
    add_heading(doc, "0b. One OPT's story — OPT4 from start to finish", level=1)
    add_para(doc, "Follow a single product, OPT4 (a TBL \"top-up\" jacket, OPT_REQ=30, OPT_MBQ=30).", italic=True, color=C_MUTED)

    add_table(doc,
        ["Phase", "What happens to OPT4", "Why"],
        [
            ["Main pass",
             "Ships 0 units. ALLOC_FLAG = 'N' (never inserted into alloc table).",
             "PRI_CT_REM = 0 — both Primary sub-shelves (MJ_RNG_SEG and MJ_M_VND_CD) were already empty when OPT4's turn came."],
            ["FB-0 Snapshot",
             "Added to the under-allocated list.",
             "Matches ALLOC_FLAG='N'. We take a photo of every shelf's original size."],
            ["FB-1 Boost",
             "Nothing yet — just shelf sizes grow.",
             "Sub-shelves grow 30%. MJ_M_VND_CD goes 50 → 65. MJ_RNG_SEG goes 80 → 104. MJ shelf untouched."],
            ["FB-2 Recompute",
             "MJ_RNG_SEG now has room. MJ_M_VND_CD also has slight room.",
             "REQ chain re-walks the new shelf sizes."],
            ["FB-3 Insert",
             "OPT4 still NOT inserted. PRI_CT_REM still 0.",
             "MJ_M_VND_CD is still tight enough that the Primary rule fails."],
            ["FB-4 Re-waterfall",
             "OPT4 considered in TBL pass but skipped (still blocked).",
             "RL pass ships OPT2 (+20), TBC pass ships OPT3 (40). TBL pass tries OPT4, can't fit."],
            ["FB-5 level 2",
             "MJ_M_VND_CD demoted from Primary to Secondary (in-memory). OPT4 inserted with FALLBACK_LVL=2.",
             "Lowest-priority Primary downgraded first. PRI_CT_REM now only checks MJ_RNG_SEG → OPT4 eligible."],
            ["FB-5 level 2 waterfall",
             "OPT4 ships 10 units (warehouse runs out).",
             "Pool was 10 entering FB-5. OPT4 wanted 30 but only 10 left. Partial alloc."],
            ["End",
             "OPT4: SHIP_QTY=10, ALLOC_FLAG='Y', FALLBACK_LVL=2, ALLOC_REMARKS=\"FB_DEMOTE(MJ_M_VND_CD)\".",
             "Anyone investigating can immediately see why OPT4 got stock — it took a level-2 demotion."],
        ],
        col_widths=[1.2, 2.3, 3.3],
    )
    add_callout(doc,
        "Read this table top-to-bottom once. The rest of the document is the detailed mechanics of each row.",
        color=C_SUB,
    )

    # -------------------------------------------------------------------
    # 1. HARD RULES
    # -------------------------------------------------------------------
    add_heading(doc, "1. Hard Rules (non-negotiable)", level=1)
    add_bullet(doc, "Fallback NEVER invents stock. It only re-uses the leftover #nre_pool from the main pass.")
    add_bullet(doc, "Fallback NEVER boosts OPT_MBQ or OPT_MBQ_WH. Growth applies only to grid MBQs.")
    add_bullet(doc, "MJ_MBQ is boosted only when fallback_include_mj=True. Default is False.")
    add_bullet(doc, "MJ grid (seq=1) is NEVER demoted. Any non-MJ Primary can be demoted in FB-5.")
    add_bullet(doc, "ARS_GRID_BUILDER on disk is NEVER mutated. Demotion is in-memory only — crash-safe.")
    add_bullet(doc, "Grid Cap (renamed from Sec-grid Cap) is per-grid in ARS_GRID_BUILDER. Global toggle is gone.")
    add_bullet(doc, "Grid Cap basis stays FROZEN at the pre-FB MBQ (cap-frozen). The boost does not move the cap.")

    # -------------------------------------------------------------------
    # 2. WORKED EXAMPLE SETUP
    # -------------------------------------------------------------------
    add_heading(doc, "2. Worked Example — setup we will follow through every phase", level=1)
    add_para(doc, "One MAJ_CAT \"HM24\" at WERKS=1116. Active grids:", space_after=4)
    add_table(doc,
        ["seq", "grid_col", "group", "grid_cap_enabled", "grid_cap_pct"],
        [
            ["1", "MJ (root)",       "Primary",   "off", "—"],
            ["2", "MJ_RNG_SEG",      "Primary",   "off", "—"],
            ["3", "MJ_M_VND_CD",     "Primary",   "on",  "130"],
            ["4", "MJ_FAB",          "Secondary", "on",  "130"],
            ["5", "MJ_MACRO_MVGR",   "Secondary", "off", "—"],
            ["6", "MJ_MICRO_MVGR",   "Secondary", "off", "—"],
            ["7", "MJ_ART_GRD",      "Secondary", "off", "—"],
        ],
        col_widths=[0.5, 1.6, 1.0, 1.3, 1.0],
    )
    add_para(doc, "Original MBQs (pre-fallback snapshot):", space_after=4)
    add_mono(doc,
        "MJ_MBQ=200  MJ_RNG_SEG_MBQ=80  MJ_M_VND_CD_MBQ=50\n"
        "MJ_FAB_MBQ=40  MJ_MACRO_MVGR_MBQ=60\n"
        "MJ_MICRO_MVGR_MBQ=40  MJ_ART_GRD_MBQ=30"
    )
    add_para(doc, "Four OPTs after main pass (enable_fallback=False):", space_after=4)
    add_table(doc,
        ["OPT", "OPT_TYPE", "OPT_MBQ", "OPT_REQ", "SHIP_after_main", "ALLOC_FLAG", "Why blocked"],
        [
            ["OPT1", "RL",  "100", "80", "80", "Y", "Fully allocated — NOT eligible"],
            ["OPT2", "RL",  "60",  "50", "30", "Y", "Hit MJ_FAB headroom (partial)"],
            ["OPT3", "TBC", "40",  "40", "0",  "N", "MJ_RNG_SEG_REM exhausted"],
            ["OPT4", "TBL", "30",  "30", "0",  "N", "PRI_CT_REM=0 (no primary headroom)"],
        ],
        col_widths=[0.6, 0.8, 0.8, 0.8, 1.1, 0.9, 2.2],
    )
    add_para(doc, "Leftover #nre_pool entering fallback: 70 units.", bold=True)
    add_para(doc, "Fallback knobs: enable_fallback=True, growth_pct=130, include_mj=False, demote_primaries=True.")

    # -------------------------------------------------------------------
    # 3. PHASES
    # -------------------------------------------------------------------
    add_heading(doc, "3. Fallback Phases (FB-0 → FB-5)", level=1)

    # ----- FB-0 -----
    add_heading(doc, "FB-0  Snapshot & Identify", level=2)
    add_callout(doc,
        "In plain language: \"Before we touch anything, take a photo of every shelf's "
        "current size, and make a list of products that did not get fully stocked in "
        "the main run.\" The photo lets us audit later; the list tells us who fallback "
        "is trying to help.",
        color=C_SUB,
    )
    add_para(doc, "What it does.", bold=True)
    add_bullet(doc, "Copy every active grid's <prefix>_MBQ into <prefix>_MBQ_PRE_FB columns.")
    add_bullet(doc, "Record MJ_REQ_PRE_FB for the POST-FB MJ_REQ guard.")
    add_bullet(doc, "Tag all main-pass alloc rows with FALLBACK_LVL = 0.")
    add_bullet(doc, "Build the under-allocated OPT set:")
    add_mono(doc, "ALLOC_FLAG = 'N'\n  OR (ALLOC_FLAG = 'Y' AND SHIP_QTY < OPT_REQ)")
    add_para(doc, "Result for the worked example.", bold=True)
    add_bullet(doc, "Under-allocated set = {OPT2, OPT3, OPT4}. OPT1 is full, skipped.")
    add_pros_cons(doc,
        pros=[
            "Audit becomes trivial — every MBQ has a before/after column.",
            "Crash-safe: any later phase can restart from *_PRE_FB.",
        ],
        cons=[
            "One extra column per active grid in ARS_LISTING_WORKING.",
            "Pandas in-memory: negligible cost. SQL Server twin: moderate cost.",
        ],
    )

    # ----- FB-1 -----
    add_heading(doc, "FB-1  Boost Grid MBQs (OPT_MBQ frozen)", level=2)
    add_callout(doc,
        "In plain language: \"Make every sub-shelf 30% bigger. The big shelf (MJ) "
        "stays the same unless include_mj is on. Per-product display caps (OPT_MBQ) "
        "are NEVER touched — that rule is sacred.\"",
        color=C_SUB,
    )
    add_para(doc, "What it does.", bold=True)
    add_bullet(doc, "Multiply every active grid's _MBQ by growth_pct/100.")
    add_bullet(doc, "MJ_MBQ is multiplied ONLY if include_mj=True.")
    add_bullet(doc, "OPT_MBQ and OPT_MBQ_WH are NEVER touched — hard rule.")
    add_para(doc, "With growth_pct=130, include_mj=False:", space_after=4)
    add_table(doc,
        ["Grid", "Before", "After", "Note"],
        [
            ["MJ_MBQ",            "200", "200", "Unchanged (include_mj=False)"],
            ["MJ_RNG_SEG_MBQ",    "80",  "104", "+30%"],
            ["MJ_M_VND_CD_MBQ",   "50",  "65",  "+30%"],
            ["MJ_FAB_MBQ",        "40",  "52",  "+30%"],
            ["MJ_MACRO_MVGR_MBQ", "60",  "78",  "+30%"],
            ["MJ_MICRO_MVGR_MBQ", "40",  "52",  "+30%"],
            ["MJ_ART_GRD_MBQ",    "30",  "39",  "+30%"],
            ["OPT_MBQ (every OPT)", "—", "unchanged", "Hard rule"],
        ],
        col_widths=[1.7, 0.8, 0.8, 2.5],
    )
    add_callout(doc,
        "Consequence for OPT2. OPT_MBQ=60, already shipped 30 → maximum new SHIP "
        "is 60 - 30 = 30 units, regardless of how much grid headroom opens up. "
        "Same logic for any OPT already pinned at OPT_MBQ.",
        color=C_ACCENT,
    )
    add_pros_cons(doc,
        pros=[
            "One predictable multiplier — operators can reason about the outcome.",
            "OPT_MBQ frozen preserves the \"growth at MJ+grid only\" invariant.",
            "include_mj switch handles HM24-style MJ-bound cases on demand.",
        ],
        cons=[
            "Uniform 130% on every grid — may relax grids that did not need relaxing (grid_cap catches over-ship).",
            "OPTs already at OPT_MBQ cannot benefit — may surprise operators (mitigation: audit log).",
        ],
    )

    # ----- FB-2 -----
    add_heading(doc, "FB-2  Recompute REQ Chain", level=2)
    add_callout(doc,
        "In plain language: \"Now that the sub-shelves are bigger, recalculate how "
        "much room each one has. Use exactly the same arithmetic the main run uses — "
        "we are not inventing new rules, just feeding the same math bigger numbers.\"",
        color=C_SUB,
    )
    add_para(doc, "What it does.", bold=True)
    add_bullet(doc, "Re-run the same REQ-chain recompute the main pass uses:")
    add_mono(doc, "<prefix>_REQ  →  _REM  →  H_REM  →  PRI_CT_REM  →  ALLOC_FLAG")
    add_bullet(doc, "Applies to every active grid — not only the boosted ones.")
    add_para(doc, "Effect on the example.", bold=True)
    add_bullet(doc, "MJ_FAB_REM was 0 → now 52 - prior_consumption → OPT2 sees headroom.")
    add_bullet(doc, "MJ_RNG_SEG_REM was 0 → now 104 - prior_consumption → OPT3 becomes eligible.")
    add_bullet(doc, "MJ_REM unchanged (MJ not boosted).")
    add_pros_cons(doc,
        pros=[
            "Reuses production REQ-chain math — zero new failure surface.",
            "Pandas runs in ms; SQL Server in seconds.",
        ],
        cons=[
            "SQL Server twin pays a CTE recompute on each fallback level (FB-2 + each FB-5 iteration).",
        ],
    )

    # ----- FB-3 -----
    add_heading(doc, "FB-3  Insert Newly-Eligible OPTs", level=2)
    add_callout(doc,
        "In plain language: \"Some products were rejected in the main run because every "
        "sub-shelf was full when their turn came. Now that the sub-shelves are bigger, "
        "check if any of them suddenly fit. Add those products to the shipping list and "
        "tag them with 'FB level 1' so we can trace them later.\"",
        color=C_SUB,
    )
    add_para(doc, "What it does.", bold=True)
    add_bullet(doc, "Any OPT in the under-allocated set whose ALLOC_FLAG flipped from N → Y "
                    "after FB-2, AND which is not yet in ARS_ALLOC_WORKING, is INSERTed with FALLBACK_LVL=1.")
    add_bullet(doc, "OPTs already in the alloc table do not get a duplicate row — they just become "
                    "eligible to ship more in FB-4.")
    add_para(doc, "Effect on the example.", bold=True)
    add_bullet(doc, "OPT3 inserted with FALLBACK_LVL=1.")
    add_bullet(doc, "OPT2 already present — no insert.")
    add_bullet(doc, "OPT4 still PRI_CT_REM=0 (MJ_M_VND_CD still tight) — NOT inserted at this level. FB-5 will handle it.")
    add_pros_cons(doc,
        pros=[
            "FALLBACK_LVL gives an audit trail: this row entered alloc via FB-1 boost.",
            "No duplicates — idempotent.",
        ],
        cons=["None significant."],
    )

    # ----- FB-4 -----
    add_heading(doc, "FB-4  Re-run RL → TBC → TBL Waterfall (the heart of fallback)", level=2)
    add_callout(doc,
        "In plain language: \"Take the under-stocked products and run them through the "
        "exact same shipping plan the main run uses — RL first, then TBC, then TBL. "
        "Now that the sub-shelves are bigger, more products fit. Take from the leftover "
        "warehouse and place on shelves until either the product hits its own display "
        "cap, the sub-shelf fills up, or the warehouse runs dry. After everyone has had "
        "a turn, round up to box multiples (pack-round) and enforce hard ceilings on "
        "any shelf that opted in to grid_cap.\"",
        color=C_SUB,
    )

    add_para(doc, "Why RL → TBC → TBL order matters.", bold=True)
    add_bullet(doc, "RL (Replenishment-Listed) = must-stock. Always given first chance at the leftover pool.")
    add_bullet(doc, "TBC (To-Be-Continued) = maintain stock. Given second chance — only what RL didn't take.")
    add_bullet(doc, "TBL (To-Be-Listed) = top-up / nice-to-have. Last chance — only what RL and TBC didn't take.")
    add_bullet(doc, "Within each type, OPTs are ordered by priority (rank). Higher-priority OPTs go first.")

    add_para(doc, "Step-by-step inside FB-4 for our worked example.", bold=True)
    add_para(doc,
        "Entering FB-4: pool=70 units. Under-allocated set = {OPT2 (RL, SHIP=30/50), "
        "OPT3 (TBC, SHIP=0/40, just inserted), OPT4 (TBL, SHIP=0/30)}.",
        italic=True, color=C_MUTED,
    )

    add_para(doc, "Step 1 — RL pass: process OPT2.", bold=True)
    add_bullet(doc, "OPT2 needs 50 - 30 = 20 more units to be full.")
    add_bullet(doc, "Check OPT_MBQ: SHIP would become 50, cap is 60 → OK.")
    add_bullet(doc, "Check MJ_FAB_REM: boosted to 52 from 40. Headroom available. → OK.")
    add_bullet(doc, "Check MJ_MBQ: still 200 (not boosted). Plenty of room. → OK.")
    add_bullet(doc, "Pool has 70 → 20 available. → Ship 20.")
    add_bullet(doc, "Update: OPT2.SHIP=50 (full), pool=70-20=50.")

    add_para(doc, "Step 2 — TBC pass: process OPT3.", bold=True)
    add_bullet(doc, "OPT3 was just inserted in FB-3 with FALLBACK_LVL=1. Needs 40.")
    add_bullet(doc, "Check OPT_MBQ: SHIP would become 40, cap is 40 → exactly at cap, OK.")
    add_bullet(doc, "Check MJ_RNG_SEG_REM: boosted to 104 from 80. Headroom. → OK.")
    add_bullet(doc, "Pool=50 → 40 available. → Ship 40.")
    add_bullet(doc, "Update: OPT3.SHIP=40 (full), pool=50-40=10.")

    add_para(doc, "Step 3 — TBL pass: process OPT4.", bold=True)
    add_bullet(doc, "OPT4 needs 30. Pool has only 10 left.")
    add_bullet(doc, "Check PRI_CT_REM: needs BOTH Primary sub-shelves (MJ_RNG_SEG AND MJ_M_VND_CD) to have room.")
    add_bullet(doc, "MJ_RNG_SEG_REM: has room after OPT3 used 40 of 104.")
    add_bullet(doc, "MJ_M_VND_CD_REM: boosted to 65 from 50, but still TIGHT — prior main-pass consumption already filled most of it.")
    add_bullet(doc, "PRI_CT_REM check FAILS. OPT4 is SKIPPED in FB-4.")
    add_bullet(doc, "Pool stays at 10. OPT4.SHIP=0. Still under-allocated.")

    add_para(doc, "Step 4 — pack-round.", bold=True)
    add_bullet(doc, "Products usually ship in packs (e.g. carton of 6). pack-round forces SHIP to a multiple.")
    add_bullet(doc, "If pack-round trims any unit, that unit returns to the pool (which is now 10 or slightly higher).")
    add_bullet(doc, "For this example, assume no trim — all SHIPs already multiples.")

    add_para(doc, "Step 5 — grid_cap enforcement (only on grids with grid_cap_enabled=1).", bold=True)
    add_bullet(doc, "MJ_FAB has cap on at 130%. Cap basis is FROZEN at original MJ_FAB_MBQ=40, so ceiling = 1.30 × 40 = 52.")
    add_bullet(doc, "Compute Sigma SHIP per MJ_FAB value across all shipped OPTs in this MJ_FAB bucket.")
    add_bullet(doc, "If Sigma SHIP ≤ 52 → no action. If > 52 → trim lowest-priority rows (FALLBACK_LVL DESC, OPT_TYPE TBL > TBC > RL, rank ASC) until Sigma ≤ 52. Trimmed units return to pool.")
    add_bullet(doc, "Same check for MJ_M_VND_CD (also opted in at 130%, ceiling = 1.30 × 50 = 65).")
    add_bullet(doc, "Other grids (MJ_MACRO_MVGR, MJ_MICRO_MVGR, MJ_ART_GRD) have grid_cap_enabled=0 → no check.")

    add_para(doc, "FB-4 result summary.", bold=True)
    add_table(doc,
        ["OPT", "OPT_TYPE", "SHIP after main", "SHIP after FB-4", "FALLBACK_LVL", "Status"],
        [
            ["OPT1", "RL",  "80",  "80",         "0", "full (not touched by fallback)"],
            ["OPT2", "RL",  "30",  "50 (+20)",   "0", "full"],
            ["OPT3", "TBC", "0",   "40 (+40)",   "1", "full (inserted in FB-3)"],
            ["OPT4", "TBL", "0",   "0",          "—", "STILL under-allocated → FB-5 next"],
        ],
        col_widths=[0.6, 0.8, 1.2, 1.3, 1.1, 2.1],
    )
    add_para(doc, "Pool entering FB-5: 10 units. Loop continues.", bold=True)

    add_para(doc, "Two subtle but important rules inside FB-4.", bold=True)
    add_bullet(doc,
        "Filter is on under-allocated OPTs only. Already-full OPTs (e.g. OPT1) are read-only — "
        "their SHIP_QTY contributes to shelf consumption but they never get a second pass.")
    add_bullet(doc,
        "Budgets are SHARED. When OPT2 ships +20, that 20 reduces MJ_FAB_REM, MJ_MACRO_MVGR_REM, "
        "etc. for every later OPT in the same pass. The REQ chain re-reads remaining budgets on "
        "every step — exactly like main pass.")
    add_bullet(doc,
        "R09 cap is unchanged. The OPT_MBQ check inside the waterfall uses the original (non-boosted) "
        "OPT_MBQ. An OPT pinned at its OPT_MBQ in main pass cannot ship more in fallback, full stop.")

    add_pros_cons(doc,
        pros=[
            "Reuses production waterfall code (_run_majcat_waterfall in pandas / _stage_c_waterfall in SQL) — no parallel implementation to maintain.",
            "Filter to under-allocated keeps runtime bounded.",
            "Same priority order (RL > TBC > TBL) makes the outcome easy to explain to operators.",
            "Pack-round and grid-cap give two layers of safety against over-ship.",
        ],
        cons=[
            "Filter requires care: gates read shared budgets (MJ_REM, sub-grid_REM). FB-2 already accounts for this, but must be regression-tested.",
            "Grid-cap trims need a deterministic priority order so we always trim the right rows. Document the trim order in code comments.",
            "If pool was already empty exiting main pass, FB-4 is a no-op. Audit log should make this visible (\"fallback ran but no pool to allocate\").",
        ],
    )

    # ----- FB-5 -----
    add_heading(doc, "FB-5  Primary → Secondary Demotion Loop", level=2)
    add_callout(doc,
        "In plain language: \"FB-4 boosted the sub-shelves but a few products are still "
        "blocked because a PRIMARY rule (a sub-shelf that is mandatory, not just a budget) "
        "still says no. Take the LOWEST-PRIORITY primary rule and downgrade it from "
        "'mandatory' to 'advisory' — just for this run, not on disk. Recalculate, run the "
        "waterfall again. If still stuck, downgrade the next-lowest primary rule. Repeat "
        "until everyone is helped OR only the big shelf (MJ) is left as primary — we "
        "NEVER downgrade MJ.\"",
        color=C_SUB,
    )
    add_para(doc, "Algorithm.", bold=True)
    add_mono(doc,
        "level = 2\n"
        "while under_allocated AND non-MJ Primaries remain:\n"
        "    g = pop highest-seq Primary  # skip seq=1 (MJ)\n"
        "    treat g as Secondary IN-MEMORY only (ARS_GRID_BUILDER untouched)\n"
        "    recompute REQ chain (FB-2 logic)\n"
        "    insert newly-eligible OPTs with FALLBACK_LVL = level\n"
        "    run RL -> TBC -> TBL -> pack_round -> grid_cap (FB-4 logic)\n"
        "    level += 1\n"
        "    if no OPTs surfaced this iteration: break"
    )
    add_para(doc, "Effect on the example (pool=10 entering FB-5):", space_after=4)
    add_para(doc,
        "Level 2 — demote seq=3 MJ_M_VND_CD (highest non-MJ Primary).",
        bold=True,
    )
    add_bullet(doc, "MJ_M_VND_CD moves from Primary set → Secondary set for this run only.")
    add_bullet(doc, "PRI_CT_REM for OPT4 now depends only on MJ_RNG_SEG (the remaining Primary).")
    add_bullet(doc, "MJ_RNG_SEG_REM has room (104 - prior consumption) → OPT4 becomes eligible.")
    add_bullet(doc, "OPT4 inserted with FALLBACK_LVL=2.")
    add_bullet(doc, "TBL pass ships OPT4 = 10 units (pool exhausts at 10, OPT_REQ was 30 — partial).")
    add_bullet(doc, "Pool=0, under-allocated set effectively done. Loop exits.")
    add_callout(doc,
        "Stop conditions. (a) No under-allocated OPTs remain. "
        "(b) Only MJ left as Primary — cannot demote. "
        "(c) An iteration surfaced zero new OPTs — break early.",
    )
    add_pros_cons(doc,
        pros=[
            "Targets the real binding constraint per OPT.",
            "Largest-seq-first ensures lowest-priority Primaries demote first.",
            "In-memory only — a mid-fallback crash leaves ARS_GRID_BUILDER clean.",
        ],
        cons=[
            "Iteration count = (non-MJ Primary count). Bounded but adds runtime.",
            "Operators need clear telemetry: log ALLOC_REMARKS=\"FB_DEMOTE(MJ_M_VND_CD)\" so the cause is visible.",
        ],
    )

    # -------------------------------------------------------------------
    # 4. GRID CAP (renamed)
    # -------------------------------------------------------------------
    add_heading(doc, "4. Grid Cap (renamed from Sec-grid Cap)", level=1)
    add_callout(doc,
        "In plain language: \"Some shelves are critical — fabric, vendor — and we never "
        "want to over-stock them, even when fallback is boosting budgets. For those "
        "shelves, set a hard ceiling (e.g. 130% of the ORIGINAL shelf size). If shipping "
        "would push the shelf past that ceiling, trim the lowest-priority products until "
        "we are back under. The operator picks per-shelf whether the cap is on — it is "
        "no longer a single global switch.\"",
        color=C_SUB,
    )
    add_para(doc, "Old behavior.", bold=True)
    add_bullet(doc, "Global toggle apply_sec_cap_in_normal on the request payload.")
    add_bullet(doc, "_apply_sec_grid_cap_pre_gate walked every Secondary grid at 130%.")
    add_para(doc, "New behavior.", bold=True)
    add_bullet(doc, "Per-grid flag in ARS_GRID_BUILDER:")
    add_table(doc,
        ["Column", "Type", "Default", "Meaning"],
        [
            ["grid_cap_enabled", "BIT",   "0",   "Cap on/off for this grid."],
            ["grid_cap_pct",     "FLOAT", "130", "Cap multiplier (e.g. 130 = 130%)."],
        ],
        col_widths=[1.8, 0.9, 0.8, 3.0],
    )
    add_bullet(doc, "Cap function (renamed _apply_grid_cap) walks ONLY grids with grid_cap_enabled=1.")
    add_bullet(doc, "Global apply_sec_cap_in_normal request flag is REMOVED.")
    add_bullet(doc, "During fallback the cap basis is FROZEN at <prefix>_MBQ_PRE_FB. The boost makes room INSIDE the cap; it does not move the cap.")
    add_para(doc, "Example with grid_cap_enabled=1 on MJ_FAB at 130%:", bold=True)
    add_mono(doc,
        "Pre-FB MBQ:           MJ_FAB_MBQ_PRE_FB = 40\n"
        "Cap basis (frozen):   1.30 * 40 = 52\n"
        "Post-FB SHIP_FAB:     <= 52   (otherwise trim lowest-priority rows)"
    )
    add_pros_cons(doc,
        pros=[
            "Operator picks exactly which grids deserve a cap.",
            "One config table, one function — main pass and fallback share the mechanism.",
            "Cap-frozen during fallback prevents the boost from defeating the circuit-breaker.",
        ],
        cons=[
            "Requires GridBuilderPage.jsx change (checkbox + % per grid).",
            "Migration choice: default all grids to off (clean slate) vs. auto-on for Secondary (behavior-preserving).",
        ],
    )

    # -------------------------------------------------------------------
    # 5. KNOBS
    # -------------------------------------------------------------------
    add_heading(doc, "5. Request-level Knobs", level=1)
    add_table(doc,
        ["Field", "Default", "Effect"],
        [
            ["enable_fallback",              "False", "Master gate."],
            ["fallback_growth_pct",          "130.0", "Single % for all grid MBQs in scope."],
            ["fallback_include_mj",          "False", "If True, MJ_MBQ is also boosted."],
            ["fallback_demote_primaries",    "True",  "If False, skip FB-5 entirely (boost-only fallback)."],
            ["fallback_max_demotion_levels", "None",  "Cap on FB-5 iterations. None = demote all non-MJ Primaries."],
        ],
        col_widths=[2.4, 0.8, 3.5],
    )
    add_callout(doc,
        "Removed: apply_sec_cap_in_normal, fallback_boost_scope, fallback_rerun_scope, "
        "fallback_boost_mode, static_growth_pct, str_tiers. Caps are now per-grid; the "
        "single growth_pct + include_mj pair replaces the old scope knobs.",
    )

    # -------------------------------------------------------------------
    # 6. SCHEMA CHANGES
    # -------------------------------------------------------------------
    add_heading(doc, "6. Schema Changes", level=1)
    add_para(doc, "ARS_LISTING_WORKING (or pandas in-memory equivalent):", bold=True)
    add_bullet(doc, "<prefix>_MBQ_PRE_FB FLOAT — one per active grid (incl. MJ).")
    add_bullet(doc, "MJ_REQ_PRE_FB FLOAT — for POST-FB MJ_REQ guard.")
    add_para(doc, "ARS_ALLOC_WORKING:", bold=True)
    add_bullet(doc, "FALLBACK_LVL INT NULL DEFAULT 0 — 0=main, 1=FB-4, 2..N=FB-5 demotion level.")
    add_bullet(doc, "FB_SHIP_QTY FLOAT — SHIP_QTY minus main_final_SHIP_QTY (FB delta accounting).")
    add_para(doc, "ARS_GRID_BUILDER:", bold=True)
    add_bullet(doc, "grid_cap_enabled BIT NOT NULL DEFAULT 0.")
    add_bullet(doc, "grid_cap_pct FLOAT NULL DEFAULT 130.")

    # -------------------------------------------------------------------
    # 7. FILE-LEVEL CHANGES
    # -------------------------------------------------------------------
    add_heading(doc, "7. File-level Changes (high level)", level=1)
    add_table(doc,
        ["File", "Change"],
        [
            ["backend/app/api/v1/endpoints/listing.py",
             "Add 5 request fields to GenerateRequest. Remove apply_sec_cap_in_normal."],
            ["backend/app/services/rule_engine_pandas.py",
             "Add _run_fallback_pandas orchestrating FB-0..FB-5. Wire in after the main sec-cap call (~line 944). "
             "Reuses _run_majcat_waterfall for FB-4 and each FB-5 iteration."],
            ["backend/app/services/rule_engine_new.py",
             "Add _run_fallback_new (SQL Server twin). Rename _apply_sec_grid_cap_pre_gate → _apply_grid_cap. "
             "Read grid_cap_enabled from ARS_GRID_BUILDER. Drop apply_sec_cap_in_normal plumbing."],
            ["backend/app/services/listing_allocator.py",
             "Delete the legacy _run_fallback block (already marked legacy; design does not match new flow)."],
            ["frontend/src/pages/ListingPage.jsx",
             "Restore \"Fallback Allocation\" ParamGroup with 5 controls; remove \"Sec-grid Cap\" toggle."],
            ["frontend/src/pages/GridBuilderPage.jsx",
             "Add grid_cap_enabled checkbox and grid_cap_pct input per grid row."],
            ["frontend/src/services/api.js",
             "Update listing-run payload mapper to include the 5 new fields."],
            ["backend/app/docs/processes/fallback_archived.md",
             "Replace with fallback.md (active spec) once shipped, linking the rebuild commit."],
        ],
        col_widths=[2.8, 4.0],
    )

    # -------------------------------------------------------------------
    # 8. POST-FB GUARDS
    # -------------------------------------------------------------------
    add_heading(doc, "8. Post-FB Guards (do not skip)", level=1)
    add_callout(doc,
        "HM24 historical bug. When include_mj=True boosted MJ_MBQ, Sigma SHIP per "
        "(WERKS, MAJ_CAT) was allowed to exceed the original MJ_REQ. Without a guard, "
        "fallback over-shipped MAJ_CATs whose actual demand was small.",
        color=C_ACCENT,
    )
    add_para(doc, "Recommended guard (re-instate from archived spec).", bold=True)
    add_bullet(doc, "After FB-4 and after every FB-5 iteration, run _stage_c_apply_mj_req_cap with phase=\"POST_FB\".")
    add_bullet(doc, "Allow Sigma SHIP > original MJ_REQ ONLY when include_mj=True AND MJ_REQ_PRE_FB >= 0.5 * ACS_D.")
    add_bullet(doc, "Otherwise trim FB-added units (FALLBACK_LVL > 0) from lowest-priority rows first.")
    add_bullet(doc, "Note: existing per-OPT_TYPE rl/tbc/tbl_mj_req_cap_pct knobs (default 100) already enforce "
                    "a hard ceiling at MJ_REQ for main pass. Decide if the POST-FB cap is additive (trim FB rows first) "
                    "or redundant (rely on existing knobs).")

    # -------------------------------------------------------------------
    # 9. AUDIT OUTPUT
    # -------------------------------------------------------------------
    add_heading(doc, "9. Audit Output (returned via API response)", level=1)
    add_mono(doc,
        'result["fallback_audit"] = {\n'
        '  "enabled":              True,\n'
        '  "growth_pct":           130.0,\n'
        '  "include_mj":           False,\n'
        '  "demote_primaries":     True,\n'
        '  "levels_run":           [1, 2, 3, ...],\n'
        '  "newly_eligible_total": <int>,\n'
        '  "fallback_ship_qty":    <float>,   # Sigma SHIP_QTY WHERE FALLBACK_LVL > 0\n'
        '  "pack_round_units":     <float>,\n'
        '  "grid_cap_units_trimmed": <float>,\n'
        '  "duration_sec":         <float>,\n'
        '}'
    )
    add_bullet(doc, "ALLOC_REMARKS on each FB-shipped row: FB_BOOST(level) or FB_DEMOTE(grid_col).")
    add_bullet(doc, "Grid-cap trim events log GRID_CAP_HIT(grid_col, trimmed=N).")

    # -------------------------------------------------------------------
    # 10. OPEN QUESTIONS
    # -------------------------------------------------------------------
    add_heading(doc, "10. Open Questions (answer before implementation)", level=1)
    add_table(doc,
        ["#", "Question", "Recommendation"],
        [
            ["1", "Under-allocated definition.",
                  "ALLOC_FLAG='N' OR (ALLOC_FLAG='Y' AND SHIP_QTY<OPT_REQ). Confirm."],
            ["2", "POST-FB MJ_REQ guard rule.",
                  "Re-instate MJ_REQ_PRE_FB >= 0.5 * ACS_D from archived spec."],
            ["3", "fallback_max_demotion_levels default.",
                  "None (demote until MJ-only). Add UI control if needed later."],
            ["4", "grid_cap_enabled migration default.",
                  "Default off. Run a one-time SQL: SET grid_cap_enabled=1 WHERE group='Secondary' "
                  "if behavior-preserving rollout is required."],
            ["5", "Partial OPT4-style outcomes.",
                  "Acceptable. Fallback is best-effort. ALLOC_REMARKS records why."],
        ],
        col_widths=[0.4, 3.0, 3.4],
    )

    # -------------------------------------------------------------------
    # 11. APPROVAL
    # -------------------------------------------------------------------
    add_heading(doc, "11. Approval Block", level=1)
    add_table(doc,
        ["Role", "Name", "Decision", "Date"],
        [
            ["Owner",      "santosh kumar", "",    ""],
            ["Reviewer 1", "",              "",    ""],
            ["Reviewer 2", "",              "",    ""],
        ],
        col_widths=[1.5, 2.0, 1.8, 1.2],
    )

    add_para(doc, "", space_after=12)
    add_para(doc,
        "End of document. Source: backend/scripts/build_fallback_doc.py. "
        "Re-run with backend venv python after edits.",
        italic=True, color=C_MUTED, size=9,
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
