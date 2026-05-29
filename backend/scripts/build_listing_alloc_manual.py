"""
Generates `ARS_Listing_Allocation_Manual.docx` — a USER-facing manual covering:
  • The full listing & allocation pipeline (Parts 1-8) with worked examples
  • OPT_TYPE classification rules with examples
  • Every cap/gate rule and toggle pattern matrix (PRI strict, sec-cap, fallback)
  • Recent fixes applied (May 2026 — PRI strict → MJ-cap, shared cross-type
    budget, live MJ_REQ_REM per-band recompute)
  • BRD: existing state vs proposed improvements
  • Performance & result improvement suggestions
  • Open items / review checklist

Run:
    python backend/scripts/build_listing_alloc_manual.py
Output:
    d:/ARS_PROD/ars_prod/ARS_Listing_Allocation_Manual.docx
"""
from __future__ import annotations
import os
from datetime import date

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


OUT = r"d:/ARS_PROD/ars_prod/ARS_Listing_Allocation_Manual.docx"

# --- palette ---
C_TITLE   = RGBColor(0x10, 0x35, 0x6B)
C_H1      = RGBColor(0x15, 0x4E, 0x9E)
C_H2      = RGBColor(0x2E, 0x70, 0xC0)
C_H3      = RGBColor(0x4A, 0x7B, 0xBA)
C_BODY    = RGBColor(0x22, 0x22, 0x22)
C_OK      = RGBColor(0x1B, 0x6E, 0x2D)
C_WARN    = RGBColor(0xB8, 0x3A, 0x3A)
C_MUTED   = RGBColor(0x66, 0x66, 0x66)
F_HEAD    = "DCE6F2"
F_NOTE    = "FFF7DC"
F_FIX     = "E8F5E9"
F_WARN    = "FDECEA"


def _shade(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    bdr = OxmlElement("w:tblBorders")
    for e in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{e}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "BFBFBF")
        bdr.append(b)
    tblPr.append(bdr)


def H(doc, text, level=1):
    sizes = {0: 26, 1: 18, 2: 14, 3: 12}
    cols = {0: C_TITLE, 1: C_H1, 2: C_H2, 3: C_H3}
    p = doc.add_paragraph()
    if level == 0:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10 if level else 0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(sizes.get(level, 12))
    r.font.color.rgb = cols.get(level, C_BODY)


def P(doc, text, bold=False, italic=False, size=10, color=None, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Inches(indent)
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color or C_BODY


def Bullet(doc, text, level=0):
    style = "List Bullet" if level == 0 else "List Bullet 2"
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = C_BODY


def Code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


def Box(doc, label, body, fill=F_NOTE, label_color=None):
    """Note/Tip/Warning callout box."""
    t = doc.add_table(rows=1, cols=1)
    _borders(t)
    cell = t.rows[0].cells[0]
    _shade(cell, fill)
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(label + ": ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = label_color or C_H1
    r2 = p1.add_run(body)
    r2.font.size = Pt(10)
    r2.font.color.rgb = C_BODY
    doc.add_paragraph()


def Table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    _borders(t)
    for j, h in enumerate(headers):
        c = t.rows[0].cells[j]
        _shade(c, F_HEAD)
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = C_H1
    for i, row in enumerate(rows, 1):
        for j, v in enumerate(row):
            c = t.rows[i].cells[j]
            p = c.paragraphs[0]
            r = p.add_run(str(v))
            r.font.size = Pt(9)
            r.font.color.rgb = C_BODY
    if col_widths:
        for row in t.rows:
            for j, w in enumerate(col_widths):
                row.cells[j].width = Inches(w)
    doc.add_paragraph()


# =============================================================================
#  Document
# =============================================================================
def build():
    doc = Document()

    # set narrow margins
    for s in doc.sections:
        s.left_margin = Inches(0.75)
        s.right_margin = Inches(0.75)
        s.top_margin = Inches(0.6)
        s.bottom_margin = Inches(0.6)

    # ----------------------------------------------------------------- title
    H(doc, "ARS Listing & Allocation — Complete Manual", 0)
    P(doc, "V2 Retail Auto Replenishment System",
      italic=True, size=11, color=C_MUTED)
    P(doc, f"Document date: {date.today().isoformat()} | "
      f"Pipeline: backend/app/api/v1/endpoints/listing.py + "
      f"rule_engine_pandas.py + rule_engine_new.py",
      italic=True, size=9, color=C_MUTED)
    doc.add_paragraph()

    P(doc, "This manual covers WHAT the pipeline does and WHY, with concrete "
      "examples for every rule and toggle. It is meant to be read end-to-end "
      "by ops + product so anyone can: (a) understand each step, (b) verify "
      "the engine's behaviour against business intent, (c) flag what still "
      "needs improvement.", size=10)
    doc.add_page_break()

    # ====================================================================
    # PART I — PROCESS MANUAL
    # ====================================================================
    H(doc, "Part I — Process Manual", 1)

    # ---------------------------------------------------------- 1. Overview
    H(doc, "1. Purpose & End-to-End Flow", 2)
    P(doc, "Each run takes the latest store stock, the MSA-driven option "
      "universe, and grid attributes; produces a per-store, per-size shipment "
      "list ready for warehouse dispatch.")
    P(doc, "There are two macro phases:", bold=True)
    Bullet(doc, "Listing — Parts 1 to 7 (SQL): build ARS_LISTING, classify "
                "every option as RL/TBC/TBL/MIX, enrich with caps & requirements.")
    Bullet(doc, "Allocation — Part 8 (pandas): Stage A filter → Stage B explode "
                "to size → Stage C waterfall (RL → TBC → TBL) → Stage D sec-grid "
                "cap → Stage E fallback (optional) → Stage F finalise.")

    Box(doc, "Output tables",
        "ARS_LISTING (option grain), ARS_LISTING_WORKING (eligible only), "
        "ARS_ALLOC_WORKING (size grain with SHIP/HOLD), ARS_STORE_RANKING. "
        "Each run also snapshots these to ARS_*_PARKED for audit.",
        fill=F_NOTE)

    # ----------------------------------------------------------- 2. Listing
    H(doc, "2. Listing Stage (Parts 1 → 7)", 2)
    Table(doc,
          ["Part", "What it does", "Source(s)", "Output column(s) added"],
          [
              ["1",   "Insert grid rows for active stores into ARS_LISTING (IS_NEW=0).",
               "ARS_LISTING_GRID, ARS_ST_MASTER", "All SLOC cols, STK_TTL, STR"],
              ["2",   "Insert MSA-missing options as IS_NEW=1 (zero stock).",
               "MSA tables", "IS_NEW=1 rows"],
              ["2.5", "Build helper indexes (WERKS, MAJ_CAT), (GEN_ART), and "
                      "(WERKS, MAJ_CAT, GEN_ART, CLR) — May 2026 addition.",
               "—", "—"],
              ["3.5", "Pull ACS_D, ALC_D, AGE, AUTO_GEN_ART_SALE from master "
                      "tables.",
               "ARS_CALC_ST_MAJ_CAT, ARS_CALC_ST_ART, MASTER_GEN_ART_SALE/AGE",
               "ACS_D, ALC_D, AGE"],
              ["3.54","Roll up open warehouse holds from prior TBL runs.",
               "ARS_NL_TBL_HOLD_TRACKING", "RL_HOLD_QTY"],
              ["3.55","Pre-populate MSA_FNL_Q, VAR_COUNT, VAR_FNL_COUNT.",
               "ARS_MSA_GEN_ART, ARS_MSA_VAR_ART", "MSA_FNL_Q, VAR_COUNT"],
              ["3.6", "Classify every option as RL / TBC / TBL / MIX.",
               "Logic CASE", "OPT_TYPE"],
              ["3.7", "Collapse MIX rows: max 1 MIX line per (store, MAJ_CAT).",
               "self-aggregate", "(rows compacted)"],
              ["4",   "Add grid columns (MJ_MBQ, RNG_SEG_MBQ, CLR_MBQ etc.) by "
                      "joining each active grid table.",
               "Grid tables", "<prefix>_STK_TTL/MBQ/CONT/OPT_CNT/DISP_Q"],
              ["4b",  "PER_OPT_SALE from the use_for_opt_sale grid.",
               "MJ_RNG_SEG (typical)", "PER_OPT_SALE"],
              ["4c",  "Compute OPT_MBQ = ACS_D + rate × ALC_D, OPT_REQ, "
                      "OPT_MBQ_WH (+hold_days for new), OPT_REQ_WH, MAX_DAILY_SALE.",
               "—", "OPT_MBQ, OPT_REQ, OPT_MBQ_WH, OPT_REQ_WH"],
              ["4d",  "ART_EXCESS = MAX(0, STK_TTL − 2 × OPT_MBQ); MIX rows skip.",
               "—", "ART_EXCESS, EXCESS_STK"],
              ["4e",  "Per-grid REQ with excess deduction.",
               "—", "<grid>_REQ for each grid"],
              ["5",   "Final indexes on ARS_LISTING.", "—", "—"],
              ["6",   "Rank stores by weighted REQ × Fill%.",
               "ARS_LISTING aggregates", "ARS_STORE_RANKING, ST_RANK"],
              ["7",   "Build ARS_LISTING_WORKING (eligible rows only: "
                      "MSA_FNL_Q>0 OR HOLD_QTY>0, OPT_REQ_WH≥1). Compute "
                      "PRI_CT% / SEC_CT% from grid headers. Set ALLOC_FLAG.",
               "—", "GH_*, H_*, PRI_CT%, SEC_CT%, ALLOC_FLAG"],
          ],
          col_widths=[0.5, 3.0, 2.0, 2.0])

    # ------------------------------------------------ 3. OPT_TYPE rules
    H(doc, "3. OPT_TYPE Classification (Part 3.6)", 2)
    P(doc, "Each option gets exactly one OPT_TYPE based on stock vs ACS_D, "
      "MSA supply, NL hold, and color-fill ratio. Order matters — first match "
      "wins.")
    Table(doc,
          ["Type", "Rule (simplified)", "Meaning"],
          [
              ["MIX (a)", "STK < 60% × ACS_D AND MSA_FNL_Q = 0 AND RL_HOLD_QTY = 0",
               "Low stock, nothing to send — collapsed into one MIX line per store×MAJ_CAT"],
              ["MIX (b)", "VAR_FNL_COUNT / VAR_COUNT < 60%  (or < min_size_count)",
               "Poor color fill — even if other rules pass, group as MIX"],
              ["RL",      "(STK ≥ 60% × ACS_D OR RL_HOLD_QTY > 0) AND MSA_FNL_Q > 0",
               "Adequate stock + fresh MSA supply — top-up dispatch"],
              ["TBC",     "0 < STK < 60% × ACS_D AND (MSA_FNL_Q > 0 OR RL_HOLD_QTY > 0)",
               "Low stock but supply exists — break-pack candidate"],
              ["TBL",     "STK ≤ 0 AND (MSA_FNL_Q > 0 OR RL_HOLD_QTY > 0)",
               "Zero stock — new listing"],
          ], col_widths=[0.7, 3.5, 3.0])
    Box(doc, "Live example",
        "Store HJ08, ACS_D=18 → 60% threshold = 10.8 units. "
        "Article 1116112692 colour M_PST, STK=4, MSA_FNL_Q=24. "
        "STK<10.8 AND MSA>0 ⇒ TBC. If STK were 0 ⇒ TBL. If STK were 15 ⇒ RL.",
        fill=F_NOTE)

    # ---------------------------------------- 4. Allocation stage walkthrough
    H(doc, "4. Allocation Stage (Part 8)", 2)

    H(doc, "4.1 Stage A — rule filtering", 3)
    P(doc, "Each option is tested against rules R01–R09. First failure marks "
      "ALLOC_STATUS='SKIPPED' with a SKIP_REASON.")
    Table(doc,
          ["Rule", "Check", "Notes"],
          [
              ["R01_LISTING",        "LISTING flag = 1",                       "Published flag must be true"],
              ["R02_NOT_MIX",        "OPT_TYPE ≠ 'MIX'",                       "MIX rows never allocate"],
              ["R04_MSA_POS",        "MSA_FNL_Q > 0 OR RL_HOLD_QTY > 0",       "Need supply OR prior hold"],
              ["R05_REQ_POS",        "OPT_REQ_WH ≥ 1",                          "Need at least 1 unit of demand"],
              ["R06_PRI_100",        "PRI_CT% = 100",                          "TBL always enforces; RL/TBC only when pri_ct_check_* flag is True"],
              ["R07_VAR_RATIO_TBL",  "VAR_FNL/VAR_COUNT ≥ 60% OR VAR_FNL ≥ 3", "TBL only — block thin-size variants"],
              ["R09_TBL_TRIVIAL",    "cap_pct × MJ_MBQ − MJ_STK_TTL ≥ 0.5 × ACS_D",
                                     "Skip when remaining headroom is sub-significant"],
          ], col_widths=[1.6, 3.0, 3.0])

    H(doc, "4.2 Stage B — explode to size grain", 3)
    P(doc, "Each surviving OPT is exploded into one row per (VAR_ART, SZ) using "
      "the VAR contribution table. Columns added: CONT (size %), SZ_MBQ, "
      "SZ_REQ, SZ_REQ_WH, FNL_Q_REM (live pool), SHIP_QTY=0, HOLD_QTY=0, "
      "ALLOC_STATUS='PENDING'.")

    H(doc, "4.3 Stage C — pandas waterfall (RL → TBC → TBL)", 3)
    P(doc, "All stores compete simultaneously inside each band. Each "
      "OPT_TYPE runs to completion (across all rounds) before the next type.")
    Bullet(doc, "Bands within an OPT_TYPE = rounds 1, 2, … up to MAX(I_ROD).")
    Bullet(doc, "Pool drain order per band: POOL_KEYS → OPT_PRIORITY_RANK ASC "
                "→ ST_RANK ASC → WERKS. Best store gets pool first.")
    Bullet(doc, "After every band, _revalidate_after_band decrements "
                "MJ_REQ_REM, MSA_FNL_Q_REM, and each primary grid's *_REQ_REM. "
                "Pre-band check skips OPTs whose PRI_CT_REM dropped below 100.")

    Box(doc, "Why both RL and TBC run before TBL",
        "RL and TBC top up existing stock — preserving service for proven sellers. "
        "TBL ships new options last so the warehouse pool doesn't get drained "
        "by speculative new listings while existing options are still short.",
        fill=F_NOTE)

    H(doc, "4.4 Stage D — secondary-grid cap pre-gate", 3)
    P(doc, "After the waterfall, for every shipped OPT we check: would the "
      "post-alloc total for this OPT push any Secondary grid (MJ_MICRO_MVGR, "
      "MJ_FAB, MJ_CLR, MJ_M_VND_CD) above SEC_CAP_DEFAULT_PCT × grid_MBQ "
      "(default 130%)?")
    Bullet(doc, "BLOCK — entire OPT's SHIP set to 0; ALLOC_REASON = "
                "'BLOCKED_SEC_CAP_PRE_<grid>'.")
    Bullet(doc, "OVERRIDE — if the OPT's OPT_REQ ≥ 100% × OPT_MBQ "
                "(SEC_CAP_PRE_OVERRIDE_OPT_REQ_PCT), the breach is allowed "
                "(remark appended).")
    P(doc, "Toggle: apply_sec_cap_in_normal (default True). Always ON inside "
      "fallback regardless.")

    Box(doc, "Live example — session 20260514_152345_428",
        "Sec-cap pre-gate logged: blocked=2 OPTs (47 units) overridden=26 "
        "OPTs (1131 units). Both blocked OPTs were MJ_CLR breaches at HD28 "
        "(L_GRY, 24u) and HJ15 (SEA_GRN, 23u). Comparing to session "
        "20260514_152624_574 (sec-cap OFF), those 47 units shipped — "
        "explaining the +47 ship-total difference.",
        fill=F_NOTE)

    H(doc, "4.5 Stage E — fallback (REMOVED 2026-05-16)", 3)
    P(doc, "The optional fallback phase (boost MBQs, recompute REQ, re-run "
      "waterfall, demote Primary grids) has been removed. The orchestrator "
      "goes straight from Stage C waterfall + caps to Stage F (reflect). "
      "To surface stock the main pass left, raise the per-OPT-type MBQ cap "
      "sliders (rl_mbq_cap_pct / tbc_mbq_cap_pct / tbl_mbq_cap_pct) instead. "
      "Historical design preserved in fallback_archived.md.")

    H(doc, "4.6 Stage F — reflect & finalise", 3)
    Bullet(doc, "ALLOC_STATUS classified: ALLOCATED / PARTIAL / SKIPPED.")
    Bullet(doc, "SKIP_REASON / ALLOC_REASON populated.")
    Bullet(doc, "park_alloc + listing snapshots → ARS_*_PARKED (audit trail).")
    Bullet(doc, "OPT_STATUS, TBL_LISTED_DATE updated. New TBL holds written to "
                "ARS_NL_TBL_HOLD_TRACKING for next run.")

    # ----------------------------------- 5. Cap & gate rules with examples
    H(doc, "5. Cap & Gate Rules — Definitions and Live Examples", 2)

    H(doc, "5.1 PRI≥100 strict (RL / TBC)", 3)
    P(doc, "When pri_ct_check_rl=True, an OPT is skipped if PRI_CT_REM < 100. "
      "PRI_CT% = Σ H_<grid> / Σ GH_<grid> across Primary grids. Each band "
      "decrements H_REM as REQ is consumed. Strict mode forces the engine "
      "to also pin the store×MJ-level cap at 100% (see 5.5).")
    Box(doc, "Example", "Store HJ08, MJ_MBQ=2000, MJ_STK_TTL=1800, MJ_REQ=200. "
        "PRI strict ON → cap = 100% × 2000 − 1800 = 200. After RL ships 150, "
        "remaining = 50. TBC can ship at most 50. TBL gets whatever is left.",
        fill=F_NOTE)

    H(doc, "5.2 MBQ-cap relaxed (PRI strict OFF)", 3)
    P(doc, "When pri_ct_check_rl=False and rl_mbq_cap_pct=130 (for example), "
      "the cap is 130% × MJ_MBQ − MJ_STK_TTL. Same math, just higher ceiling.")
    Box(doc, "Example", "Same store as above. PRI OFF, cap_pct=130. "
        "cap = 130% × 2000 − 1800 = 800. RL+TBC+TBL combined can dispatch up "
        "to 800 units.", fill=F_NOTE)

    H(doc, "5.3 TBL MBQ-cap", 3)
    P(doc, "Independent dial tbl_mbq_cap_pct (default 100). Same formula. "
      "TBL is never under PRI strict — it always uses the % cap.")

    H(doc, "5.4 Sec-grid cap pre-gate (130%)", 3)
    P(doc, "Already covered in 4.4. Applies AFTER waterfall, only to "
      "Secondary grids, OPT-grain block/override.")

    H(doc, "5.5 Store×MJ live budget (the May-2026 fix)", 3)
    P(doc, "At the start of EVERY band, the engine rebuilds a per-WERKS budget "
      "from working_df['MJ_REQ_REM'] which is decremented after every prior "
      "band. Formula:")
    Code(doc, "budget[WERKS] = max(0, MJ_REQ_REM + (cap_pct − 100)/100 × MJ_MBQ)")
    Bullet(doc, "cap_pct = 100 (PRI strict)  → budget = MJ_REQ_REM")
    Bullet(doc, "cap_pct = 130 (relaxed)     → budget = MJ_REQ_REM + 30% × MJ_MBQ")
    Bullet(doc, "Inside a band, multiple OPTs at the same store compete for "
                "this budget in OPT_PRIORITY_RANK order — highest rank wins.")

    H(doc, "5.6 Pre-band PRI_CT_REM check", 3)
    P(doc, "Before each band, OPTs whose PRI_CT_REM < 100 are skipped for the "
      "active enforced types (TBL always; RL/TBC when flags on). Reason "
      "appended: SKIP_PRI_BROKEN(pri=NN.N).")

    H(doc, "5.7 Store-broken across types", 3)
    P(doc, "When ENABLE_STORE_BROKEN=True, if MJ_REQ_REM < 0.5 × ACS_D after "
      "a band, the store is dropped for the CURRENT type and ALL LATER types "
      "(it's saturated). Reason: SKIP_STORE_BROKEN(mj_rem=NN).")

    # ---------------------------------------------- 6. Toggle matrix
    H(doc, "6. Toggle Pattern Matrix", 2)
    P(doc, "Every combination of PRI strict (RL), PRI strict (TBC), "
      "apply_sec_cap_in_normal — and what each pattern means in practice. "
      "(Fallback was removed 2026-05-16 and is no longer in the matrix.)")
    Table(doc,
          ["PRI RL", "PRI TBC", "Sec-cap", "Outcome / when to use"],
          [
              ["ON",  "ON",  "ON",  "Strict baseline (recommended). Store×MJ ≤ 100%. Secondary grids policed at their per-grid % (default 130) with OPT_REQ-based override."],
              ["ON",  "ON",  "OFF", "Strict on primary, no secondary policing. Risk: heavy colour/MVGR skew."],
              ["ON",  "OFF", "ON",  "RL strict, TBC relaxed at tbc_mbq_cap_pct. Mixed enforcement — usually for diagnostics, NOT production."],
              ["OFF", "OFF", "ON",  "RL/TBC both at user-set cap (e.g. 110/130). Use when MBQ underestimates demand."],
              ["OFF", "OFF", "OFF", "Pure pool-limited dispatch. Will over-ship favoured stores. Avoid."],
          ], col_widths=[0.7, 0.7, 0.7, 5.3])
    Box(doc, "Recommendation",
        "Default daily run: PRI RL=ON, PRI TBC=ON, Sec-cap=ON, Fallback=OFF. "
        "Switch Fallback=ON when ship/REQ at aggregate drops below 90% AND "
        "warehouse has confirmed extra pool.", fill=F_NOTE)

    # ---------------------------------------------- 7. Worked example
    H(doc, "7. Worked Example — Store HU84 / MAJ_CAT M_TEES_HS / PRI strict", 2)
    P(doc, "Real data from session 20260514_220843_986 (pre-fix) to illustrate "
      "what the cap does.")
    Table(doc,
          ["Stage", "Value", "Note"],
          [
              ["MJ_MBQ", "1659", "Store-MAJ_CAT capacity"],
              ["MJ_STK_TTL", "1623", "Pre-run on-hand"],
              ["MJ_REQ", "36", "= MBQ − STK"],
              ["Pre-fix SHIP_RL", "24", ""],
              ["Pre-fix SHIP_TBC", "149", ""],
              ["Pre-fix SHIP_TBL", "0", ""],
              ["Pre-fix SHIP_TOTAL", "173", "≈ 480% of MJ_REQ — BUG"],
              ["Post-fix expected", "≤ 36", "Budget capped at MJ_REQ_REM"],
          ], col_widths=[2.0, 1.5, 3.0])

    doc.add_page_break()

    # ====================================================================
    # PART II — RECENT FIXES (May 2026)
    # ====================================================================
    H(doc, "Part II — Recent Fixes (May 2026)", 1)

    Box(doc, "Fix 1",
        "PRI≥100 strict was previously a per-OPT eligibility gate ONLY. It "
        "set rl_mbq_cap_pct=0 which disabled _build_mbq_budget entirely — so "
        "stores could ship 300%+ of MJ_MBQ. Now PRI strict pins the effective "
        "cap to 100%. File: rule_engine_pandas.py — `eff_rl_cap = 100.0 if "
        "pri_ct_check_rl else rl_mbq_cap_pct`.", fill=F_FIX,
        label_color=C_OK)

    Box(doc, "Fix 2",
        "Each OPT_TYPE used to get its own independent budget — RL, TBC, TBL "
        "each fresh 100%. The three stacked → 3× overshoot possible. Now the "
        "budget is shared across types via cumulative SHIP_QTY across ALL "
        "OPT_TYPEs at the WERKS.", fill=F_FIX, label_color=C_OK)

    Box(doc, "Fix 3",
        "Budget is rebuilt FRESH at the start of every band from the live "
        "working_df['MJ_REQ_REM'] (which _revalidate_after_band decrements "
        "after every band). New helper: _live_mbq_budget(). No more stale "
        "static dicts; single source of truth.", fill=F_FIX, label_color=C_OK)

    Box(doc, "Fix 4 (listing-side)",
        "Part 2.5 now creates a 4-column composite index "
        "(WERKS, MAJ_CAT, GEN_ART_NUMBER, CLR) which speeds up Parts 3.5a, "
        "3.5b, 3.5c, and 3.54 at 6M+ row scale.", fill=F_FIX, label_color=C_OK)

    Box(doc, "Fix 5 (listing-side)",
        "Part 3.5a Step 2 OR-WHERE clause was preventing index seek on "
        "ARS_CALC_ST_ART join. Rewritten to use CASE-in-SET so the planner "
        "can use the new composite index. ~5-10x speed-up at full-universe "
        "scale.", fill=F_FIX, label_color=C_OK)

    H(doc, "How to verify the fixes", 2)
    Bullet(doc, "Engine header log now reads: \"RL: PRI>=100 strict (MJ-cap "
                "100%) | TBC: PRI>=100 strict (MJ-cap 100%) | TBL: MBQ-cap "
                "100.0%\".")
    Bullet(doc, "Run: `python backend/scripts/excess_breakdown.py "
                "<session_id>` — confirm \"Stores where SHIP > MJ_REQ: 0/N\".")
    Bullet(doc, "Compare two sessions: `python backend/scripts/"
                "compare_sessions.py <A> <B>` shows OPT-level diffs.")

    doc.add_page_break()

    # ====================================================================
    # PART III — BRD
    # ====================================================================
    H(doc, "Part III — BRD: Existing vs Proposed", 1)

    H(doc, "Existing State — As Is", 2)
    Table(doc,
          ["Area", "Current behaviour", "Status"],
          [
              ["Listing build", "Sequential SQL parts 1–7 over ARS_LISTING. "
                                "Scales linearly with grid_rows.", "Works"],
              ["OPT_TYPE classification", "MIX → RL → TBC → TBL CASE. "
                                          "Stable.", "Works"],
              ["MIX aggregation", "Always groups by (WERKS, MAJ_CAT). "
                                  "max 1 MIX line per store×MAJ_CAT.", "Works"],
              ["PRI≥100 gate", "Per-OPT skip when PRI_CT_REM<100.", "Works"],
              ["Store×MJ cap (strict)", "Pinned to 100% via live MJ_REQ_REM "
                                        "(May 2026).", "Recently fixed — re-verify"],
              ["Sec-grid cap", "130% pre-gate post-waterfall with OPT_REQ "
                               "override.", "Works"],
              ["Fallback", "Optional. Boosts MBQ by 130% and re-runs.", "Works, rarely used"],
              ["Hold tracking", "ARS_NL_TBL_HOLD_TRACKING written on TBL "
                                "ship; read into RL_HOLD_QTY next run.", "Works"],
              ["Determinism", "Mostly deterministic; one known per-store "
                              "drift (HO15↔HO24 in May tests, 3 units).", "Open"],
              ["Performance",  "55 s for 209K-row listing, ~2 min for 6.3M "
                               "row full universe (post-index fix).",   "Acceptable"],
              ["Audit",        "ARS_*_PARKED snapshots every run.", "Works"],
              ["Reproducibility", "Same inputs ≠ bit-identical per-store "
                                  "output (pool tie-break drift).", "Improve"],
          ], col_widths=[2.0, 4.0, 1.5])

    H(doc, "Proposed Improvements — To Be", 2)
    Table(doc,
          ["Area", "Improvement", "Why / Impact"],
          [
              ["Determinism",
               "Add (ST_RANK, WERKS, OPT_PRIORITY_RANK) as the final tie-break "
               "in every pool-consumption sort. Eliminate the HO15↔HO24 drift.",
               "Bit-identical output across re-runs → trustworthy audit + "
               "fewer support questions."],
              ["Primary-grid cap parity",
               "Apply the same pre-gate logic to Primary grids (MJ_MACRO_MVGR, "
               "MJ_RNG_SEG) at OPT grain — like Sec-cap does for secondary.",
               "Stops single-grid skew within MJ. Today only Sec-grids are "
               "policed at OPT grain."],
              ["MJ_REQ_REM init",
               "Initialise MJ_REQ_REM in Part 7 (SQL) rather than relying on "
               "pandas null-fallback at line 1055. Saves an early-pass guard.",
               "Cleaner separation between Listing and Allocation."],
              ["UI cap settings",
               "Surface the 'effective cap %' in the UI (= live cap_pct for "
               "RL/TBC/TBL after PRI flag coercion). Today user has to read "
               "the log to see what cap actually fired.",
               "Removes ambiguity. Users currently confused why 'PRI strict' "
               "and 'MBQ-cap' are different toggles when they should be one "
               "dial."],
              ["Unified dial",
               "Replace pri_ct_check_rl + rl_mbq_cap_pct with one dial: "
               "'RL store×MJ cap %' that defaults to 100 (strict). Drop the "
               "boolean flag entirely.",
               "Single source of truth in API + UI. Cleaner code."],
              ["TBL hold visibility",
               "Surface RL_HOLD_QTY on the listing UI per OPT so planners "
               "can see why TBC/RL is firing 'against an empty pool'.",
               "Today the hold is invisible until you query the DB."],
              ["Pre-band cap pre-gate",
               "Apply the cap as a PRE-band filter (zero-budget WERKS get "
               "their rows skipped before pool calc) — saves CPU.",
               "Tiny speed-up (~5% in dense MAJ_CATs)."],
              ["Multi-MAJ_CAT parallelism",
               "Today inline-runs for ≤1 MAJ_CAT. Lower the threshold so "
               "more runs use the process pool. Or move pool-take to numba.",
               "10–30% Part-8 speed-up on multi-MAJ_CAT runs."],
              ["Schema deduplication",
               "ARS_LISTING_WORKING duplicates most ARS_LISTING columns. "
               "Either replace with VIEW or store delta only.",
               "DB storage cut, snapshot size cut. Maintenance burden lower."],
              ["Sec-cap fairness",
               "Sec-cap pre-gate today is all-or-nothing per OPT. Switch to "
               "partial-trim (reduce SHIP to just-under-cap) to avoid losing "
               "entire 24u/23u OPTs as in the May test.",
               "Smoother allocation; fewer customer-facing zeros for valid "
               "OPTs."],
              ["Audit lineage",
               "Currently ARS_ALLOC_PARKED has ALLOC_REASON / SKIP_REASON / "
               "FB_REASON. Add a JSON column 'audit_trail' with every "
               "round-level decision per row.",
               "Lets reviewers trace 'why did this row ship 4u' without "
               "guessing."],
              ["Configurable thresholds",
               "stock_threshold_pct (0.6), excess_multiplier (2.0), "
               "age_threshold (15) are passed as request fields. Surface "
               "them in a settings UI with per-MAJ_CAT overrides.",
               "Replenishment varies by category — flat thresholds are "
               "overly blunt."],
          ], col_widths=[1.8, 3.2, 2.5])

    H(doc, "Acceptance Criteria for the Fix Already Applied", 2)
    Bullet(doc, "All sessions with pri_ct_check_rl=True must report 0 stores "
                "where SHIP > MJ_REQ for RL+TBC combined.")
    Bullet(doc, "Aggregate ship / total-MJ_REQ ratio ≥ 90% on full-universe runs.")
    Bullet(doc, "Engine header log shows '(MJ-cap 100%)' when strict is on.")
    Bullet(doc, "HOLD_QTY across two identical re-runs matches to the unit.")
    Bullet(doc, "compare_sessions.py reports 0 OPT-grain mismatches between "
                "two identical-input sessions (this surfaces the remaining "
                "non-determinism).")

    doc.add_page_break()

    # ====================================================================
    # PART IV — PERFORMANCE & RESULT IMPROVEMENTS
    # ====================================================================
    H(doc, "Part IV — Performance & Result Improvements", 1)

    H(doc, "Database / SQL", 2)
    Bullet(doc, "Add composite index on ARS_CALC_ST_ART(ST_CD, MAJ_CAT, "
                "GEN_ART_NUMBER, CLR). Today 6.3M × ARS_CALC_ST_ART join in "
                "Part 3.5a is hash-join because the source side has no index "
                "on this combo.")
    Bullet(doc, "Add covering index on MASTER_GEN_ART_SALE(ST_CD, MAJ_CAT, "
                "GEN_ART_NUMBER, CLR) INCLUDE (SAL_PD) — Part 3.5b.")
    Bullet(doc, "Add covering index on MASTER_GEN_ART_AGE(ST_CD, MAJ_CAT, "
                "GEN_ART_NUMBER, CLR) INCLUDE (AGE) — Part 3.5c.")
    Bullet(doc, "Move ARS_LISTING from heap to clustered index on "
                "(WERKS, MAJ_CAT, GEN_ART_NUMBER, CLR). Today every UPDATE "
                "scans; clustered key would let SQL Server seek + locality.")
    Bullet(doc, "Replace per-grid Part 4 UPDATEs with a single MERGE that "
                "joins all grids in one statement (~7 grids × 6M-row UPDATE "
                "→ 1 × 6M-row pass).")

    H(doc, "Pandas / Allocation Engine", 2)
    Bullet(doc, "Replace working_df.iterrows() in _build_mbq_budget / "
                "_live_mbq_budget with vectorised .to_dict(orient='records'). "
                "iterrows is the slowest pandas idiom.")
    Bullet(doc, "Pre-compute pool_keys tuples ONCE in alloc_df and reuse — "
                "currently rebuilt in every _run_band and _snapshot_fnl_q_rem.")
    Bullet(doc, "Use numpy structured arrays for the pool-take cumulative "
                "window instead of pandas groupby+cumsum. ~3× faster in profile.")
    Bullet(doc, "Switch process-pool threshold from 1 MAJ_CAT to 4 — "
                "current inline-run wastes 3 workers on multi-MAJ_CAT runs.")
    Bullet(doc, "Use category dtype for WERKS / MAJ_CAT / VAR_ART / CLR in "
                "alloc_df. Memory and join speed up significantly.")

    H(doc, "Result Quality", 2)
    Bullet(doc, "Tie-break ties in OPT_PRIORITY_RANK by (ST_RANK, WERKS, "
                "GEN_ART_NUMBER, CLR) globally. Today some sort paths fall "
                "back to row index → drift between runs.")
    Bullet(doc, "Add a 'reserved-for-low-priority' pool slice when high-rank "
                "stores would consume >80% of a pool key — protects long-tail "
                "stores. Configurable per MAJ_CAT.")
    Bullet(doc, "Treat sec-cap-blocked OPTs as 'pending review' rather than "
                "silently SKIPPED. Surface them in a daily report so planners "
                "can override one-by-one.")
    Bullet(doc, "Add a 'min ship per OPT' threshold: if SHIP < min_ship, "
                "drop to 0 to avoid splintered shipments.")

    H(doc, "Monitoring & Audit", 2)
    Bullet(doc, "Persist step_timings into ARS_LISTING_SESSIONS — today only "
                "logged.")
    Bullet(doc, "Add ARS_RUN_STATS table: per-run aggregate of "
                "ship/REQ/hold/skip/override + count of each ALLOC_REASON.")
    Bullet(doc, "Wire a Grafana board on ARS_RUN_STATS for week-over-week "
                "trend.")
    Bullet(doc, "Generate a diff-report between two sessions automatically "
                "via the compare_sessions helper; surface in UI.")

    doc.add_page_break()

    # ====================================================================
    # PART V — OPEN ITEMS / REVIEW CHECKLIST
    # ====================================================================
    H(doc, "Part V — Open Items / Review Checklist", 1)
    P(doc, "Items that need confirmation from product / ops before they can "
      "be considered closed.")

    Table(doc,
          ["#", "Item", "Owner", "Status"],
          [
              ["1", "Re-run a full session with PRI strict ON post-fix; "
                    "confirm excess_breakdown reports 0 overshoots.",
               "Engineering", "Pending"],
              ["2", "Decide: should fallback mode also rebuild budgets from "
                    "live MJ_REQ_REM, or keep static-cap math? Currently "
                    "fallback uses a separate path.",
               "Product", "Open"],
              ["3", "Confirm tbl_mbq_cap_pct=100 is the intended TBL ceiling "
                    "even when fallback is ON (current behaviour lifts it).",
               "Product", "Open"],
              ["4", "HO15 ↔ HO24 3-unit drift across re-runs — accept as "
                    "harmless or treat as bug? See compare_sessions output.",
               "Engineering + Product", "Open"],
              ["5", "Sec-cap pre-gate: confirm 130% default and OPT_REQ ≥ "
                    "100% override threshold. Move to per-MAJ_CAT config?",
               "Product", "Open"],
              ["6", "Add unit tests for: (a) PRI strict cap, (b) shared "
                    "budget across types, (c) live MJ_REQ_REM rebuild.",
               "Engineering", "Pending"],
              ["7", "Documentation: developer doc (build_listing_alloc_doc.py) "
                    "vs this user manual — keep both, or merge?",
               "Engineering", "Open"],
              ["8", "User-facing labelling: 'PRI≥100 strict' is confusing. "
                    "Rename to 'Cap RL/TBC at 100% of MBQ'?",
               "Product / UX", "Open"],
              ["9", "Add a 'why this row' tooltip in the listing UI that "
                    "reads ALLOC_REASON + SKIP_REASON + FB_REASON.",
               "Frontend", "Pending"],
              ["10","Daily diff report between today's session and "
                    "yesterday's — surface unusual swings.",
               "Engineering", "Pending"],
          ], col_widths=[0.4, 4.8, 1.6, 1.0])

    H(doc, "Details Still Missing — Required for Sign-off", 2)
    Bullet(doc, "Business definition of 'service level' (target ship/REQ %) "
                "per category. Currently the engine has no target; it just "
                "ships what fits.")
    Bullet(doc, "Per-MAJ_CAT cap overrides (e.g. fashion = 110%, basics = "
                "100%). Today one global flag.")
    Bullet(doc, "Holiday / promotion override hooks — bursts not modelled.")
    Bullet(doc, "Cross-RDC pool sharing rules under cross-RDC mode — only "
                "scaffolded in code, never validated end-to-end.")
    Bullet(doc, "TBL hold expiry policy — today HOLD_REM stays open until "
                "manually closed; no max age.")
    Bullet(doc, "What happens to MIX rows in the next run — re-evaluated "
                "from scratch? Confirm with planners.")
    Bullet(doc, "Audit retention SLA for ARS_*_PARKED tables — today they "
                "grow unbounded.")
    Bullet(doc, "Roll-up reporting needs: per region / store-cluster / "
                "RDC — not currently exposed.")

    # ----- final note -----
    doc.add_paragraph()
    P(doc, "This manual is a living document. Re-run "
      "build_listing_alloc_manual.py after any rule change. Pair it with the "
      "developer doc (ARS_Listing_Allocation_Process.docx) for file:line "
      "references.", italic=True, color=C_MUTED, size=9)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
